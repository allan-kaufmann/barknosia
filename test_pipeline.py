"""
Unit-Tests für pipeline.py

Ausführung:
    python -m pytest test_pipeline.py -v

Jede Funktion, die mit test_ beginnt, wird von pytest automatisch gefunden und ausgeführt.
"""
import pytest
from docx import Document

from pipeline import (
    _clean_for_hidden,
    _clean_heading_text,
    normalize_heading,
    normalize_heading_levels,
    parse_sections,
    parse_qa_response,
    serialize_qa_items,
    rework_partial_answers,
    _find_para_by_quote,
    _normalize_quote,
    _advance_counter,
    _prefix_chapter_number,
    _is_skip_heading,
    split_into_level1_chapters,
    add_formatted_text,
    process_markdown_to_docx,
    _strip_kontrollliste,
    _compress_heading_levels,
    _fix_concatenated_stats_cells,
    _drop_empty_columns,
    _hide_paragraph,
    add_markdown_table_to_doc,
    _set_cell_text,
    _strip_ocr_y_prefix,
    build_interleaved_word_document,
    _split_at_level2,
    _split_at_level,
    _detect_chapter_level,
    _find_sublevel,
    _group_into_chunks,
    _summarize_chapter_by_sections,
    _summarize_single_chapter,
    generate_summary_by_chapter,
    _is_box_heading,
    _split_paragraphs,
    repair_box_structure,
    _classify_box_boundaries,
    _rebase_chapter_number,
    extract_chapter,
    strip_front_matter,
)
import pipeline


# ---------------------------------------------------------------------------
# Gruppe 1: _clean_for_hidden  (Fix: Bilder werden nicht entfernt)
# ---------------------------------------------------------------------------

def test_clean_for_hidden_preserves_images():
    result = _clean_for_hidden("Bild: ![Abbildung 1](figures/img.png) im Text.")
    assert "![Abbildung 1](figures/img.png)" in result


def test_clean_for_hidden_removes_text_links():
    result = _clean_for_hidden("[Quelle](https://example.com)")
    assert result == "Quelle"


def test_clean_for_hidden_text_link_does_not_affect_image():
    md = "[Link](https://url.com) und ![Bild](bild.jpeg)"
    result = _clean_for_hidden(md)
    assert "Link" in result           # Link-Text bleibt
    assert "https://url.com" not in result  # URL weg
    assert "![Bild](bild.jpeg)" in result   # Bild unverändert


def test_clean_for_hidden_removes_html_tags():
    result = _clean_for_hidden('<span id="page-12">Text</span>')
    assert "<span" not in result
    assert "Text" in result


def test_clean_for_hidden_removes_standalone_page_refs():
    result = _clean_for_hidden("[1]\n\nText dahinter")
    assert "[1]" not in result
    assert "Text dahinter" in result


# ---------------------------------------------------------------------------
# Gruppe 2: _clean_heading_text + normalize_heading  (Fix: Lookup-Key)
# ---------------------------------------------------------------------------

def test_clean_heading_removes_html():
    raw = '<span id="page-12-0"></span>**1.2 Definition**'
    assert _clean_heading_text(raw) == "1.2 Definition"


def test_clean_heading_removes_markdown_link():
    raw = '**1.2 [Definition](https://dorsch.hogrefe.com/stichwort/definition)**'
    assert _clean_heading_text(raw) == "1.2 Definition"


def test_clean_heading_combined_html_and_link():
    raw = '<span id="page-12-0"></span><span id="page-12-1"></span>**1.2 [Definition](https://dorsch.de)**'
    assert _clean_heading_text(raw) == "1.2 Definition"


def test_normalize_heading_strips_bold_and_lowercases():
    assert normalize_heading("**1.2 Definition**") == "1.2 definition"


def test_normalize_heading_strips_italic():
    assert normalize_heading("*Abschnitt*") == "abschnitt"


def test_lookup_key_ocr_heading_matches_summary_heading():
    """
    Kerntest für den Empty-Chapter-Fix:
    Ein OCR-Heading mit HTML und Links muss denselben lookup_key ergeben
    wie der saubere Summary-Heading.
    """
    ocr_heading     = '<span id="p">**1.2 [Definition](https://dorsch.de)**'
    summary_heading = "**1.2 Definition**"
    assert normalize_heading(_clean_heading_text(ocr_heading)) == normalize_heading(summary_heading)


# ---------------------------------------------------------------------------
# Gruppe 3: add_formatted_text  (Fixes: italic, triple asterisk)
# ---------------------------------------------------------------------------

def test_add_formatted_text_bold():
    para = Document().add_paragraph()
    add_formatted_text(para, "**fett**")
    bold_runs = [r for r in para.runs if r.bold]
    assert len(bold_runs) == 1
    assert bold_runs[0].text == "fett"


def test_add_formatted_text_italic():
    para = Document().add_paragraph()
    add_formatted_text(para, "*kursiv*")
    italic_runs = [r for r in para.runs if r.italic]
    assert len(italic_runs) == 1
    assert italic_runs[0].text == "kursiv"


def test_add_formatted_text_bold_italic_triple_asterisk():
    """Fix: ***Ausbildung:*** wurde früher als Literal-Text mit Sternchen ausgegeben."""
    para = Document().add_paragraph()
    add_formatted_text(para, "***Ausbildung:***")
    bi_runs = [r for r in para.runs if r.bold and r.italic]
    assert len(bi_runs) == 1
    assert bi_runs[0].text == "Ausbildung:"
    assert "*" not in bi_runs[0].text


def test_add_formatted_text_plain():
    para = Document().add_paragraph()
    add_formatted_text(para, "normaler Text")
    assert para.runs[0].text == "normaler Text"
    assert not para.runs[0].bold
    assert not para.runs[0].italic


def test_add_formatted_text_mixed():
    para = Document().add_paragraph()
    add_formatted_text(para, "Vor ***bold+italic*** Nach")
    texts = [r.text for r in para.runs]
    assert "Vor " in texts
    assert "bold+italic" in texts
    assert " Nach" in texts
    bi_runs = [r for r in para.runs if r.bold and r.italic]
    assert len(bi_runs) == 1
    assert bi_runs[0].text == "bold+italic"


def test_add_formatted_text_no_asterisk_in_output():
    """Kein Sternchen darf als sichtbarer Text im Dokument landen."""
    para = Document().add_paragraph()
    add_formatted_text(para, "***fett-kursiv*** und **nur-fett** und *nur-kursiv*")
    for run in para.runs:
        assert "*" not in run.text, f"Sternchen in Run: '{run.text}'"


# ---------------------------------------------------------------------------
# Gruppe 4: process_markdown_to_docx – Bullet-Handling (mehrere Fixes)
# ---------------------------------------------------------------------------

def test_bullet_no_leading_spaces():
    """Fix: '*   Text' (Stern + 3 Leerzeichen) erzeugte '  Text' mit führenden Spaces."""
    doc = Document()
    process_markdown_to_docx(doc, "*   Lernen in Unternehmen")
    bullets = [p for p in doc.paragraphs if p.style.name == 'List Bullet']
    assert len(bullets) == 1
    assert bullets[0].text == "Lernen in Unternehmen"


def test_bullet_single_space_after_star():
    doc = Document()
    process_markdown_to_docx(doc, "* Einfacher Bullet")
    bullets = [p for p in doc.paragraphs if p.style.name == 'List Bullet']
    assert bullets[0].text == "Einfacher Bullet"


def test_empty_bullet_skipped():
    """Fix: '*   ' ohne Text erzeugte leeren Bullet-Absatz."""
    doc = Document()
    process_markdown_to_docx(doc, "*   \n* Inhalt")
    bullets = [p for p in doc.paragraphs if p.style.name == 'List Bullet']
    assert len(bullets) == 1
    assert bullets[0].text == "Inhalt"


def test_numbered_sub_item_as_sub_bullet():
    """Fix: '    1.  Text' landete als Plaintext statt eingerücktem Sub-Bullet."""
    doc = Document()
    process_markdown_to_docx(doc, "*   Liste:\n    1.  Erster Punkt\n    2.  Zweiter Punkt")
    sub_bullets = [p for p in doc.paragraphs if p.style.name in ('List Bullet 2', 'List Bullet 3')]
    assert len(sub_bullets) == 2
    assert sub_bullets[0].text == "Erster Punkt"
    assert sub_bullets[1].text == "Zweiter Punkt"


def test_sub_bullet_with_asterisk():
    doc = Document()
    process_markdown_to_docx(doc, "*   Eltern\n    *   Kind-Bullet")
    sub = [p for p in doc.paragraphs if p.style.name in ('List Bullet 2', 'List Bullet 3')]
    assert len(sub) == 1
    assert sub[0].text == "Kind-Bullet"


def test_triple_asterisk_in_bullet():
    """Fix: '***Ausbildung:*** Text' in Bullet durfte keine Sternchen zeigen."""
    doc = Document()
    process_markdown_to_docx(doc, "*   ***Ausbildung:*** Curricular organisierte Aktivitäten")
    bullets = [p for p in doc.paragraphs if p.style.name == 'List Bullet']
    assert len(bullets) == 1
    for run in bullets[0].runs:
        assert "*" not in run.text


# ---------------------------------------------------------------------------
# Gruppe 5: normalize_heading_levels + parse_sections
# ---------------------------------------------------------------------------

def test_normalize_level1_to_h1():
    result = normalize_heading_levels("# 1 Einleitung\nText")
    assert result.startswith("# 1 Einleitung")


def test_normalize_level2_to_h2():
    result = normalize_heading_levels("# 1.1 Abschnitt\nText")
    assert result.startswith("## 1.1 Abschnitt")


def test_normalize_level3_to_h3():
    result = normalize_heading_levels("# 1.1.1 Unterabschnitt\nText")
    assert result.startswith("### 1.1.1 Unterabschnitt")


def test_normalize_non_numbered_unchanged():
    result = normalize_heading_levels("# Einleitung ohne Nummer\nText")
    assert result.startswith("# Einleitung ohne Nummer")


def test_parse_sections_finds_headings():
    md = "## 2 Kapitel\nErster Absatz.\n\n## 3 Folgekapitel\nZweiter Absatz."
    sections = parse_sections(md)
    heads = [s['heading'] for s in sections if s['heading'] != '__preamble__']
    assert "2 Kapitel" in heads
    assert "3 Folgekapitel" in heads


def test_parse_sections_preamble():
    md = "Einleitungstext\n## 1 Kapitel\nInhalt"
    sections = parse_sections(md)
    assert sections[0]['heading'] == '__preamble__'
    assert "Einleitungstext" in sections[0]['body']


def test_parse_sections_body_content():
    md = "## 1 Kapitel\nZeile 1\nZeile 2"
    sections = parse_sections(md)
    kap = next(s for s in sections if s['heading'] != '__preamble__')
    assert "Zeile 1" in kap['body']
    assert "Zeile 2" in kap['body']


# ---------------------------------------------------------------------------
# Bonus: Weitere pure Helfer
# ---------------------------------------------------------------------------

def test_advance_counter_increments_level():
    counters = [0] * 9
    result = _advance_counter(counters, 1)
    assert result == "1"
    assert counters[0] == 1


def test_advance_counter_resets_deeper_levels():
    counters = [3, 2, 1, 0, 0, 0, 0, 0, 0]
    _advance_counter(counters, 2)
    assert counters[1] == 3
    assert counters[2] == 0


def test_advance_counter_skips_zero_segments():
    counters = [4, 0, 0, 0, 0, 0, 0, 0, 0]
    _advance_counter(counters, 2)
    result = _advance_counter(counters, 2)
    assert '.' in result
    assert '0' not in result.split('.')[0]


def test_prefix_chapter_number_numbered():
    assert _prefix_chapter_number("1 EINLEITUNG", "4.2") == "4.2.1 EINLEITUNG"


def test_prefix_chapter_number_numbered_sub():
    assert _prefix_chapter_number("1.1 Abschnitt", "4.2") == "4.2.1.1 Abschnitt"


def test_prefix_chapter_number_non_numbered_unchanged():
    assert _prefix_chapter_number("Einleitung ohne Zahl", "4.2") == "Einleitung ohne Zahl"


def test_is_skip_heading_referenzen():
    assert _is_skip_heading("Referenzen") is True
    assert _is_skip_heading("referenzen") is True


def test_is_skip_heading_literaturverzeichnis():
    assert _is_skip_heading("Literaturverzeichnis") is True


def test_is_skip_heading_normal_chapter():
    assert _is_skip_heading("1 Einleitung") is False
    assert _is_skip_heading("Methoden") is False


def test_parse_qa_response_basic():
    qa = (
        "Frage 1\n"
        "Antwort: Die Antwort auf Frage 1.\n"
        "Textgrundlage: 1.1 Abschnitt\n"
        "Schlüsselbegriffe: Begriff A, Begriff B\n"
        "Abdeckung: vollständig\n\n"
        "Frage 2\n"
        "Antwort: Die Antwort auf Frage 2.\n"
        "Textgrundlage: 1.2 Unterabschnitt\n"
        "Schlüsselbegriffe: Begriff C\n"
        "Abdeckung: teilweise\n"
    )
    items = parse_qa_response(qa)
    assert len(items) == 2
    assert items[0]['num'] == 1
    assert "Frage 1" in items[0]['antwort'] or "Antwort" in items[0]['antwort']
    assert items[0]['textgrundlage'] == "1.1 Abschnitt"
    assert items[1]['num'] == 2
    assert items[1]['abdeckung'] == "teilweise"


def test_parse_qa_response_parses_beleg_and_bold_labels():
    qa = (
        "**Frage 9**\n"
        "**Antwort:** Kompetenzen umfassen Wissen und Fertigkeiten.\n"
        "**Textgrundlage:** 2.1 Wozu wird gelernt?\n"
        "**Schlüsselbegriffe:** Wissen, Fertigkeiten\n"
        "**Beleg:** Kompetenzen umfassen alle Wissensbestände\n"
        "**Abdeckung:** vollständig\n"
    )
    items = parse_qa_response(qa)
    assert len(items) == 1
    assert items[0]['beleg'] == "Kompetenzen umfassen alle Wissensbestände"
    assert items[0]['textgrundlage'] == "2.1 Wozu wird gelernt?"
    assert items[0]['antwort'].startswith("Kompetenzen umfassen Wissen")


def test_serialize_qa_items_roundtrip():
    qa = (
        "Frage 1\n"
        "Antwort: Antwort eins.\n"
        "Textgrundlage: 1.1 Abschnitt\n"
        "Schlüsselbegriffe: A, B\n"
        "Beleg: ein wörtliches Zitat hier\n"
        "Abdeckung: vollständig\n"
    )
    items = parse_qa_response(qa)
    reparsed = parse_qa_response(serialize_qa_items(items))
    assert reparsed[0]['num'] == 1
    assert reparsed[0]['beleg'] == "ein wörtliches Zitat hier"
    assert reparsed[0]['abdeckung'] == "vollständig"
    assert reparsed[0]['textgrundlage'] == "1.1 Abschnitt"


def test_normalize_quote_strips_markdown_and_case():
    assert _normalize_quote("**Hallo**  Welt") == "hallo welt"
    assert _normalize_quote('„Zitat"') == "zitat"


def test_find_para_by_quote_matches_prefix():
    class FakePara:
        def __init__(self, text):
            self.text = text
    paras = [FakePara("Ein einleitender Satz ohne Bezug."),
             FakePara("Kompetenzen umfassen alle Wissensbestände und Fertigkeiten.")]
    hit = _find_para_by_quote(paras, "Kompetenzen umfassen alle Wissensbestände")
    assert hit is paras[1]
    assert _find_para_by_quote(paras, "etwas völlig anderes das nicht vorkommt") is None
    assert _find_para_by_quote(paras, "kurz") is None  # zu kurz


def test_rework_partial_answers_supplements_partial(monkeypatch):
    qa = (
        "Frage 1\n"
        "Antwort: Der Zyklus hat acht Schritte.\n"
        "Textgrundlage: 2.2 Modelle\n"
        "Schlüsselbegriffe: Zyklus\n"
        "Beleg: Der Zyklus hat acht Schritte\n"
        "Abdeckung: teilweise\n\n"
        "Frage 2\n"
        "Antwort: Vollständige Antwort.\n"
        "Textgrundlage: 2.1 Grundlagen\n"
        "Schlüsselbegriffe: X\n"
        "Beleg: Vollständige Antwort\n"
        "Abdeckung: vollständig\n"
    )
    working = "## 2.2 Modelle\nDie acht Schritte sind A, B, C, D, E, F, G und H.\n"

    class FakeResp:
        text = "Die acht Schritte sind A bis H."

    calls = []
    def fake_call(**kwargs):
        calls.append(kwargs)
        return FakeResp()
    monkeypatch.setattr('pipeline.call_gemini_with_retry', fake_call)

    new_qa, supplements = rework_partial_answers(qa, working)
    items = parse_qa_response(new_qa)
    by_num = {it['num']: it for it in items}
    # Nur die teilweise-Frage wurde nachbearbeitet → genau ein Gemini-Call
    assert len(calls) == 1
    assert by_num[1]['abdeckung'] == "vollständig durch Nachbearbeitung"
    assert "A bis H" in by_num[1]['antwort']
    assert by_num[2]['abdeckung'] == "vollständig"  # unverändert
    assert any("A bis H" in s for v in supplements.values() for s in v)


def test_rework_partial_answers_keeps_item_when_no_supplement(monkeypatch):
    qa = (
        "Frage 1\n"
        "Antwort: Teilantwort.\n"
        "Textgrundlage: 9.9 Unbekannt\n"
        "Schlüsselbegriffe: X\n"
        "Beleg: Teilantwort\n"
        "Abdeckung: nicht enthalten\n"
    )

    class FakeResp:
        text = "KEINE ERGÄNZUNG MÖGLICH"
    monkeypatch.setattr('pipeline.call_gemini_with_retry', lambda **k: FakeResp())

    new_qa, supplements = rework_partial_answers(qa, "## Anderes\nIrrelevant.\n")
    items = parse_qa_response(new_qa)
    assert items[0]['abdeckung'] == "nicht enthalten"
    assert supplements == {}


def test_split_into_level1_chapters():
    md = (
        "## 1 Einleitung\nText Kap 1\n"
        "### 1.1 Unterkapitel\nText 1.1\n"
        "## 2 Methoden\nText Kap 2\n"
    )
    chapters = split_into_level1_chapters(md)
    assert len(chapters) == 2
    headings = [c['heading'] for c in chapters]
    assert any("1 Einleitung" in h for h in headings)
    assert any("2 Methoden" in h for h in headings)


# ---------------------------------------------------------------------------
# Gruppe 7: _strip_kontrollliste
# ---------------------------------------------------------------------------

def test_strip_kontrollliste_removes_block_with_hr():
    text = "Übersetzungstext.\n\n---\n**Kontrollliste**\n- Absätze: 5\n- Hinweise: Keine.\n\n## Kapitel 2\nWeiterer Text."
    result = _strip_kontrollliste(text)
    assert "Kontrollliste" not in result
    assert "Kapitel 2" in result
    assert "Übersetzungstext" in result


def test_strip_kontrollliste_no_hr():
    text = "Text.\n\n**Kontrollliste**\n- Absätze: 3\n- Hinweise: Keine.\n\n## Nächstes Kapitel\nText."
    result = _strip_kontrollliste(text)
    assert "Kontrollliste" not in result
    assert "Nächstes Kapitel" in result


def test_strip_kontrollliste_no_kontrollliste():
    text = "# Kapitel\n\nNormaler Text ohne Kontrollliste."
    result = _strip_kontrollliste(text)
    assert result == text


# ---------------------------------------------------------------------------
# Gruppe 8: _compress_heading_levels
# ---------------------------------------------------------------------------

def test_compress_heading_levels_siblings():
    """H1 gefolgt von drei H4-Geschwistern → H1, H2, H2, H2."""
    md = "# Titel\n\n#### Sub A\n\n#### Sub B\n\n#### Sub C\n"
    result = _compress_heading_levels(md)
    lines = [l for l in result.split('\n') if l.startswith('#')]
    assert lines[0] == '# Titel'
    assert lines[1] == '## Sub A'
    assert lines[2] == '## Sub B'
    assert lines[3] == '## Sub C'


def test_compress_heading_levels_no_change_needed():
    """Korrekte Hierarchie bleibt unverändert."""
    md = "# H1\n\n## H2\n\n### H3\n"
    result = _compress_heading_levels(md)
    assert result == md


def test_compress_heading_levels_resets_on_new_h1():
    """Nach einem neuen H1 wird die Mapping-Logik neu berechnet."""
    md = "# Abschnitt A\n\n#### Sub A1\n\n# Abschnitt B\n\n#### Sub B1\n"
    result = _compress_heading_levels(md)
    lines = [l for l in result.split('\n') if l.startswith('#')]
    assert lines[0] == '# Abschnitt A'
    assert lines[1] == '## Sub A1'
    assert lines[2] == '# Abschnitt B'
    assert lines[3] == '## Sub B1'


# ---------------------------------------------------------------------------
# Gruppe 9: _is_skip_heading Erweiterungen
# ---------------------------------------------------------------------------

def test_is_skip_heading_history():
    assert _is_skip_heading("History") is True
    assert _is_skip_heading("Historie") is True


def test_is_skip_heading_author_email():
    assert _is_skip_heading("Timo Kortsch t.kortsch@tu-bs.de") is True


def test_is_skip_heading_msc():
    assert _is_skip_heading("Timo Kortsch, M.Sc. Prof. Dr. Simone Kauffeld") is True


def test_is_skip_heading_normal_not_skipped():
    assert _is_skip_heading("Theoretischer Hintergrund") is False


# ---------------------------------------------------------------------------
# Gruppe 10: process_markdown_to_docx – HTML-Stripping
# ---------------------------------------------------------------------------

def test_process_strips_sup_tags():
    doc = Document()
    process_markdown_to_docx(doc, "Text mit Fußnote<sup>1</sup> hier.")
    texts = [p.text for p in doc.paragraphs]
    combined = ' '.join(texts)
    assert '<sup>' not in combined
    assert 'Text mit Fußnote' in combined
    assert '1' in combined


def test_process_strips_span_tags():
    doc = Document()
    process_markdown_to_docx(doc, "Ein <span id='x'>wichtiger</span> Begriff.")
    texts = ' '.join(p.text for p in doc.paragraphs)
    assert '<span' not in texts
    assert 'wichtiger' in texts


# ---------------------------------------------------------------------------
# Gruppe 11: Bildfilterung in process_markdown_to_docx
# ---------------------------------------------------------------------------

def test_image_skip_images_flag(tmp_path):
    """skip_images=True: Bildzeile wird komplett ignoriert, kein Platzhalter."""
    doc = Document()
    process_markdown_to_docx(doc, "![Logo](logo.png)", base_path=str(tmp_path), skip_images=True)
    texts = ' '.join(p.text for p in doc.paragraphs)
    assert 'logo.png' not in texts
    assert 'Bild nicht' not in texts


def test_image_skip_images_false_adds_placeholder(tmp_path):
    """skip_images=False (Standard): fehlendes Bild → Platzhalter-Text."""
    doc = Document()
    process_markdown_to_docx(doc, "![Abb](missing.png)", base_path=str(tmp_path), skip_images=False)
    texts = ' '.join(p.text for p in doc.paragraphs)
    assert 'missing.png' in texts


# ---------------------------------------------------------------------------
# Gruppe 12: _fix_concatenated_stats_cells
# ---------------------------------------------------------------------------

def test_fix_concatenated_stats_cells_left_right_empty():
    """Werte links und rechts leer → auf drei Spalten verteilen."""
    rows = [['Item', 'Übersetzung', '', '3.44 1.00 .49', '']]
    result = _fix_concatenated_stats_cells(rows)
    assert result[0][2] == '3.44'
    assert result[0][3] == '1.00'
    assert result[0][4] == '.49'


def test_fix_concatenated_stats_cells_right_empty():
    """Werte nur rechts leer → auf rechte drei Spalten verteilen."""
    rows = [['Item', '3.44 1.00 .49', '', '']]
    result = _fix_concatenated_stats_cells(rows)
    assert result[0][1] == '3.44'
    assert result[0][2] == '1.00'
    assert result[0][3] == '.49'


def test_fix_concatenated_stats_cells_no_change_needed():
    """Zeile ohne OCR-Artefakt bleibt unverändert."""
    rows = [['Item', 'Übersetzung', '3.44', '1.00', '.49']]
    result = _fix_concatenated_stats_cells(rows)
    assert result[0] == ['Item', 'Übersetzung', '3.44', '1.00', '.49']


# ---------------------------------------------------------------------------
# Gruppe 13: normalize_heading_levels – beliebige Tiefe
# ---------------------------------------------------------------------------

def test_normalize_4level_depth():
    """4.1.1.1 (3 Punkte) → H4 (####)."""
    result = normalize_heading_levels("# 4.1.1.1 Tief\nText")
    assert result.startswith("#### 4.1.1.1 Tief")


def test_normalize_5level_depth():
    """4 Punkte → H5 (#####)."""
    result = normalize_heading_levels("# 1.2.3.4.5 Sehr tief\nText")
    assert result.startswith("##### 1.2.3.4.5 Sehr tief")


def test_normalize_tail_number_reordered():
    """'#### Titel 7.2.2' (Nummer am Ende, 2 Punkte) → '### 7.2.2 Titel'."""
    result = normalize_heading_levels("#### Motivation als Zielverfolgung 7.2.2\nText")
    assert result.startswith("### 7.2.2 Motivation als Zielverfolgung"), (
        f"Tail-Normalisierung fehlgeschlagen: {result!r}"
    )


def test_normalize_tail_number_false_positive_year():
    """'Studie aus dem Jahr 2023' (kein Punkt in Zahl) bleibt unverändert."""
    result = normalize_heading_levels("#### Studie aus dem Jahr 2023\nText")
    assert result.startswith("#### Studie aus dem Jahr 2023"), (
        f"Jahreszahl darf nicht als Kapitelnummer behandelt werden: {result!r}"
    )


def test_normalize_tail_number_false_positive_one_dot():
    """'Abbildung 18.4' (nur 1 Punkt) bleibt unverändert — Mindest 2 Punkte erforderlich."""
    result = normalize_heading_levels("#### Abbildung 18.4\nText")
    assert result.startswith("#### Abbildung 18.4"), (
        f"Einfache Abbildungsnummer darf nicht normalisiert werden: {result!r}"
    )


# ---------------------------------------------------------------------------
# Gruppe 14: _compress_heading_levels – nummerierte Headings unverändert
# ---------------------------------------------------------------------------

def test_compress_skips_numbered():
    """Nummerierte Headings (nach normalize) werden von _compress nicht verändert."""
    md = "## 4 Kapitel\n### 4.1 Sub\n#### 4.1.1 SubSub\n### 4.2 Sub2\n"
    result = _compress_heading_levels(md)
    assert "## 4 Kapitel" in result
    assert "### 4.1 Sub" in result
    assert "#### 4.1.1 SubSub" in result
    assert "### 4.2 Sub2" in result


# ---------------------------------------------------------------------------
# Gruppe 15: split_into_level1_chapters – adaptiver Split
# ---------------------------------------------------------------------------

def test_split_chapters_adaptive_h3():
    """Extrahiertes Kapitel: einziges ## → Unterkapitel auf ### werden gesplittet."""
    md = (
        "## 4 Vorgehen\nText\n"
        "### 4.1 Bedarfserhebung\nText 4.1\n"
        "### 4.2 Konzeptentwicklung\nText 4.2\n"
    )
    chapters = split_into_level1_chapters(md)
    headings = [c['heading'] for c in chapters]
    assert len(chapters) == 2
    assert any("4.1" in h for h in headings)
    assert any("4.2" in h for h in headings)


def test_split_chapters_multi_top_stays_at_h2():
    """Mehrere ## Kapitel → weiterhin an ## splitten, nicht tiefer."""
    md = (
        "## 1 Einleitung\nText Kap 1\n"
        "### 1.1 Unterkapitel\nText 1.1\n"
        "## 2 Methoden\nText Kap 2\n"
    )
    chapters = split_into_level1_chapters(md)
    assert len(chapters) == 2
    headings = [c['heading'] for c in chapters]
    assert any("1 Einleitung" in h for h in headings)
    assert any("2 Methoden" in h for h in headings)


# ---------------------------------------------------------------------------
# Gruppe 16: _drop_empty_columns
# ---------------------------------------------------------------------------

def test_drop_empty_columns_removes_blank():
    """Spalte, die in allen Zeilen leer ist, wird entfernt."""
    rows = [['A', 'B', ''], ['1', '2', ''], ['3', '4', '']]
    result = _drop_empty_columns(rows)
    assert all(len(r) == 2 for r in result)
    assert result[0] == ['A', 'B']


def test_drop_empty_columns_keeps_nonempty():
    """Spalte mit mindestens einem nicht-leeren Wert bleibt erhalten."""
    rows = [['A', 'B', 'C'], ['1', '', '3']]
    result = _drop_empty_columns(rows)
    assert all(len(r) == 3 for r in result)


def test_drop_empty_columns_empty_input():
    assert _drop_empty_columns([]) == []


# ---------------------------------------------------------------------------
# Gruppe 17: _hide_paragraph – Aufzählungszeichen ausblenden
# ---------------------------------------------------------------------------

def test_hide_paragraph_hides_bullet():
    """_hide_paragraph setzt w:vanish in pPr/rPr, um das Bullet-Zeichen auszublenden."""
    from docx.oxml.ns import qn as _qn
    doc = Document()
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Bullet Text")
    _hide_paragraph(p)
    pPr = p._p.find(_qn('w:pPr'))
    assert pPr is not None, "pPr fehlt"
    rPr = pPr.find(_qn('w:rPr'))
    assert rPr is not None, "rPr in pPr fehlt"
    vanish = rPr.find(_qn('w:vanish'))
    assert vanish is not None, "w:vanish fehlt – Bullet-Zeichen wird nicht ausgeblendet"


# ---------------------------------------------------------------------------
# Gruppe 18: _hide_paragraph – w:numPr wird entfernt
# ---------------------------------------------------------------------------

def test_hide_paragraph_removes_numpr():
    """_hide_paragraph entfernt w:numPr, damit das List-Label (• + Tab) nicht sichtbar bleibt."""
    from docx.oxml.ns import qn as _qn
    doc = Document()
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Bullet Text")
    # List Bullet-Style setzt w:numPr → muss nach _hide_paragraph weg sein
    _hide_paragraph(p)
    pPr = p._p.find(_qn('w:pPr'))
    assert pPr is not None, "pPr fehlt"
    numPr = pPr.find(_qn('w:numPr'))
    assert numPr is None, "w:numPr ist noch vorhanden – Bullet-Label bleibt sichtbar"


# ---------------------------------------------------------------------------
# Gruppe 19: add_markdown_table_to_doc – Titel-Zeile erkennen
# ---------------------------------------------------------------------------

def test_table_title_row_merged():
    """Erste Zeile mit nur 1 Non-Empty-Cell → wird zur gemergten Titelzeile; Zeile 1 bekommt Shading."""
    from docx.oxml.ns import qn as _qn
    doc = Document()
    table_lines = [
        "| Lange Fragestellung |  |  |",
        "|---------------------|--|--|",
        "| Spalte A | Spalte B | Spalte C |",
        "| Wert 1   | Wert 2   | Wert 3   |",
    ]
    add_markdown_table_to_doc(doc, table_lines)
    # Zuletzt eingefügte Tabelle
    tbl = doc.tables[-1]
    # Zeile 0 soll nach merge nur noch 1 Zelle haben (alle zusammengeführt)
    row0_cells = tbl.rows[0].cells
    assert row0_cells[0] is row0_cells[-1], "Titelzeile wurde nicht zusammengeführt (Zellen nicht identisch)"
    # Zeile 1 (echter Header) soll Shading haben
    tc1 = tbl.rows[1].cells[0]._tc
    tcPr = tc1.find(_qn('w:tcPr'))
    assert tcPr is not None, "tcPr in Header-Zeile fehlt"
    shd = tcPr.find(_qn('w:shd'))
    assert shd is not None, "w:shd fehlt – Header-Zeile hat kein Shading"


def test_table_normal_header_unchanged():
    """Tabelle mit mehreren Non-Empty-Cells in Zeile 0 → normale Header-Logik (Zeile 0 = Header)."""
    from docx.oxml.ns import qn as _qn
    doc = Document()
    table_lines = [
        "| Name | Wert |",
        "|------|------|",
        "| A    | 1    |",
    ]
    add_markdown_table_to_doc(doc, table_lines)
    tbl = doc.tables[-1]
    # Zeile 0 soll Shading haben (normaler Header)
    tc0 = tbl.rows[0].cells[0]._tc
    tcPr = tc0.find(_qn('w:tcPr'))
    assert tcPr is not None
    shd = tcPr.find(_qn('w:shd'))
    assert shd is not None, "w:shd fehlt in Zeile 0 – normaler Header nicht erkannt"


# ---------------------------------------------------------------------------
# Gruppe 20: _set_cell_text – Bullet-Concatenation
# ---------------------------------------------------------------------------

def test_set_cell_text_splits_bullets():
    """•A•B•C in einer Zelle → 3 separate Paragraphen."""
    doc = Document()
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    _set_cell_text(cell, "•Wiedererkennen•Benennen•Abrufen")
    texts = [p.text.strip() for p in cell.paragraphs if p.text.strip()]
    assert len(texts) == 3, f"Erwartet 3 Paragraphen, erhalten: {texts}"
    assert "Wiedererkennen" in texts
    assert "Benennen" in texts
    assert "Abrufen" in texts


def test_set_cell_text_single_bullet_unchanged():
    """Einzelnes • am Anfang (kein Concat) → kein Split."""
    doc = Document()
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    _set_cell_text(cell, "•Eintrag ohne Split")
    texts = [p.text.strip() for p in cell.paragraphs if p.text.strip()]
    assert len(texts) == 1, f"Kein Split erwartet, erhalten: {texts}"


# ---------------------------------------------------------------------------
# Gruppe 21: process_markdown_to_docx – headings_as_bold
# ---------------------------------------------------------------------------

def test_process_markdown_headings_as_bold():
    """headings_as_bold=True → ### Titel wird als Normal-Paragraph mit bold=True eingefügt."""
    doc = Document()
    process_markdown_to_docx(doc, "### Durchführung\nText darunter.", headings_as_bold=True)
    # Erster Nicht-Leer-Paragraph soll Normal-Style (kein Heading) sein
    paras = [p for p in doc.paragraphs if p.text.strip()]
    assert len(paras) >= 1
    first = paras[0]
    assert first.style.name == 'Normal', f"Erwartet Normal, erhalten: {first.style.name}"
    assert first.runs and first.runs[0].bold, "Heading-Text soll bold=True sein"


def test_process_markdown_headings_default_uses_heading_style():
    """headings_as_bold=False (Standard) → ### Titel wird als Heading 3 eingefügt."""
    doc = Document()
    process_markdown_to_docx(doc, "### Titel", headings_as_bold=False)
    paras = [p for p in doc.paragraphs if p.text.strip()]
    assert len(paras) >= 1
    assert 'Heading' in paras[0].style.name or paras[0].style.name.startswith('berschrift'), \
        f"Erwartet Heading-Style, erhalten: {paras[0].style.name}"


# ---------------------------------------------------------------------------
# Gruppe 22: _strip_ocr_y_prefix – Headings
# ---------------------------------------------------------------------------

def test_strip_y_prefix_heading():
    """'## y **Titel**' → '## **Titel**' (y-Präfix entfernt)."""
    text = "## y **Woran erkenne ich das?**"
    result = _strip_ocr_y_prefix(text)
    assert result == "## **Woran erkenne ich das?**"


def test_strip_y_prefix_all_heading_levels():
    """y-Präfix wird bei allen Heading-Levels (# bis ######) entfernt."""
    for level in range(1, 7):
        hashes = '#' * level
        text = f"{hashes} y Titel"
        result = _strip_ocr_y_prefix(text)
        assert result == f"{hashes} Titel", f"Fehler bei Level {level}: {result!r}"


# ---------------------------------------------------------------------------
# Gruppe 23: _strip_ocr_y_prefix – Bullets und Sonderfälle
# ---------------------------------------------------------------------------

def test_strip_y_prefix_bullet():
    """'- y Inhalt' → '- Inhalt' (y-Präfix bei Bullet-Items entfernt)."""
    text = "- y erfasst neue Situationen"
    result = _strip_ocr_y_prefix(text)
    assert result == "- erfasst neue Situationen"


def test_strip_y_prefix_no_false_positive_word():
    """'- young person' bleibt unverändert (y ist Teil eines Wortes, kein isoliertes Präfix)."""
    text = "- young person"
    result = _strip_ocr_y_prefix(text)
    assert result == "- young person"


def test_strip_y_prefix_regular_text_unchanged():
    """Fließtext mit 'y' bleibt unverändert (kein Heading/Bullet-Präfix)."""
    text = "Das Wort yesterday enthält y, aber das ist kein Präfix."
    result = _strip_ocr_y_prefix(text)
    assert result == text


def test_strip_y_prefix_multiline():
    """Mehrere Zeilen: nur Zeilen mit y-Präfix werden bereinigt."""
    text = "## y Kapitel\nNormaler Text mit y drin.\n- y Bullet\n- Normales Bullet"
    result = _strip_ocr_y_prefix(text)
    lines = result.split('\n')
    assert lines[0] == "## Kapitel"
    assert lines[1] == "Normaler Text mit y drin."
    assert lines[2] == "- Bullet"
    assert lines[3] == "- Normales Bullet"


# ---------------------------------------------------------------------------
# Gruppe 24: build_interleaved_word_document – unnummerierte Headings in Nav
# ---------------------------------------------------------------------------

def _run_interleaved(orig_md: str, summary_md: str) -> list:
    """Hilfsfunktion: build_interleaved_word_document in Temp-Datei ausführen, Paragraphs zurückgeben."""
    import tempfile, os
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = tmp.name
    try:
        build_interleaved_word_document(
            translated_text=orig_md,
            summary_text=summary_md,
            qa_text="",
            output_path=tmp_path,
        )
        from docx import Document as _Doc
        doc = _Doc(tmp_path)
        return [p for p in doc.paragraphs if p.text.strip()]
    finally:
        os.unlink(tmp_path)


def test_interleaved_unnumbered_heading_not_in_nav():
    """Wiederkehrende unnummerierte Überschrift (freq>1) → Normal+Bold, nicht in Nav."""
    # "Was ist das?" erscheint unter 5.1 UND 5.2 → freq=2 → kein Auto-Nummerierung
    orig_md = (
        "## 5.1 Kapitel\n\n### Was ist das?\n\nDefinition A.\n\n"
        "## 5.2 Anderes\n\n### Was ist das?\n\nDefinition B."
    )
    sum_md = "## 5.1 Kapitel\n\nZusammenfassung.\n\n### Was ist das?\n\nKurze Zusammenfassung."
    paras = _run_interleaved(orig_md, sum_md)
    heading_para = next((p for p in paras if "Was ist das" in p.text), None)
    assert heading_para is not None, "Überschrift 'Was ist das?' nicht im Dokument gefunden"
    assert heading_para.style.name == 'Normal', (
        f"Wiederkehrende Überschrift soll Normal-Style haben, erhalten: {heading_para.style.name!r}"
    )
    assert heading_para.runs and heading_para.runs[0].bold, \
        "Wiederkehrende Überschrift soll bold=True sein"


def test_interleaved_numbered_heading_in_nav():
    """Nummerierte Sektion mit Summary → Heading-Style (erscheint in Nav)."""
    orig_md = "## 5.1 Einführung\n\nHier steht der Originaltext."
    sum_md = "## 5.1 Einführung\n\nKurze Zusammenfassung."
    paras = _run_interleaved(orig_md, sum_md)
    heading_para = next((p for p in paras if "5.1 Einführung" in p.text), None)
    assert heading_para is not None, "Überschrift '5.1 Einführung' nicht im Dokument gefunden"
    assert 'Heading' in heading_para.style.name or heading_para.style.name.startswith('berschrift'), (
        f"Nummerierte Überschrift soll Heading-Style haben, erhalten: {heading_para.style.name!r}"
    )


def test_interleaved_unnumbered_h2_in_nav_when_no_numbered():
    """Unnummerierte H2 in Dokument OHNE nummerierte Kapitel → Heading-Style (smart nav)."""
    orig_md = "## Einführung\n\nHier steht der Originaltext."
    sum_md = "## Einführung\n\nKurze Zusammenfassung."
    paras = _run_interleaved(orig_md, sum_md)
    heading_para = next((p for p in paras if "Einführung" in p.text), None)
    assert heading_para is not None, "Überschrift 'Einführung' nicht im Dokument gefunden"
    assert 'Heading' in heading_para.style.name or heading_para.style.name.startswith('berschrift'), (
        f"Unnummerierte H2 (kein nummeriertes Dokument) soll Heading-Style haben, erhalten: {heading_para.style.name!r}"
    )


def test_interleaved_unnumbered_h2_not_in_nav_when_numbered_doc():
    """Wiederkehrende unnummerierte H2 mit Summary in numm. Dokument → Normal+Bold (nicht in Nav)."""
    # "Woran erkenne ich das?" erscheint zweimal → freq=2 → keine Auto-Nummerierung.
    # Summary enthält "Woran erkenne ich das?" → wird angezeigt, aber als Normal+Bold (nicht in Nav).
    orig_md = (
        "## 5.1 Einführung\n\nKapiteltext.\n\n## Woran erkenne ich das?\n\nText A.\n\n"
        "## 5.2 Weiteres\n\nKapiteltext.\n\n## Woran erkenne ich das?\n\nText B."
    )
    sum_md = (
        "## 5.1 Einführung\n\nZusammenfassung.\n\n"
        "## 5.2 Weiteres\n\nZusammenfassung 2.\n\n"
        "## Woran erkenne ich das?\n\nKompetenz-Merkmale."
    )
    paras = _run_interleaved(orig_md, sum_md)
    heading_para = next((p for p in paras if "Woran erkenne ich" in p.text), None)
    assert heading_para is not None, "Überschrift 'Woran erkenne ich das?' nicht gefunden"
    assert heading_para.style.name == 'Normal', (
        f"Wiederkehrende H2 in numm. Dokument soll Normal-Style haben, erhalten: {heading_para.style.name!r}"
    )
    assert heading_para.runs and heading_para.runs[0].bold, "Soll bold=True sein"


# ---------------------------------------------------------------------------
# Gruppe 25: has_children erkennt unnummerierte Folge-Sektion als konzept. Kind
# ---------------------------------------------------------------------------

def test_numbered_section_not_skipped_when_followed_by_unnumbered():
    """Nummerierte Sektion ohne Body/Summary erscheint, wenn direkt danach unnummerierte Sektion folgt."""
    # 5.3 hat keinen Body und kein Summary, aber "Agilität" (unnummeriert) folgt direkt
    orig_md = "## 5.3 Überblick\n\n## Agilität\n\nKompetenz-Inhalt hier."
    sum_md = "## Agilität\n\nAgilität-Zusammenfassung."
    paras = _run_interleaved(orig_md, sum_md)
    # 5.3 muss im Dokument erscheinen (nicht übersprungen)
    heading_para = next((p for p in paras if "5.3 Überblick" in p.text), None)
    assert heading_para is not None, "5.3 Überblick soll nicht übersprungen werden (has_children via unnummerierte Folgesektion)"
    assert 'Heading' in heading_para.style.name or heading_para.style.name.startswith('berschrift'), \
        f"5.3 soll Heading-Style haben, erhalten: {heading_para.style.name!r}"


# ---------------------------------------------------------------------------
# Gruppe 26: Auto-Nummerierung einzigartiger Unterkapitel + Skip-Fix für Orig-Body
# ---------------------------------------------------------------------------

def test_auto_numbering_unique_heading_gets_number_and_nav():
    """Einzigartige unnummerierte Überschrift nach numm. Kapitel → auto 5.3.1, in Nav."""
    orig_md = "## 5.3 Überblick\n\n## Agilität\n\nAgilität ist die Fähigkeit..."
    sum_md = "## Agilität\n\nZusammenfassung Agilität."
    paras = _run_interleaved(orig_md, sum_md)
    heading_para = next((p for p in paras if "5.3.1" in p.text and "Agilität" in p.text), None)
    assert heading_para is not None, "Einzigartige Überschrift soll als '5.3.1 Agilität' erscheinen"
    assert 'Heading' in heading_para.style.name or heading_para.style.name.startswith('berschrift'), \
        f"Auto-nummerierte Überschrift soll Heading-Style haben, erhalten: {heading_para.style.name!r}"


def test_auto_numbering_recurring_heading_stays_bold_not_nav():
    """Wiederkehrende unnummerierte Überschrift mit Summary → Normal+Bold, keine Auto-Nummerierung."""
    # "Was ist das?" erscheint zweimal → freq=2 → NICHT auto-nummeriert → Normal+Bold (nicht Heading-Style).
    # Summary enthält "Was ist das?" damit es nicht übersprungen wird.
    orig_md = (
        "## 5.3 Überblick\n\n"
        "## Agilität\n\n"
        "## Was ist das?\n\nDefinition Agilität.\n\n"
        "## Analytisches Denken\n\n"
        "## Was ist das?\n\nDefinition Analytik."
    )
    sum_md = (
        "## Agilität\n\nAgilität-Zusammenfassung.\n\n"
        "## Analytisches Denken\n\nAnalytik-Zusammenfassung.\n\n"
        "## Was ist das?\n\nDefinition."
    )
    paras = _run_interleaved(orig_md, sum_md)
    was_paras = [p for p in paras if "Was ist das?" in p.text]
    assert len(was_paras) >= 1, "Mindestens eine 'Was ist das?'-Überschrift soll im Dokument erscheinen"
    for wp in was_paras:
        assert wp.style.name == 'Normal', \
            f"'Was ist das?' soll Normal-Style haben (wiederkehrend), erhalten: {wp.style.name!r}"


def test_unnumbered_sibling_heading_hidden_when_no_summary():
    """Unnummerierte Geschwister-Sektion ohne Summary → Überschrift ausgeblendet.

    'Off the job' und 'On the job' erscheinen mehrfach (freq>1) → keine Auto-Nummerierung.
    Ohne Summary → show_heading_visible=False → Überschrift ausgeblendet (w:vanish).
    Die vorherige falsche Logik setzte has_children=True für jede unnummerierte Folge-Sektion
    (auch Geschwister), was diese sichtbar machte.
    """
    orig_md = (
        "## 5.3.1 Durchsetzen\n\n"
        "## Off the job\n\nMaßnahme A.\n\n"
        "## On the job\n\nMaßnahme B.\n\n"
        "## 5.3.2 Empathie\n\n"
        "## Off the job\n\nMaßnahme C."
    )
    sum_md = (
        "## 5.3.1 Durchsetzen\n\nZusammenfassung Durchsetzen.\n\n"
        "## 5.3.2 Empathie\n\nZusammenfassung Empathie."
    )
    paras = _run_interleaved(orig_md, sum_md)

    from lxml import etree
    from docx.oxml.ns import qn

    off_paras = [p for p in paras if "Off the job" in p.text]
    assert len(off_paras) >= 1, "'Off the job' muss im Dokument vorhanden sein (Originaltext erhalten)"
    for p in off_paras:
        pPr = p._p.find(qn('w:pPr'))
        has_vanish = False
        if pPr is not None:
            rPr = pPr.find(qn('w:rPr'))
            if rPr is not None and rPr.find(qn('w:vanish')) is not None:
                has_vanish = True
        assert has_vanish, \
            f"'Off the job' ohne Summary soll ausgeblendet (w:vanish) sein, Style: {p.style.name!r}"


# ---------------------------------------------------------------------------
# Gruppe 27: _any_visible_desc – kein sichtbarer Elternknoten ohne sichtbare Kinder
# ---------------------------------------------------------------------------

def test_numbered_parent_hidden_when_all_children_have_no_summary():
    """Nummerierter Parent ohne eigenes Summary und ohne Summary in Kindern → nicht sichtbar.

    Szenario analog zu '5.3.27 Rückmeldung zu Konfliktverhalten' im echten Dokument:
    Das OCR hat 'Rückmeldung zu Konfliktverhalten' als eigene Kompetenz mit Originaltext,
    die KI-Summary enthält sie NICHT als eigene Überschrift. Auch Kinder ('Off the job',
    'On the job') haben kein Summary.
    Früher: has_children=True → sichtbare goldene Leer-Überschrift.
    Nach Fix: _any_visible_desc=False → has_children=False → Überschrift ausgeblendet.

    Außerdem: '5.3 Überblick' hat 'Agilität' als konzeptuelles Kind (gleiche Ebene,
    nummerierter Parent + unnummerierter Folger). 'Agilität' hat Summary →
    _any_visible_desc[5.3]=True → '5.3 Überblick' bleibt sichtbar (kein Rückschritt).
    """
    orig_md = (
        "## 5.3 Überblick\n\n"
        "## Rückmeldung zu Konfliktverhalten\n\nWas ist das? Originaltext hier.\n\n"
        "## Off the job\n\nMaßnahme A.\n\n"
        "## On the job\n\nMaßnahme B.\n\n"
        "## Agilität\n\nAgilität ist die Fähigkeit...\n\n"
        "## Off the job\n\nMaßnahme C.\n\n"
        "## On the job\n\nMaßnahme D."
    )
    # Summary hat KEINE Einträge für Rückmeldung, Off/On the job.
    # Agilität hat ein Summary → Elternknoten "5.3 Überblick" bleibt sichtbar (hat sichtbares Kind).
    sum_md = (
        "## 5.3 Überblick\n\nÜberblick-Text.\n\n"
        "## Agilität\n\nAgilität-Zusammenfassung."
    )
    paras = _run_interleaved(orig_md, sum_md)

    from docx.oxml.ns import qn

    # "Rückmeldung zu Konfliktverhalten" muss im Dokument vorhanden sein (Originaltext erhalten)
    rueckmeldung_paras = [p for p in paras if "Rückmeldung zu Konfliktverhalten" in p.text]
    assert len(rueckmeldung_paras) >= 1, "'Rückmeldung zu Konfliktverhalten' fehlt im Dokument"

    # Muss als hidden (w:vanish) markiert sein – keine sichtbare Leer-Überschrift
    for p in rueckmeldung_paras:
        pPr = p._p.find(qn('w:pPr'))
        has_vanish = False
        if pPr is not None:
            rPr = pPr.find(qn('w:rPr'))
            if rPr is not None and rPr.find(qn('w:vanish')) is not None:
                has_vanish = True
        assert has_vanish, (
            "'Rückmeldung zu Konfliktverhalten' ohne Summary und ohne sichtbare Kinder "
            f"soll ausgeblendet sein, Style: {p.style.name!r}"
        )

    # Zur Sicherheit: "Agilität" MUSS sichtbar sein (hat Summary)
    agilitat_paras = [p for p in paras if "Agilität" in p.text and "Zusammenfassung" not in p.text]
    assert any(
        p._p.find(qn('w:pPr')) is None or
        (p._p.find(qn('w:pPr')).find(qn('w:rPr')) is None) or
        (p._p.find(qn('w:pPr')).find(qn('w:rPr')).find(qn('w:vanish')) is None)
        for p in agilitat_paras
    ), "Agilität mit Summary muss sichtbar sein"


# ---------------------------------------------------------------------------
# Gruppe 28: Sequentielle Nummerierung (keine Lücken)
# ---------------------------------------------------------------------------

def test_sequential_numbering_no_gaps():
    """3 Kompetenzen, nur 2 in Summary → Navigation zeigt 5.3.1, 5.3.2 statt 5.3.1, 5.3.3."""
    orig_md = (
        "## 5.3 Überblick\n\nEinleitungstext.\n\n"
        "## Agilität\n\nAgilität Originaltext.\n\n"
        "## Belastbarkeit\n\nBelastbarkeit Originaltext.\n\n"
        "## Durchsetzungsvermögen\n\nDurchsetzung Originaltext."
    )
    sum_md = (
        "## 5.3 Überblick\n\nÜberblick-Zusammenfassung.\n\n"
        "## Agilität\n\nAgilität-Zusammenfassung.\n\n"
        # Belastbarkeit fehlt im Summary
        "## Durchsetzungsvermögen\n\nDurchsetzung-Zusammenfassung."
    )
    paras = _run_interleaved(orig_md, sum_md)
    visible_texts = [p.text for p in paras]

    # 5.3.1 und 5.3.2 müssen vorhanden sein (keine Lücke 5.3.1, 5.3.3)
    assert any("5.3.1" in t for t in visible_texts), "5.3.1 fehlt"
    assert any("5.3.2" in t for t in visible_texts), "5.3.2 fehlt (Lücke statt sequentiell)"
    # 5.3.3 darf nicht existieren (es gibt nur 2 sichtbare Kompetenzen)
    assert not any("5.3.3" in t for t in visible_texts), "5.3.3 sollte nicht existieren"
    # Belastbarkeit bekommt keine sichtbare Nummer – darf nicht in Nav erscheinen
    belastbarkeit_nav = [t for t in visible_texts if "5.3.2 Belastbarkeit" in t or "5.3.3 Belastbarkeit" in t]
    assert len(belastbarkeit_nav) == 0, "Belastbarkeit ohne Summary darf keine Nummer haben"


# ---------------------------------------------------------------------------
# Gruppe 29: Scoped Lookup (Sub-Sections kompetenzbezogen)
# ---------------------------------------------------------------------------

def test_scoped_lookup_off_the_job_scoped_per_competency():
    """'Off the job' unter Kompetenz A soll Inhalt von A zeigen, nicht von B."""
    orig_md = (
        "## 5.3 Überblick\n\nEinleitungstext.\n\n"
        "## Agilität\n\nAgilität Originaltext.\n\n"
        "## Off the job\n\nMaßnahme Agilität original.\n\n"
        "## Belastbarkeit\n\nBelastbarkeit Originaltext.\n\n"
        "## Off the job\n\nMaßnahme Belastbarkeit original."
    )
    sum_md = (
        "## 5.3 Überblick\n\nÜberblick-Zusammenfassung.\n\n"
        "## Agilität\n\nAgilität-Zusammenfassung.\n\n"
        "#### Off the job\n\nAgilität-Maßnahme-A.\n\n"
        "## Belastbarkeit\n\nBelastbarkeit-Zusammenfassung.\n\n"
        "#### Off the job\n\nBelastbarkeit-Maßnahme-B."
    )
    paras = _run_interleaved(orig_md, sum_md)
    all_texts = " | ".join(p.text for p in paras)

    # Beide Kompetenzen müssen vorhanden sein
    assert "Agilität" in all_texts, "Agilität fehlt"
    assert "Belastbarkeit" in all_texts, "Belastbarkeit fehlt"
    # Beide Off-the-job-Inhalte müssen vorhanden sein (nicht nur der letzte)
    assert "Agilität-Maßnahme-A" in all_texts, "Agilität Off-the-job-Inhalt fehlt"
    assert "Belastbarkeit-Maßnahme-B" in all_texts, "Belastbarkeit Off-the-job-Inhalt fehlt"


# ---------------------------------------------------------------------------
# Gruppe 30: Platzhalter-Sektionen überspringen
# ---------------------------------------------------------------------------

def test_placeholder_section_skipped():
    """Section mit \\_\\_\\_ im Body (OCR-Notizlinien) wird komplett übersprungen."""
    orig_md = (
        "## 5.3 Überblick\n\nEinleitungstext.\n\n"
        "## Agilität\n\nAgilität Originaltext.\n\n"
        "## Meine persönlichen Anmerkungen\n\n\\_\\_\\_\\_\\_\\_\\_\\_\\_\\_\n\n"
        "## Wo finde ich noch mehr darüber\n\nLiteraturliste."
    )
    sum_md = (
        "## 5.3 Überblick\n\nÜberblick-Zusammenfassung.\n\n"
        "## Agilität\n\nAgilität-Zusammenfassung.\n\n"
        "## Wo finde ich noch mehr darüber\n\nBuch A, Buch B."
    )
    paras = _run_interleaved(orig_md, sum_md)
    all_texts = [p.text for p in paras]

    # "Meine persönlichen Anmerkungen" darf weder sichtbar noch als Heading erscheinen
    assert not any("Meine persönlichen Anmerkungen" in t for t in all_texts), (
        "'Meine persönlichen Anmerkungen' (Platzhalter) darf nicht im Dokument erscheinen"
    )
    # "Wo finde ich" darf erscheinen (hat Summary)
    assert any("Wo finde ich" in t for t in all_texts), "'Wo finde ich' fehlt im Dokument"


# ---------------------------------------------------------------------------
# Gruppe 31: _split_at_level2 – Splitting-Funktion
# ---------------------------------------------------------------------------

def test_split_at_level2_with_preamble():
    """Preamble-Text vor erstem ## wird korrekt separiert."""
    text = "Intro text.\n\n## Agilität\nContent A.\n\n## Belastbarkeit\nContent B."
    preamble, sections = _split_at_level2(text)
    assert "Intro text" in preamble
    assert len(sections) == 2
    assert sections[0]['heading'] == 'Agilität'
    assert sections[1]['heading'] == 'Belastbarkeit'
    assert 'Content A' in sections[0]['text']
    assert 'Content B' in sections[1]['text']


def test_split_at_level2_no_preamble():
    """Text ohne Preamble liefert leeren Preamble-String."""
    text = "## Agilität\nContent A.\n\n## Belastbarkeit\nContent B."
    preamble, sections = _split_at_level2(text)
    assert preamble.strip() == ''
    assert len(sections) == 2


def test_split_at_level2_preserves_sub_headings():
    """### Unter-Headings bleiben innerhalb ihrer Section (kein Split daran)."""
    text = "## Agilität\n### On the job\nOTJ.\n### Off the job\nOFJ.\n\n## Belastbarkeit\nText."
    preamble, sections = _split_at_level2(text)
    assert len(sections) == 2
    assert '### On the job' in sections[0]['text']
    assert '### Off the job' in sections[0]['text']


def test_split_at_level2_empty_text():
    """Leerer Text liefert leeren Preamble und keine Sections."""
    preamble, sections = _split_at_level2('')
    assert preamble == ''
    assert sections == []


# ---------------------------------------------------------------------------
# Gruppe 32: _summarize_chapter_by_sections – Pro-Section-Summarisierung
# ---------------------------------------------------------------------------

def test_summarize_chapter_by_sections_calls_per_section(tmp_path, monkeypatch):
    """Jede Level-2-Section bekommt einen eigenen _summarize_single_chapter-Call."""
    calls = []

    def fake_summarize(heading, text, output_lang="de"):
        calls.append(heading)
        return f"## {heading}\n- Zusammenfassung."

    monkeypatch.setattr('pipeline._summarize_single_chapter', fake_summarize)
    monkeypatch.setattr('pipeline.time.sleep', lambda _: None)

    ch = {
        'index': 0,
        'heading': '5.3 Kompetenzen',
        'full_text': (
            '## Agilität\nText.\n\n'
            '## Belastbarkeit\nText.\n\n'
            '## Durchsetzung\nText.\n\n'
            '## Einfühlungsvermögen\nText.'
        )
    }
    result = _summarize_chapter_by_sections(ch, tmp_path)
    assert len(calls) == 4
    assert 'Agilität' in calls
    assert 'Belastbarkeit' in calls
    assert 'Agilität' in result
    assert 'Belastbarkeit' in result


def test_summarize_chapter_by_sections_preamble_included(tmp_path, monkeypatch):
    """Preamble-Text vor erster Section wird separat zusammengefasst."""
    calls = []

    def fake_summarize(heading, text, output_lang="de"):
        calls.append(heading)
        return f"## {heading}\n- Summary."

    monkeypatch.setattr('pipeline._summarize_single_chapter', fake_summarize)
    monkeypatch.setattr('pipeline.time.sleep', lambda _: None)

    ch = {
        'index': 1,
        'heading': 'Kapitel 5',
        'full_text': 'Einleitungstext zum Kapitel.\n\n## Agilität\nText.\n\n## Belastbarkeit\nText.\n\n## Durchsetzung\nText.\n\n## Ausdauer\nText.'
    }
    result = _summarize_chapter_by_sections(ch, tmp_path)
    # 4 Sections + 1 Preamble = 5 calls
    assert len(calls) == 5
    assert calls[0] == 'Kapitel 5'  # Preamble-Call nutzt Kapitel-Heading


def test_summarize_chapter_by_sections_caching(tmp_path, monkeypatch):
    """Zweiter Aufruf überspringt API-Calls für bereits gecachte Sub-Sections."""
    call_count = [0]

    def counting_summarize(heading, text):
        call_count[0] += 1
        return f"## {heading}\n- Cached."

    monkeypatch.setattr('pipeline._summarize_single_chapter', counting_summarize)
    monkeypatch.setattr('pipeline.time.sleep', lambda _: None)

    ch = {
        'index': 0,
        'heading': 'X',
        'full_text': '## A\nT.\n\n## B\nT.\n\n## C\nT.\n\n## D\nT.'
    }
    _summarize_chapter_by_sections(ch, tmp_path)
    assert call_count[0] == 4

    call_count[0] = 0
    _summarize_chapter_by_sections(ch, tmp_path)  # alle gecacht → kein API-Call
    assert call_count[0] == 0


def test_summarize_chapter_by_sections_combined_contains_all(tmp_path, monkeypatch):
    """Kombinierter Output enthält alle Section-Headings."""
    monkeypatch.setattr('pipeline._summarize_single_chapter',
                        lambda h, t: f"## {h}\n- Punkt 1.\n- Punkt 2.")
    monkeypatch.setattr('pipeline.time.sleep', lambda _: None)

    headings = [f'Kompetenz{i}' for i in range(1, 8)]
    full_text = '\n\n'.join(f'## {h}\nText.' for h in headings)
    ch = {'index': 3, 'heading': '5.3', 'full_text': full_text}
    result = _summarize_chapter_by_sections(ch, tmp_path)
    for h in headings:
        assert h in result, f"'{h}' fehlt im kombinierten Output"


# ---------------------------------------------------------------------------
# Gruppe 33: max_output_tokens – Token-Limit je Kapitelgröße
# ---------------------------------------------------------------------------

def test_summarize_single_chapter_large_uses_high_token_limit(monkeypatch):
    """Bei > 50 Level-2-Sections wird max_output_tokens=65536 gesetzt."""
    captured_config = []

    def fake_call(model_name, contents, config, **kwargs):
        captured_config.append(config)
        class R:
            text = "## Agilität\n- Punkt."
        return R()

    monkeypatch.setattr('pipeline.call_gemini_with_retry', fake_call)
    chapter_text = '\n\n'.join(f'## Kompetenz{i}\nText.' for i in range(55))
    _summarize_single_chapter('5.3 Test', chapter_text)
    assert len(captured_config) == 1
    assert captured_config[0].max_output_tokens == 65536


def test_summarize_single_chapter_small_uses_default_token_limit(monkeypatch):
    """Bei ≤ 50 Level-2-Sections wird max_output_tokens=32768 gesetzt."""
    captured_config = []

    def fake_call(model_name, contents, config, **kwargs):
        captured_config.append(config)
        class R:
            text = "## Intro\n- Punkt."
        return R()

    monkeypatch.setattr('pipeline.call_gemini_with_retry', fake_call)
    chapter_text = '\n\n'.join(f'## Abschnitt{i}\nText.' for i in range(5))
    _summarize_single_chapter('Kap. 1', chapter_text)
    assert len(captured_config) == 1
    assert captured_config[0].max_output_tokens == 32768


# ---------------------------------------------------------------------------
# Gruppe 35: Strenge Längenbeschränkung im Prompt für große Kapitel
# ---------------------------------------------------------------------------

def test_summarize_single_chapter_large_has_strict_length_instruction(monkeypatch):
    """Bei > 50 Level-2-Sections enthält der Prompt die strenge Mengenbeschränkung."""
    captured_prompt = []

    def fake_call(model_name, contents, config, **kwargs):
        captured_prompt.append(contents)
        class R:
            text = "## A\n- Punkt."
        return R()

    monkeypatch.setattr('pipeline.call_gemini_with_retry', fake_call)
    chapter_text = '\n\n'.join(f'## Kompetenz{i}\nText.' for i in range(55))
    _summarize_single_chapter('5.3 Test', chapter_text)
    assert "Maximal 3 Stichpunkte pro Unterabschnitt" in captured_prompt[0]
    assert "Vollständigkeit" in captured_prompt[0]


def test_summarize_single_chapter_small_no_strict_length(monkeypatch):
    """Bei ≤ 10 Level-2-Sections gibt es keine strenge Mengenbeschränkung."""
    captured_prompt = []

    def fake_call(model_name, contents, config, **kwargs):
        captured_prompt.append(contents)
        class R:
            text = "## A\n- Punkt."
        return R()

    monkeypatch.setattr('pipeline.call_gemini_with_retry', fake_call)
    chapter_text = '\n\n'.join(f'## Abschnitt{i}\nText.' for i in range(5))
    _summarize_single_chapter('Kap. 1', chapter_text)
    assert "Maximal 3 Stichpunkte" not in captured_prompt[0]


# ---------------------------------------------------------------------------
# Gruppe 34: generate_summary_by_chapter – Preamble-Inklusion
# ---------------------------------------------------------------------------

def test_generate_summary_includes_preamble_in_first_chapter(tmp_path, monkeypatch):
    """Preamble-Text vor erstem nummerierten Heading landet im ersten Kapitel-Call."""
    captured_texts = []

    def fake_summarize(heading, text, output_lang="de"):
        captured_texts.append(text)
        return f"## {heading}\n- Summary."

    monkeypatch.setattr('pipeline._summarize_single_chapter', fake_summarize)
    monkeypatch.setattr('pipeline.time.sleep', lambda _: None)

    md = (
        "# Kompetenzen wirksam entwickeln\n\nEinleitungstext zum Buch.\n\n"
        "## 5.1 Anregungen\n\nAnregungstext.\n\n"
        "## 5.2 Kompetenzen\n\nKompetenztext."
    )
    generate_summary_by_chapter(md, tmp_path)

    assert len(captured_texts) >= 1, "Kein Kapitel wurde zusammengefasst"
    assert (
        "Kompetenzen wirksam entwickeln" in captured_texts[0]
        or "Einleitungstext" in captured_texts[0]
    ), "Preamble-Text fehlt im ersten Kapitel-Call"


def test_generate_summary_no_preamble_unchanged(tmp_path, monkeypatch):
    """Text ohne Preamble (beginnt direkt mit nummiertem Heading) läuft normal durch."""
    captured_texts = []

    def fake_summarize(heading, text, output_lang="de"):
        captured_texts.append(text)
        return f"## {heading}\n- Summary."

    monkeypatch.setattr('pipeline._summarize_single_chapter', fake_summarize)
    monkeypatch.setattr('pipeline.time.sleep', lambda _: None)

    md = "## 5.1 Anregungen\n\nAnregungstext.\n\n## 5.2 Kompetenzen\n\nKompetenztext."
    generate_summary_by_chapter(md, tmp_path)

    assert len(captured_texts) == 2, f"Erwartet 2 Kapitel-Calls, bekam {len(captured_texts)}"
    assert "Anregungstext" in captured_texts[0]


# ---------------------------------------------------------------------------
# Gruppe 35: _split_at_level – verallgemeinerter Level-Split
# ---------------------------------------------------------------------------

def test_split_at_level_level3_basic():
    """Splittet korrekt an ### Überschriften."""
    text = "## 5.3 Thema\n\nEinleitung\n\n### Agilität\n\nText A\n\n### Resilienz\n\nText B"
    preamble, sections = _split_at_level(text, 3)
    assert "Einleitung" in preamble
    assert len(sections) == 2
    assert sections[0]['heading'] == "Agilität"
    assert sections[1]['heading'] == "Resilienz"
    assert "Text A" in sections[0]['text']
    assert "Text B" in sections[1]['text']


def test_split_at_level_no_sections():
    """Nur Preamble, keine Sections – gibt leere sections-Liste zurück."""
    text = "## 5.3 Thema\n\nNur Fließtext, keine Unterkapitel."
    preamble, sections = _split_at_level(text, 3)
    assert sections == []
    assert "Nur Fließtext" in preamble


def test_split_at_level_does_not_split_deeper():
    """#### wird nicht als ### erkannt – bleibt im Text der übergeordneten Section."""
    text = "### Agilität\n\n#### Unterebene\n\nTieferer Text\n\n### Resilienz\n\nText R"
    preamble, sections = _split_at_level(text, 3)
    assert len(sections) == 2
    assert "#### Unterebene" in sections[0]['text']
    assert "#### Unterebene" not in sections[1]['text']


def test_split_at_level2_equivalent():
    """_split_at_level(text, 2) muss identisch zu _split_at_level2(text) sein."""
    text = "Preamble\n\n## A\n\nText A\n\n## B\n\nText B"
    p1, s1 = _split_at_level2(text)
    p2, s2 = _split_at_level(text, 2)
    assert p1 == p2
    assert [s['heading'] for s in s1] == [s['heading'] for s in s2]


# ---------------------------------------------------------------------------
# Gruppe 36: _detect_chapter_level
# ---------------------------------------------------------------------------

def test_detect_chapter_level_double_hash():
    assert _detect_chapter_level("## 5.3 Großes Thema\n\nText") == 2


def test_detect_chapter_level_triple_hash():
    assert _detect_chapter_level("### Agilität\n\nText") == 3


def test_detect_chapter_level_fallback():
    """Kein Heading → Fallback 2."""
    assert _detect_chapter_level("Nur Fließtext ohne Heading.") == 2


def test_detect_chapter_level_ignores_deeper_first():
    """Erster Heading bestimmt Level, auch wenn tiefere folgen."""
    text = "## Kapitel\n\n### Unterkapitel\n\nText"
    assert _detect_chapter_level(text) == 2


# ---------------------------------------------------------------------------
# Gruppe 37: _find_sublevel
# ---------------------------------------------------------------------------

def test_find_sublevel_finds_triple_hash():
    text = "## 5.3 Thema\n\n### Agilität\n\nText"
    assert _find_sublevel(text, 2) == 3


def test_find_sublevel_finds_quad_hash():
    text = "### Agilität\n\n#### Detail\n\nText"
    assert _find_sublevel(text, 3) == 4


def test_find_sublevel_returns_none_when_no_deeper():
    text = "## 5.3 Thema\n\nNur Fließtext, keine Unterkapitel."
    assert _find_sublevel(text, 2) is None


def test_find_sublevel_skips_same_level():
    """Gleiche Ebene wird nicht als sublevel erkannt."""
    text = "## 5.1 A\n\n## 5.2 B\n\nText"
    assert _find_sublevel(text, 2) is None


# ---------------------------------------------------------------------------
# Gruppe 38: _group_into_chunks – Greedy Grouping
# ---------------------------------------------------------------------------

def _make_sections(names_and_sizes):
    """Hilfsfunktion: erstellt sections-Liste mit kontrollierten Textgrößen."""
    return [{'heading': name, 'text': 'x' * size} for name, size in names_and_sizes]


def test_group_into_chunks_single_chunk_small():
    """Kleine sections → ein Chunk."""
    sections = _make_sections([("A", 1000), ("B", 1000), ("C", 1000)])
    chunks = _group_into_chunks("", sections, max_chars=10_000)
    assert len(chunks) == 1
    assert len(chunks[0]['sections']) == 3


def test_group_into_chunks_splits_at_boundary():
    """Splittet sobald Limit überschritten wird."""
    sections = _make_sections([("A", 8000), ("B", 8000), ("C", 8000)])
    chunks = _group_into_chunks("", sections, max_chars=10_000)
    assert len(chunks) == 3
    assert chunks[0]['sections'][0]['heading'] == "A"
    assert chunks[1]['sections'][0]['heading'] == "B"
    assert chunks[2]['sections'][0]['heading'] == "C"


def test_group_into_chunks_never_splits_single_section():
    """Eine einzelne Section > max_chars darf nicht geteilt werden."""
    sections = _make_sections([("Riesig", 20_000)])
    chunks = _group_into_chunks("", sections, max_chars=10_000)
    assert len(chunks) == 1
    assert chunks[0]['sections'][0]['heading'] == "Riesig"


def test_group_into_chunks_preamble_in_first_chunk():
    """Preamble landet immer im ersten Chunk."""
    sections = _make_sections([("A", 8000), ("B", 8000)])
    chunks = _group_into_chunks("Einleitungstext", sections, max_chars=10_000)
    assert chunks[0]['preamble'] == "Einleitungstext"
    assert chunks[1]['preamble'] == ""


def test_group_into_chunks_preamble_only():
    """Nur Preamble, keine Sections → ein Chunk mit leerem sections-Array."""
    chunks = _group_into_chunks("Nur Preamble", [], max_chars=10_000)
    assert len(chunks) == 1
    assert chunks[0]['preamble'] == "Nur Preamble"
    assert chunks[0]['sections'] == []


def test_group_into_chunks_sections_never_duplicated():
    """Jede Section erscheint in genau einem Chunk."""
    sections = _make_sections([(f"S{i}", 3000) for i in range(6)])
    chunks = _group_into_chunks("", sections, max_chars=10_000)
    all_headings = [s['heading'] for c in chunks for s in c['sections']]
    assert len(all_headings) == 6
    assert len(set(all_headings)) == 6  # keine Duplikate


# ---------------------------------------------------------------------------
# Gruppe 39: split_into_level1_chapters – Preamble-Fix (adaptiver Modus)
# ---------------------------------------------------------------------------

def test_split_adaptive_preserves_intro_text():
    """
    Einleitungstext nach # 5 und vor ## 5.1 darf nicht verloren gehen.
    Fix: pre_split_lines werden ans erste Kapitel gehängt.
    """
    md = (
        "# 5 Kompetenzen wirksam entwickeln\n\n"
        "Dieser Einleitungstext muss erhalten bleiben.\n\n"
        "## 5.1 Lerntransfer\n\nTransfertext.\n\n"
        "## 5.2 Methoden\n\nMethodentext."
    )
    chapters = split_into_level1_chapters(md)
    assert len(chapters) == 2
    # Einleitungstext muss im full_text des ersten Kapitels stecken
    assert "Einleitungstext muss erhalten bleiben" in chapters[0]['full_text']


def test_split_adaptive_intro_contains_parent_heading():
    """Der # 5-Heading selbst landet ebenfalls im Preamble des ersten Kapitels."""
    md = (
        "# 5 Kompetenzen\n\nEinleitung.\n\n"
        "## 5.1 Abschnitt\n\nText."
    )
    chapters = split_into_level1_chapters(md)
    assert "# 5 Kompetenzen" in chapters[0]['full_text']


def test_split_non_adaptive_no_regression():
    """Normaler (nicht-adaptiver) Split bleibt unverändert – kein Preamble-Overhead."""
    md = "# 4 Thema A\n\nText A.\n\n# 5 Thema B\n\nText B."
    chapters = split_into_level1_chapters(md)
    assert len(chapters) == 2
    assert "Text A" in chapters[0]['full_text']
    assert "Text B" in chapters[1]['full_text']


def test_split_adaptive_substantial_intro_becomes_own_chapter():
    """Kapitel-Root mit substanziellem Intro (>300 Zeichen) vor dem ersten Unterkapitel
    wird ein eigenes führendes Kapitel (statt in Kap. 1 eingeschmolzen), damit seine
    Zusammenfassung dem Root-Heading zugeordnet werden kann."""
    intro = ("Unternehmen gestalten Lernprozesse systematisch. " * 12).strip()  # >300 Zeichen
    md = (
        "# 3 Analyse und Handlungsempfehlungen\n\n"
        f"{intro}\n\n"
        "## 3.1 Der Personalentwicklungszyklus\n\nText 3.1.\n\n"
        "## 3.2 Bedarfsanalyse\n\nText 3.2."
    )
    chapters = split_into_level1_chapters(md)
    assert chapters[0]['heading'] == "3 Analyse und Handlungsempfehlungen", \
        f"Erstes Kapitel soll der Root sein, ist: {chapters[0]['heading']!r}"
    assert intro[:40] in chapters[0]['full_text'], "Intro-Text muss im Root-Kapitel stecken"
    # Unterkapitel folgen separat und enthalten den Intro NICHT mehr.
    sub_headings = [c['heading'] for c in chapters[1:]]
    assert any("3.1" in h for h in sub_headings) and any("3.2" in h for h in sub_headings)
    assert intro[:40] not in chapters[1]['full_text'], \
        "Intro darf nicht mehr doppelt in Kapitel 3.1 stecken"


def test_split_adaptive_short_intro_still_folded():
    """Kurzer Intro (<300 Zeichen) bleibt wie bisher in Kapitel 1 eingeschmolzen –
    kein leeres/separates Root-Kapitel."""
    md = (
        "# 3 Analyse\n\nKurzer Intro.\n\n"
        "## 3.1 Abschnitt\n\nText 3.1.\n\n"
        "## 3.2 Abschnitt\n\nText 3.2."
    )
    chapters = split_into_level1_chapters(md)
    assert len(chapters) == 2, f"Kein separates Root-Kapitel erwartet: {[c['heading'] for c in chapters]}"
    assert "Kurzer Intro" in chapters[0]['full_text']


# ---------------------------------------------------------------------------
# Gruppe 40: _summarize_single_chapter – _required erfasst ### und ####
# ---------------------------------------------------------------------------

def test_summarize_required_includes_triple_hash(monkeypatch):
    """_required muss auch ### Überschriften erfassen, nicht nur ##."""
    captured_prompts = []

    def fake_call(model_name, contents, config):
        captured_prompts.append(contents)
        class R:
            text = "## Heading\n- bullet"
        return R()

    monkeypatch.setattr('pipeline.call_gemini_with_retry', fake_call)

    chapter_text = (
        "## 5.3 Großes Thema\n\n"
        "### Agilität\n\nText zu Agilität.\n\n"
        "### Resilienz\n\nText zu Resilienz."
    )
    _summarize_single_chapter("5.3 Großes Thema", chapter_text)

    prompt = captured_prompts[0]
    assert "Agilität" in prompt, "### Heading fehlt in _required-Liste"
    assert "Resilienz" in prompt, "### Heading fehlt in _required-Liste"
    assert "PFLICHT-VOLLSTÄNDIGKEIT" in prompt


def test_summarize_required_triggers_at_one_heading(monkeypatch):
    """Bereits bei 1 Unterkapitel wird PFLICHT-VOLLSTÄNDIGKEIT eingefügt (>= 1)."""
    captured_prompts = []

    def fake_call(model_name, contents, config):
        captured_prompts.append(contents)
        class R:
            text = "## Heading\n- bullet"
        return R()

    monkeypatch.setattr('pipeline.call_gemini_with_retry', fake_call)

    chapter_text = "## 5.1 Thema\n\n### Einziges Unterkapitel\n\nText."
    _summarize_single_chapter("5.1 Thema", chapter_text)

    assert "PFLICHT-VOLLSTÄNDIGKEIT" in captured_prompts[0]


def test_summarize_required_empty_for_flat_chapter(monkeypatch):
    """Kapitel ohne Unterkapitel → kein PFLICHT-VOLLSTÄNDIGKEIT-Block."""
    captured_prompts = []

    def fake_call(model_name, contents, config):
        captured_prompts.append(contents)
        class R:
            text = "## Heading\n- bullet"
        return R()

    monkeypatch.setattr('pipeline.call_gemini_with_retry', fake_call)

    # Nur ## Kapitel-Heading, keine ### Unterkapitel → _required leer
    chapter_text = "## 5.1 Thema\n\nNur Fließtext, keine Unterkapitel hier."
    _summarize_single_chapter("5.1 Thema", chapter_text)

    assert "PFLICHT-VOLLSTÄNDIGKEIT" not in captured_prompts[0], (
        "Für ein Kapitel ohne Sub-Headings darf kein PFLICHT-Block erscheinen"
    )


# ---------------------------------------------------------------------------
# Gruppe 41: generate_summary_by_chapter – Chunk-Splitting
# ---------------------------------------------------------------------------

def test_generate_chunks_large_chapter(tmp_path, monkeypatch):
    """Ein Kapitel > 15.000 Zeichen mit ### Sections wird in mehrere Chunks aufgeteilt."""
    call_log = []

    def fake_summarize(heading, text, output_lang="de"):
        call_log.append({'heading': heading, 'len': len(text)})
        return f"## {heading}\n- Summary."

    monkeypatch.setattr('pipeline._summarize_single_chapter', fake_summarize)
    monkeypatch.setattr('pipeline.time.sleep', lambda _: None)

    # Dokumentstruktur: # 5 → adaptive split bei ## → 5.3 als eigenes Kapitel
    # 5.3 hat 6 Sections à ~4000 Zeichen = ~24.000 gesamt → sollte in Chunks aufgeteilt werden
    sections_md = "\n\n".join(
        f"### Section {i}\n\n{'x' * 4000}" for i in range(1, 7)
    )
    md = (
        "# 5 Hauptthema\n\nEinleitung Hauptthema.\n\n"
        f"## 5.3 Großes Thema\n\nEinleitung Abschnitt.\n\n{sections_md}"
    )

    generate_summary_by_chapter(md, tmp_path)

    assert len(call_log) >= 2, f"Erwartet >= 2 Calls für großes Kapitel, bekam {len(call_log)}"
    # Kein einzelner Call darf den vollen Text (~24k) bekommen haben
    assert all(c['len'] < 20_000 for c in call_log), "Ein Chunk ist zu groß"


def test_generate_chunks_small_chapter_single_call(tmp_path, monkeypatch):
    """Ein Kapitel < 15.000 Zeichen → genau ein API-Call."""
    call_log = []

    def fake_summarize(heading, text, output_lang="de"):
        call_log.append(heading)
        return f"## {heading}\n- Summary."

    monkeypatch.setattr('pipeline._summarize_single_chapter', fake_summarize)
    monkeypatch.setattr('pipeline.time.sleep', lambda _: None)

    md = "## 5.1 Kleines Thema\n\n### Unter A\n\nText.\n\n### Unter B\n\nText."
    generate_summary_by_chapter(md, tmp_path)

    assert len(call_log) == 1, f"Erwartet 1 Call, bekam {len(call_log)}"


def test_generate_chunks_all_sections_covered(tmp_path, monkeypatch):
    """Nach chunk-weiser Verarbeitung sind alle Section-Headings im kombinierten Output."""
    import re as _re

    def fake_summarize(heading, text, output_lang="de"):
        found = _re.findall(r'^#{2,6}\s+(.+)$', text, _re.MULTILINE)
        lines = [f"### {h}\n- bullet" for h in found] if found else [f"## {heading}\n- bullet"]
        return "\n\n".join(lines)

    monkeypatch.setattr('pipeline._summarize_single_chapter', fake_summarize)
    monkeypatch.setattr('pipeline.time.sleep', lambda _: None)

    sections_md = "\n\n".join(
        f"### Section {i}\n\n{'x' * 4000}" for i in range(1, 7)
    )
    md = (
        "# 5 Hauptthema\n\nEinleitung Hauptthema.\n\n"
        f"## 5.3 Großes Thema\n\nEinleitung Abschnitt.\n\n{sections_md}"
    )

    result = generate_summary_by_chapter(md, tmp_path)

    for i in range(1, 7):
        assert f"Section {i}" in result, f"Section {i} fehlt im kombinierten Output"


# ---------------------------------------------------------------------------
# Gruppe 30: Kapitel-Extraktion & Nummern-Rebase (--chapter + --parent-chapter)
# ---------------------------------------------------------------------------

def test_extract_chapter_keeps_unnumbered_subsections():
    """extract_chapter bricht NICHT bei unnummerierten '#'-Zwischenüberschriften ab."""
    md = (
        "# 1.3 Lernen als Wissenserwerb\n\nIntro.\n\n"
        "#### Wie wird Wissen erworben?\n\nAbsatz A.\n\n"
        "# Wie ist Wissen repräsentiert?\n\nAbsatz B.\n\n"
        "# 1.4 Konstruktion\n\nNächstes Kapitel.\n"
    )
    out = extract_chapter(md, "1.3")
    assert "Wie ist Wissen repräsentiert?" in out
    assert "Absatz B." in out
    assert "1.4 Konstruktion" not in out
    assert "Nächstes Kapitel." not in out


def test_rebase_chapter_number_root_and_descendants():
    assert _rebase_chapter_number("1.3 Titel", "1.3", "2.4.1.1") == "2.4.1.1 Titel"
    assert _rebase_chapter_number("1.3.2 Unterkapitel", "1.3", "2.4.1.1") == "2.4.1.1.2 Unterkapitel"
    assert _rebase_chapter_number("1.3", "1.3", "2.4.1.1") == "2.4.1.1"


def test_rebase_chapter_number_non_descendant_fallback():
    """Nicht-Nachfahren werden präfixiert (Fallback)."""
    assert _rebase_chapter_number("2 Anderes", "1.3", "2.4.1.1") == "2.4.1.1.2 Anderes"


# ---------------------------------------------------------------------------
# Gruppe 31: Box-Strukturreparatur (Fokus/Studie/Definition/Beispiel)
# ---------------------------------------------------------------------------

def test_is_box_heading_detects_labels():
    assert _is_box_heading("Fokus: Bedeutung des Hippocampus")
    assert _is_box_heading("**Studie: Bildrätsel**")
    assert _is_box_heading("Definition: Proposition")
    assert _is_box_heading("Beispiel: Advance Organizer")
    assert not _is_box_heading("Wie ist Wissen repräsentiert?")
    assert not _is_box_heading("1.3 Lernen als Wissenserwerb")


def test_split_paragraphs_basic():
    assert _split_paragraphs("A\n\nB\n\n\nC") == ["A", "B", "C"]
    assert _split_paragraphs("") == []


def test_repair_box_structure_no_box_unchanged(monkeypatch):
    """Ohne Kasten: Text unverändert, kein Klassifikator-Aufruf."""
    called = []
    monkeypatch.setattr('pipeline._classify_box_boundaries', lambda b: called.append(b) or [])
    md = "# Kapitel\n\nText.\n\n## Unterkapitel\n\nMehr Text.\n"
    assert repair_box_structure(md) == md
    assert called == []


def test_repair_box_structure_reattaches_running_text(monkeypatch):
    """Fließtext-Suffix wandert zum Eltern, Kasten behält nur Kasteninhalt."""
    monkeypatch.setattr('pipeline._classify_box_boundaries', lambda boxes: [2])
    md = (
        "# Wie ist Wissen repräsentiert?\n\nIntro-Absatz.\n\n"
        "#### Fokus: Hippocampus\n\n"
        "Kasten Absatz 1.\n\nKasten Absatz 2.\n\n"
        "Fließtext Absatz 1.\n\nFließtext Absatz 2.\n\n"
        "# Nächster Abschnitt\n\nEgal.\n"
    )
    out = repair_box_structure(md)
    # Parent enthält jetzt den Fließtext
    parent_part = out.split("#### Fokus")[0]
    assert "Intro-Absatz." in parent_part
    assert "Fließtext Absatz 1." in parent_part
    assert "Fließtext Absatz 2." in parent_part
    # Kasten enthält NUR den Kasteninhalt
    box_part = out.split("#### Fokus: Hippocampus")[1].split("# Nächster Abschnitt")[0]
    assert "Kasten Absatz 1." in box_part
    assert "Kasten Absatz 2." in box_part
    assert "Fließtext Absatz 1." not in box_part


def test_repair_box_structure_woven_two_boxes(monkeypatch):
    """Zwei Kästen, beide Suffixe landen in Reihenfolge im selben Parent."""
    monkeypatch.setattr('pipeline._classify_box_boundaries', lambda boxes: [1, 1])
    md = (
        "# Eltern\n\nIntro.\n\n"
        "#### Studie: Eins\n\nStudieinhalt.\n\nSuffix Eins.\n\n"
        "#### Definition: Zwei\n\nDefinhalt.\n\nSuffix Zwei.\n\n"
        "# Ende\n\nfertig.\n"
    )
    out = repair_box_structure(md)
    parent_part = out.split("#### Studie")[0]
    assert "Suffix Eins." in parent_part
    assert "Suffix Zwei." in parent_part
    # Reihenfolge: Suffix Eins vor Suffix Zwei
    assert parent_part.index("Suffix Eins.") < parent_part.index("Suffix Zwei.")
    # Beide Kästen erscheinen sauber nach dem Elterntext
    assert "#### Studie: Eins" in out
    assert "#### Definition: Zwei" in out


def test_repair_box_structure_fallback_on_error(monkeypatch):
    """Klassifikator-Fehler → Originaltext unverändert (kein Crash)."""
    def boom(boxes):
        raise RuntimeError("API down")
    monkeypatch.setattr('pipeline._classify_box_boundaries', boom)
    md = (
        "# Eltern\n\nIntro.\n\n"
        "#### Fokus: X\n\nKasten.\n\nFließtext.\n\n"
    )
    assert repair_box_structure(md) == md


def test_classify_box_boundaries_clamps(monkeypatch):
    """Klassifikator clampt n auf [1, len(paragraphs)]."""
    class FakeResp:
        text = '{"0": 0, "1": 99}'

    monkeypatch.setattr('pipeline.call_gemini_with_retry',
                        lambda **kwargs: FakeResp())
    boxes = [
        {'title': 'Fokus: A', 'paragraphs': ['p1', 'p2', 'p3']},
        {'title': 'Studie: B', 'paragraphs': ['q1', 'q2']},
    ]
    assert _classify_box_boundaries(boxes) == [1, 2]


# ---------------------------------------------------------------------------
# Gruppe 32: Heading-Level im Einbette-Modus (--chapter + --parent-chapter)
# ---------------------------------------------------------------------------

def _run_interleaved_embedded(orig_md: str, summary_md: str,
                              parent_chapter: str, extracted_chapter: str,
                              qa_text: str = "") -> list:
    """Hilfsfunktion: build_interleaved_word_document im Einbette-Modus."""
    import tempfile, os
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = tmp.name
    try:
        build_interleaved_word_document(
            translated_text=orig_md,
            summary_text=summary_md,
            qa_text=qa_text,
            output_path=tmp_path,
            parent_chapter=parent_chapter,
            extracted_chapter=extracted_chapter,
        )
        from docx import Document as _Doc
        doc = _Doc(tmp_path)
        return [p for p in doc.paragraphs if p.text.strip()]
    finally:
        os.unlink(tmp_path)


def test_embed_unnumbered_sections_get_counter_numbers():
    """Extrahiertes Kapitel (has_numbered_chapters=True) mit unnummerierten Unterabschnitten:
    Counter-Nummern werden vergeben (2.4.1.2.1, 2.4.1.2.2, …)."""
    orig_md = (
        "# 1.3 Lernen\n\nIntrotext.\n\n"
        "# Wie wird Wissen erworben?\n\nText A.\n\n"
        "# Wie ist Wissen repräsentiert?\n\nText B.\n\n"
    )
    sum_md = (
        "## 1.3 Lernen\n\nZusammenfassung.\n\n"
        "## Wie wird Wissen erworben?\n\nSummary A.\n\n"
        "## Wie ist Wissen repräsentiert?\n\nSummary B.\n\n"
    )
    paras = _run_interleaved_embedded(orig_md, sum_md,
                                     parent_chapter="2.4.1.2", extracted_chapter="1.3")
    texts = [p.text for p in paras]
    assert any("2.4.1.2.1" in t and "Wie wird Wissen erworben" in t for t in texts), \
        f"Erster Unterabschnitt soll '2.4.1.2.1' bekommen. Gefunden:\n{texts}"
    assert any("2.4.1.2.2" in t and "Wie ist Wissen repräsentiert" in t for t in texts), \
        f"Zweiter Unterabschnitt soll '2.4.1.2.2' bekommen. Gefunden:\n{texts}"


def test_embed_box_headings_not_numbered():
    """Kästen (Fokus/Studie) erhalten keine Counter-Nummer im Einbette-Modus."""
    orig_md = (
        "# 1.3 Lernen\n\nIntrotext.\n\n"
        "# Erster Abschnitt\n\nText.\n\n"
        "#### Fokus: Hippocampus\n\nKasteninhalt.\n\n"
    )
    sum_md = (
        "## 1.3 Lernen\n\nZusammenfassung.\n\n"
        "## Erster Abschnitt\n\nSummary.\n\n"
        "## Fokus: Hippocampus\n\nFokus-Summary.\n\n"
    )
    paras = _run_interleaved_embedded(orig_md, sum_md,
                                     parent_chapter="2.4.1.2", extracted_chapter="1.3")
    texts = [p.text for p in paras]
    fokus_texts = [t for t in texts if "Fokus" in t and "Hippocampus" in t]
    assert fokus_texts, f"Fokus-Kasten-Überschrift soll im Dokument erscheinen. Texte:\n{texts}"
    # Kasten darf KEINE Nummer der Form "2.4.1.2.N" tragen
    for ft in fokus_texts:
        assert "2.4.1.2." not in ft, \
            f"Fokus-Kasten soll keine Kapitelnummer erhalten, aber: {ft!r}"


def test_embed_box_heading_level_is_lvlshift_plus_2():
    """Kasten-Überschrift im Einbette-Modus: Word-Heading = lvl_shift+2 statt level+lvl_shift.
    parent_chapter='2.4.1.2' → lvl_shift=5 → box display_level=7 (nicht 9)."""
    orig_md = (
        "# 1.3 Lernen\n\nIntrotext.\n\n"
        "#### Fokus: Hippocampus\n\nKasteninhalt.\n\n"
    )
    sum_md = (
        "## 1.3 Lernen\n\nZusammenfassung.\n\n"
        "## Fokus: Hippocampus\n\nFokus-Summary.\n\n"
    )
    paras = _run_interleaved_embedded(orig_md, sum_md,
                                     parent_chapter="2.4.1.2", extracted_chapter="1.3")
    box_para = next((p for p in paras if "Fokus" in p.text and "Hippocampus" in p.text), None)
    assert box_para is not None, "Fokus-Kasten-Überschrift nicht gefunden"
    style = box_para.style.name
    assert "7" in style, (
        f"Fokus-Kasten soll Heading 7 sein (lvl_shift=5 + 2), erhalten: {style!r}. "
        f"Früher wäre es level(4)+lvl_shift(5)=9 gewesen."
    )


def test_embed_parent_chapter_heading_level_from_number():
    """Rebasiertes Kapitel 'X.Y.Z.W' (4 Teile) → display_level=5 (Überschrift 5)."""
    orig_md = "# 1.3 Lernen\n\nIntrotext.\n\n"
    sum_md = "## 1.3 Lernen\n\nZusammenfassung.\n\n"
    paras = _run_interleaved_embedded(orig_md, sum_md,
                                     parent_chapter="2.4.1.2", extracted_chapter="1.3")
    chapter_para = next((p for p in paras if "2.4.1.2" in p.text and "Lernen" in p.text), None)
    assert chapter_para is not None, "Rebasiertes Kapitel '2.4.1.2 Lernen' nicht gefunden"
    style = chapter_para.style.name
    assert "5" in style, f"4-stelliges Kapitel soll Heading 5 sein, erhalten: {style!r}"


# ---------------------------------------------------------------------------
# Gruppe 33: extract_chapter erkennt "Kapitel N"-Format
# ---------------------------------------------------------------------------

def test_extract_chapter_label_format_found():
    """'Kapitel 7'-Heading (ohne führende Zahl) wird als Kapitelstart erkannt."""
    md = (
        "# Kapitel 6\n\n# 6.1 Vorheriges\n\nText.\n\n"
        "# Kapitel 7\n\n# 7.1 Bedeutung\n\nText A.\n\n# 7.2 Theorien\n\nText B.\n\n"
        "# Kapitel 8\n\n# 8.1 Nächstes\n\nText C.\n\n"
    )
    result = extract_chapter(md, "7")
    assert "Kapitel 7" in result
    assert "7.1 Bedeutung" in result
    assert "7.2 Theorien" in result
    assert "Kapitel 8" not in result
    assert "8.1 Nächstes" not in result


def test_extract_chapter_label_format_with_html_spans():
    """HTML-Spans im Kapitel-Heading werden korrekt ignoriert."""
    md = (
        '# <span id="p1"></span>Kapitel 7\n\n'
        '# 7.1 Unterkapitel\n\nText.\n\n'
        '# <span id="p2"></span>Kapitel 8\n\nNot included.\n\n'
    )
    result = extract_chapter(md, "7")
    assert "7.1 Unterkapitel" in result
    assert "Kapitel 8" not in result


def test_extract_chapter_numeric_format_unchanged():
    """Numerisches Format '7 Titel' funktioniert weiterhin unverändert."""
    md = "# 7 Lernen\n\nText.\n\n# 8 Next\n\nNot included.\n\n"
    result = extract_chapter(md, "7")
    assert "7 Lernen" in result
    assert "Not included" not in result


def test_extract_chapter_label_not_found_raises():
    """Wenn weder numerisch noch 'Kapitel N' gefunden: ValueError."""
    md = "# Kapitel 8\n\nText.\n\n"
    import pytest
    with pytest.raises(ValueError, match="Kapitel '7' nicht im Markdown gefunden"):
        extract_chapter(md, "7")


def test_extract_chapter_missing_toplevel_uses_first_subchapter():
    """Wenn Top-Level-Heading fehlt (OCR-Artefakt), startet Extraktion beim ersten Unterkapitel."""
    md = (
        "# 7 Vorherige Kapitel\n\nText.\n\n"
        "### 8.1 Das Lernen lernen\n\nUnterkapitel-Text.\n\n"
        "## 8.2 Modelle\n\nMehr Text.\n\n"
        "# 9 Nächstes Kapitel\n\nEnde.\n"
    )
    result = extract_chapter(md, "8")
    assert "8.1 Das Lernen lernen" in result
    assert "8.2 Modelle" in result
    assert "7 Vorherige Kapitel" not in result
    assert "9 Nächstes Kapitel" not in result


# ---------------------------------------------------------------------------
# Gruppe 34: Tail-Normalisierung + Level-2+-Heading-Behandlung im Einbette-Modus
# ---------------------------------------------------------------------------

def test_tail_numbered_heading_rebased_correctly():
    """Heading 'Titel 7.2.2' (Nummer am Ende) wird normalisiert zu '7.2.2 Titel'
    und erscheint als '2.4.2.1.2.2 Titel' (Heading 7), nicht als '2.4.2.1.2.1.1'."""
    orig_md = (
        "# 7.2 Theorien\n\nEinleitungstext.\n\n"
        "# 7.2.1 Motivation als Abwägen\n\nText.\n\n"
        "#### Motivation als Zielverfolgung 7.2.2\n\nText zu 7.2.2.\n\n"
    )
    sum_md = (
        "## 7.2 Theorien\n\nZusammenfassung.\n\n"
        "## 7.2.1 Motivation als Abwägen\n\nSummary.\n\n"
        "## Motivation als Zielverfolgung 7.2.2\n\nSummary 7.2.2.\n\n"
    )
    paras = _run_interleaved_embedded(orig_md, sum_md,
                                     parent_chapter="2.4.2.1", extracted_chapter="7")
    texts = [p.text for p in paras]
    ziel_paras = [p for p in paras if "Zielverfolgung" in p.text]
    assert ziel_paras, f"'Motivation als Zielverfolgung' nicht gefunden. Texte:\n{texts}"
    ziel_text = ziel_paras[0].text
    assert ziel_text.startswith("2.4.2.1.2.2"), (
        f"Heading soll mit '2.4.2.1.2.2' beginnen, gefunden: {ziel_text!r}"
    )
    assert "2.4.2.1.2.1.1" not in ziel_text, f"Fälschliche Tiefnummer gefunden: {ziel_text!r}"


def test_tail_normalization_false_positive_year():
    """'Studie aus dem Jahr 2023' darf bei extracted_chapter='7' NICHT umgeordnet werden,
    da '2023' nicht mit '7.' beginnt."""
    orig_md = (
        "# 7.1 Empirische Befunde\n\nText.\n\n"
        "#### Studie aus dem Jahr 2023\n\nStudieninhalt.\n\n"
    )
    sum_md = (
        "## 7.1 Empirische Befunde\n\nZusammenfassung.\n\n"
        "## Studie aus dem Jahr 2023\n\nStudien-Summary.\n\n"
    )
    paras = _run_interleaved_embedded(orig_md, sum_md,
                                     parent_chapter="2.4.2.1", extracted_chapter="7")
    studie_paras = [p for p in paras if "Studie" in p.text and "2023" in p.text]
    assert studie_paras, "Studie-Heading nicht gefunden"
    for sp in studie_paras:
        assert not sp.text.startswith("2023"), (
            f"Jahreszahl wurde fälschlich als Kapitelnummer behandelt: {sp.text!r}"
        )


def test_embed_level3_heading_not_auto_numbered_not_in_nav():
    """Level-3-Heading (###) ohne Kapitelnummer erhält im Einbette-Modus keine Auto-Nummer
    und erscheint als Normal+Bold (nicht als Word-Heading im Navigationsbereich)."""
    orig_md = (
        "# 7.2.1 Motivation als Abwägen\n\nText.\n\n"
        "### Die Skalen zur Erfassung der Lern- und Leistungsmotivation (SELLMO)\n\nSELLMO-Inhalt.\n\n"
    )
    sum_md = (
        "## 7.2.1 Motivation als Abwägen\n\nSummary.\n\n"
        "## Die Skalen zur Erfassung der Lern- und Leistungsmotivation (SELLMO)\n\nSELLMO-Summary.\n\n"
    )
    paras = _run_interleaved_embedded(orig_md, sum_md,
                                     parent_chapter="2.4.2.1", extracted_chapter="7")
    texts = [p.text for p in paras]
    sellmo_paras = [p for p in paras if "SELLMO" in p.text]
    assert sellmo_paras, f"SELLMO-Heading nicht gefunden. Texte:\n{texts}"
    for sp in sellmo_paras:
        assert not sp.text.startswith("2.4.2.1."), (
            f"SELLMO-Heading darf keine Kapitelnummer erhalten, gefunden: {sp.text!r}"
        )
    assert "Heading" not in sellmo_paras[0].style.name, (
        f"SELLMO soll als Normal+Bold erscheinen (nicht im Navigationsbereich), Style: {sellmo_paras[0].style.name!r}"
    )


def test_embed_level4_nonbox_heading_not_auto_numbered_not_in_nav():
    """Level-4-Heading (####) ohne Kapitelnummer und ohne Fokus/Studie-Label
    erhält im Einbette-Modus keine Auto-Nummer und erscheint als Normal+Bold
    (nicht als Word-Heading im Navigationsbereich)."""
    orig_md = (
        "# 7.3 Bezugsnormorientierungen\n\nText.\n\n"
        "#### Drei Arten von Bezugsnormen\n\nDefinitionstext.\n\n"
    )
    sum_md = (
        "## 7.3 Bezugsnormorientierungen\n\nSummary.\n\n"
        "## Drei Arten von Bezugsnormen\n\nSummary.\n\n"
    )
    paras = _run_interleaved_embedded(orig_md, sum_md,
                                     parent_chapter="2.4.2.1", extracted_chapter="7")
    drei_paras = [p for p in paras if "Drei Arten" in p.text]
    assert drei_paras, "Heading 'Drei Arten von Bezugsnormen' nicht gefunden"
    for dp in drei_paras:
        assert not dp.text.startswith("2.4.2.1."), (
            f"Level-4-Heading darf keine Kapitelnummer erhalten, gefunden: {dp.text!r}"
        )
    assert "Heading" not in drei_paras[0].style.name, (
        f"Level-4-Heading soll als Normal+Bold erscheinen (nicht im Navigationsbereich), Style: {drei_paras[0].style.name!r}"
    )


# ---------------------------------------------------------------------------
# Gruppe 35: Front-Matter-Stripping (Impressum, Inhaltsverzeichnis)
# ---------------------------------------------------------------------------

def test_strip_front_matter_removes_impressum_and_toc_keeps_title_and_image():
    """Titel + Titelbild bleiben, Impressum (inkl. Unterabschnitte) und Inhalt verschwinden,
    der erste echte Inhalt bleibt erhalten."""
    md = (
        "# Arbeitsschutz in der Praxis\n\n"
        "Untertitel-Zeile\n\n"
        "![](_page_0_Picture_3.jpeg)\n\n"
        "# **Impressum**\n\n"
        "## **Berücksichtigung psychischer Belastung**\n\n"
        "### **Herausgeber:**\n\nFoo Bar\n\n"
        "#### **Autorinnen und Autoren:**\n\nA, B, C\n\n"
        "# Inhalt\n\n"
        "1 Einleitung ........ 3\n\n"
        "# 1 Einleitung\n\nEchter Inhalt.\n"
    )
    out = strip_front_matter(md)
    assert "# Arbeitsschutz in der Praxis" in out
    assert "Untertitel-Zeile" in out
    assert "![](_page_0_Picture_3.jpeg)" in out
    assert "# 1 Einleitung" in out and "Echter Inhalt." in out
    # Front-Matter und alle seine Unterabschnitte sind weg.
    assert "Impressum" not in out
    assert "Herausgeber" not in out
    assert "Autorinnen und Autoren" not in out
    assert "# Inhalt" not in out
    assert "Einleitung ........ 3" not in out


def test_strip_front_matter_noop_without_front_matter():
    """Ohne Front-Matter bleibt der Text unverändert."""
    md = "# 1 Einleitung\n\nText.\n\n# 2 Methoden\n\nMehr Text.\n"
    assert strip_front_matter(md) == md


def test_strip_front_matter_english_labels():
    """Englische Front-Matter-Labels (Contents/Table of Contents) werden ebenfalls entfernt."""
    md = (
        "# The Title\n\nIntro line.\n\n"
        "# Table of Contents\n\n1 Foo .... 2\n\n"
        "# 1 Introduction\n\nReal content.\n"
    )
    out = strip_front_matter(md)
    assert "# The Title" in out
    assert "# 1 Introduction" in out and "Real content." in out
    assert "Table of Contents" not in out


def _run_interleaved_title_as_parent(orig_md: str, summary_md: str,
                                     parent_chapter: str, qa_text: str = "") -> list:
    """Hilfsfunktion: build_interleaved_word_document im strukturellen Einbette-Modus
    mit title_as_parent (kein extracted_chapter)."""
    import tempfile, os
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = tmp.name
    try:
        build_interleaved_word_document(
            translated_text=orig_md,
            summary_text=summary_md,
            qa_text=qa_text,
            output_path=tmp_path,
            parent_chapter=parent_chapter,
            title_as_parent=True,
        )
        doc = Document(tmp_path)
        return [p for p in doc.paragraphs if p.text.strip()]
    finally:
        os.unlink(tmp_path)


def test_embed_title_as_parent_root_carries_parent_chapter_number():
    """title_as_parent: Der Artikeltitel (Wurzel) trägt direkt die Elternkapitelnummer
    (z.B. '6.2.4.3 Arbeitsschutz in der Praxis')."""
    orig_md = (
        "# Arbeitsschutz in der Praxis\n\nUntertitel.\n\n"
        "# 1 Einleitung\n\nEinleitungstext.\n\n"
        "# 2 Gestaltungsbereiche\n\nText.\n\n"
    )
    sum_md = (
        "## Einleitung\n\nIntro-Zusammenfassung.\n\n"
        "## 1 Einleitung\n\nSummary 1.\n\n"
        "## 2 Gestaltungsbereiche\n\nSummary 2.\n\n"
    )
    paras = _run_interleaved_title_as_parent(orig_md, sum_md, parent_chapter="6.2.4.3")
    texts = [p.text for p in paras]
    assert any(t.startswith("6.2.4.3 ") and "Arbeitsschutz in der Praxis" in t for t in texts), \
        f"Titel soll '6.2.4.3 Arbeitsschutz in der Praxis' lauten. Gefunden:\n{texts}"


def test_embed_title_as_parent_root_is_visible_heading():
    """title_as_parent: Der Titel (Wurzel) ist ein sichtbarer Gliederungspunkt
    (Heading-Style, nicht ausgeblendet) – nicht als Normal/hidden Text."""
    orig_md = (
        "# Arbeitsschutz in der Praxis\n\nUntertitel.\n\n"
        "# 1 Einleitung\n\nEinleitungstext.\n\n"
    )
    sum_md = (
        "## 1 Einleitung\n\nSummary 1.\n\n"
    )
    paras = _run_interleaved_title_as_parent(orig_md, sum_md, parent_chapter="6.2.4.3")
    root = next((p for p in paras if "Arbeitsschutz in der Praxis" in p.text), None)
    assert root is not None, "Titel-Absatz nicht gefunden"
    assert "Heading" in root.style.name, \
        f"Titel soll Heading-Style haben, ist aber {root.style.name!r}"
    assert not any(r.font.hidden for r in root.runs), \
        "Titel-Überschrift darf nicht ausgeblendet sein"


def test_embed_summary_heading_with_span_still_matches():
    """Summary-Überschriften mit kopierten OCR-<span>-Tags müssen trotzdem auf die
    Original-Section matchen (Key wird via _clean_heading_text bereinigt)."""
    orig_md = (
        "# 1.3 Lernen\n\nIntro.\n\n"
        "# <span id=\"page-3-1\"></span>**Wie wird Wissen erworben?**\n\nOriginaltext A.\n\n"
    )
    sum_md = (
        "## 1.3 Lernen\n\nZusammenfassung.\n\n"
        "## <span id=\"page-3-1\"></span>**Wie wird Wissen erworben?**\n\n"
        "Stichpunkt-Erwerb-Inhalt.\n\n"
    )
    paras = _run_interleaved_embedded(orig_md, sum_md,
                                     parent_chapter="2.4.1.2", extracted_chapter="1.3")
    texts = [p.text for p in paras]
    assert any("Stichpunkt-Erwerb-Inhalt" in t for t in texts), \
        f"Summary-Stichpunkt soll trotz <span> im Heading erscheinen. Gefunden:\n{texts}"
