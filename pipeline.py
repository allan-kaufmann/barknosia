import os
import argparse
import subprocess
import sys
import time
import re
import json
from pathlib import Path
from dotenv import load_dotenv
import httpx  # transitive Abhängigkeit von google-genai; für Retry auf Netzwerk-Timeouts
from google import genai
from google.genai import types
from google.genai.errors import APIError
from lxml import etree
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.part import Part
from docx.opc.packuri import PackURI

# Heading-Farben aus MM-Skript MM_36633_AuG_SS2026
MM_HEADING_COLORS = {
    1: RGBColor(0xC4, 0x9A, 0x00),  # H1: dunkles Gold
    2: RGBColor(0xFF, 0xCA, 0x08),  # H2: helles Gold
    3: RGBColor(0xFF, 0xCA, 0x08),  # H3: helles Gold
    4: RGBColor(0x82, 0x66, 0x00),  # H4: dunkles Amber
    5: RGBColor(0x82, 0x66, 0x00),  # H5: dunkles Amber
    6: RGBColor(0x82, 0x66, 0x00),  # H6: dunkles Amber
    7: RGBColor(0x82, 0x66, 0x00),  # H7: dunkles Amber
    8: RGBColor(0x27, 0x27, 0x27),  # H8: fast Schwarz
    9: RGBColor(0x27, 0x27, 0x27),  # H9: fast Schwarz
}

# Lädt die Umgebungsvariablen aus der .env-Datei
load_dotenv()

# Gemini API-Client: lazy initialisiert beim ersten API-Aufruf (nicht beim Import).
# Verhindert dass Tests ohne API-Key scheitern.
_gemini_client = None

# Pro-Anfrage-Timeout (ms): verhindert, dass ein hängender Socket den Prozess endlos
# blockiert (z.B. wenn der Server die Antwort-Header nie sendet). Großzügig genug, dass
# auch lange Gemini-2.5-Pro-Generierungen für große Abschnitte nicht abgewürgt werden.
GEMINI_TIMEOUT_MS = 600_000  # 10 Minuten

def _get_gemini_client():
    global _gemini_client
    if _gemini_client is None:
        _gemini_client = genai.Client(
            http_options=types.HttpOptions(timeout=GEMINI_TIMEOUT_MS)
        )
    return _gemini_client

def call_gemini_with_retry(model_name: str, contents, config, max_retries: int = 5, delay: int = 5):
    """Hilfsfunktion: Ruft Gemini auf und wiederholt den Versuch bei Serverüberlastung (503/429)
    sowie bei Netzwerk-Timeouts/Verbindungsfehlern (hängender Socket)."""
    client = _get_gemini_client()
    for attempt in range(1, max_retries + 1):
        try:
            response = client.models.generate_content(
                model=model_name,
                contents=contents,
                config=config
            )
            return response
        except APIError as e:
            if e.code in [503, 429] and attempt < max_retries:
                print(f"      [Server ausgelastet] Fehler {e.code}. Warte {delay} Sekunden (Versuch {attempt}/{max_retries})...")
                time.sleep(delay)
                delay *= 2
            else:
                raise e
        except (httpx.TimeoutException, httpx.TransportError) as e:
            # Hängende oder abgebrochene Verbindung: erneut versuchen statt endlos zu warten.
            if attempt < max_retries:
                print(f"      [Netzwerk-Problem] {type(e).__name__}. Warte {delay} Sekunden (Versuch {attempt}/{max_retries})...")
                time.sleep(delay)
                delay *= 2
            else:
                raise e
    raise APIError("Maximale Anzahl an Wiederholungsversuchen erreicht.")


def load_or_run(path: Path, generator_fn, label: str) -> str:
    """Resume-Hilfsfunktion: Lädt Ergebnis aus Datei wenn vorhanden, sonst generator_fn() ausführen und speichern."""
    if path.exists():
        print(f"[SKIP] {label} – bereits vorhanden: {path}")
        return path.read_text(encoding="utf-8")
    print(f"[RUN]  {label} – starte Berechnung...")
    result = generator_fn()
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(result, encoding="utf-8")
    print(f"       Ergebnis gespeichert: {path}")
    return result


def run_marker_ocr(input_pdf_path: str, output_dir: str) -> str:
    """Schritt 1: Konvertiert das PDF über den isolierten venv-Aufruf in Markdown."""
    print(f"--- Schritt 1: Starte Marker-OCR für {input_pdf_path} ---")
    os.makedirs(output_dir, exist_ok=True)

    abs_pdf_path = os.path.abspath(input_pdf_path)
    abs_output_dir = os.path.abspath(output_dir)

    if not os.path.exists(abs_pdf_path):
        raise FileNotFoundError(f"Die PDF-Datei wurde unter '{abs_pdf_path}' nicht gefunden.")

    venv_bin_dir = Path(sys.executable).parent
    local_marker_executable = venv_bin_dir / "marker_single"

    if local_marker_executable.exists():
        command = [str(local_marker_executable), str(abs_pdf_path), "--output_dir", str(abs_output_dir)]
        print(f"Nutze isolierten venv-Marker: {local_marker_executable}")
    else:
        command = ["marker_single", str(abs_pdf_path), "--output_dir", str(abs_output_dir)]
        print("Nutze Standard-Pfad für marker_single...")

    env = os.environ.copy()

    print("Führe OCR aus (das kann einen Moment dauern)...")
    try:
        subprocess.run(command, check=True, env=env, stdout=subprocess.DEVNULL, stderr=sys.stderr, shell=False)
        print("Marker erfolgreich ausgeführt.\n")
    except subprocess.CalledProcessError as e:
        print("\n[FEHLER] Marker-OCR fehlgeschlagen.")
        raise e

    pdf_stem = Path(input_pdf_path).stem
    expected_md_path = Path(abs_output_dir) / pdf_stem / f"{pdf_stem}.md"

    if expected_md_path.exists():
        return str(expected_md_path)
    else:
        found_md_files = list(Path(abs_output_dir).glob("**/*.md"))
        if found_md_files:
            return str(found_md_files[0])
        raise FileNotFoundError("Marker hat den Prozess beendet, aber es wurde keine .md-Datei gefunden.")


def _clean_heading_text(line: str) -> str:
    """Extrahiert reinen Text einer Markdown-Überschrift (ohne HTML, Links, Bold-Marker, #)."""
    text = re.sub(r'<[^>]+>', '', line)                    # HTML-Tags (<span ...>) entfernen
    text = re.sub(r'\[([^\]]*)\]\([^)]*\)', r'\1', text)  # [Text](url) → Text
    text = re.sub(r'\*+', '', text)                        # **bold** → text
    text = re.sub(r'^#+\s*', '', text)                     # führende # entfernen
    return text.strip()


def strip_ocr_page_headers(text: str) -> str:
    """Entfernt von Marker-OCR fälschlich als Überschrift gesetzte Seiten-Kolumnentitel
    (laufende Kopf-/Fußzeilen) der Form '#### **316** | 4 Schritt 3: …'.

    Erkennung: Heading-Zeile, deren bereinigter Text mit 'Seitenzahl |' beginnt
    (^\\d{1,4}\\s*\\|). Eine Zahl direkt gefolgt von '|' ist das typische Kolumnentitel-
    Muster und nie eine echte Kapitelüberschrift. Das Entfernen verhindert einen
    vorzeitigen Abbruch der Kapitel-Extraktion (die Artefakt-Zahl wirkt sonst wie der
    Beginn eines neuen Kapitels) und hält die Störzeilen aus allen Ausgaben heraus
    (saubere Navigationsleiste). Arbeitet zeilenbasiert; alle übrigen Zeilen bleiben
    unverändert.
    """
    out = []
    for line in text.split('\n'):
        if re.match(r'^#{1,6}\s', line) and re.match(r'^\d{1,4}\s*\|', _clean_heading_text(line)):
            continue
        out.append(line)
    return '\n'.join(out)


def extract_chapter(text: str, chapter_id: str, from_section: str = None,
                    to_section: str = None) -> str:
    """
    Extrahiert ein bestimmtes Kapitel (inkl. aller Unterkapitel) aus einem Markdown-Text.
    chapter_id: z.B. '4.2' oder '1' — robuste Erkennung auch bei HTML-Tags in Überschriften.
    Unterstützt zwei Heading-Formate:
      - Numerisch:  "7 Titel…"  /  "7.2 Sub…"
      - Label:      "Kapitel 7" / "Chapter 7" / "Teil 7" / "Abschnitt 7"
    Extraktion endet bei der nächsten Überschrift gleicher/höherer Ebene die kein Unterkapitel ist.

    from_section (optional, z.B. '4.2'): Extrahiere das Kapitel erst AB diesem Unterkapitel.
    Der Kapiteltitel bleibt erhalten, alles davor (Einleitung + frühere Unterkapitel wie 4.1)
    wird übersprungen – der Rest läuft bis zum Kapitelende. Nur sinnvoll mit chapter_id als
    Oberkapitel (z.B. chapter_id='4', from_section='4.2').

    to_section (optional, z.B. '4.6.3'): Extrahiere das Kapitel nur BIS EINSCHLIESSLICH dieses
    Unterkapitels. Die eigenen Unterkapitel von to_section (z.B. '4.6.3.1') bleiben erhalten;
    erst beim nächsten gleichrangigen Kapitel (z.B. '4.6.4') wird abgeschnitten. Kombinierbar
    mit from_section (z.B. chapter_id='4.6', from_section='4.6.1', to_section='4.6.3').
    """
    escaped = re.escape(chapter_id)
    _label_pat = r'(?:Kapitel|Chapter|Teil|Abschnitt)'
    lines = text.split('\n')
    start_idx = None
    heading_level = None
    label_mode = False   # True wenn Kapitel via "Kapitel N"-Format gefunden

    for i, line in enumerate(lines):
        m = re.match(r'^(#{1,6})\s', line)
        if m:
            clean = _clean_heading_text(line)
            # Primär: "7 Titel" oder "7" allein — NICHT "7.1" oder "71"
            if re.match(rf'^{escaped}(\s|$)', clean):
                start_idx = i
                heading_level = len(m.group(1))
                break
            # Fallback: "Kapitel 7", "Chapter 7", "Teil 7", "Abschnitt 7"
            if re.match(rf'^{_label_pat}\s+{escaped}(\s|$)', clean, re.IGNORECASE):
                start_idx = i
                heading_level = len(m.group(1))
                label_mode = True
                break

    # Fallback: Kein Top-Level-Heading vorhanden (OCR hat es übersprungen),
    # aber Unterkapitel existieren (z.B. "8.1 ..." wenn "8 ..." fehlt).
    if start_idx is None:
        for i, line in enumerate(lines):
            m = re.match(r'^(#{1,6})\s', line)
            if m:
                clean = _clean_heading_text(line)
                if re.match(rf'^{escaped}\.', clean):
                    sub_idx = i
                    sub_level = len(m.group(1))
                    start_idx = sub_idx
                    heading_level = sub_level
                    # Unnummerierten Kapiteltitel mitnehmen: Bücher mit nummerierten
                    # UNTER-, aber unnummerierten Hauptkapiteln (z.B. "## Lerntheorien"
                    # über "#### 3.1 …"). Die nächste vorangehende Heading-Zeile, wenn sie
                    # unnummeriert UND flacher ist, ist die Kapitel-Wurzel → Titel +
                    # Einleitung bleiben erhalten und heading_level wird flacher (wichtig
                    # für die Abbruchlogik gegen tiefe OCR-Artefakt-Headings).
                    for j in range(sub_idx - 1, -1, -1):
                        mj = re.match(r'^(#{1,6})\s', lines[j])
                        if not mj:
                            continue
                        cj = _clean_heading_text(lines[j])
                        if len(mj.group(1)) < sub_level and not re.match(r'^\d', cj):
                            start_idx = j
                            heading_level = len(mj.group(1))
                        break  # nur das unmittelbar vorangehende Heading prüfen
                    break

    if start_idx is None:
        raise ValueError(f"Kapitel '{chapter_id}' nicht im Markdown gefunden. "
                         f"Tipp: --chapter mit exakter Nummer angeben (z.B. '1' oder '4.2').")

    result = []
    for i, line in enumerate(lines):
        if i < start_idx:
            continue
        if i > start_idx:
            m2 = re.match(r'^(#{1,6})\s', line)
            if m2:
                clean2 = _clean_heading_text(line)
                level2 = len(m2.group(1))
                # Ende erst beim nächsten Heading mit echter Abschnittsnummer (z.B. '5 …'
                # oder '4.7 …'), das weder das Kapitel selbst noch ein Unterkapitel
                # (chapter_id.x) ist. KEIN Abbruch bei nummerierten Listen-Headings
                # ('1. Vorphase' → '1.' ist keine Abschnittsnummer) oder bei OCR-Kolumnentitel-
                # Resten ('316 | 4 Schritt 3 …') – beide würden sonst mitten im Kapitel
                # (z.B. vor einer Abbildung) fälschlich beenden.
                _num = re.match(r'^(\d+(?:\.\d+)*)(?:\s|$)', clean2)
                if _num \
                        and not re.match(r'^\d{1,4}\s*\|', clean2) \
                        and not re.match(rf'^{escaped}(\.|\s|$)', clean2):
                    # Gepunktete Nummer ('4.1') = echtes (Unter-)Kapitel → immer Ende.
                    # Nackte Ganzzahl ('4') nur, wenn sie auf Kapitel-Ebene (level2 <=
                    # heading_level) steht; tiefere bloße Zahlen sind OCR-Icon-Listenmarker
                    # ('#### 4 Modelllernen:') und beenden das Kapitel NICHT.
                    if '.' in _num.group(1) or level2 <= heading_level:
                        break
                # Im label_mode: auch bei "Kapitel M" (M ≠ chapter_id) stoppen.
                if label_mode and re.match(rf'^{_label_pat}\s+\d', clean2, re.IGNORECASE):
                    if not re.match(rf'^{_label_pat}\s+{escaped}(\s|$)', clean2, re.IGNORECASE):
                        break
        result.append(line)

    # Optional: nur bis einschließlich eines bestimmten Unterkapitels (inkl. dessen Unterkapitel).
    if to_section:
        ts_esc = re.escape(to_section)
        ts_idx = next((j for j, line in enumerate(result)
                       if re.match(r'^#{1,6}\s', line)
                       and re.match(rf'^{ts_esc}(\s|$)', _clean_heading_text(line))), None)
        if ts_idx is None:
            raise ValueError(f"--to '{to_section}' nicht innerhalb von Kapitel "
                             f"'{chapter_id}' gefunden (oder liegt außerhalb der Kapitelgrenzen).")
        # Ende beim ersten Heading nach to_section, das weder to_section selbst noch
        # ein Unterkapitel davon (to_section.x) ist – z.B. '4.6.4' nach '4.6.3'.
        end_idx = len(result)
        for j in range(ts_idx + 1, len(result)):
            if re.match(r'^#{1,6}\s', result[j]) and \
                    not re.match(rf'^{ts_esc}(\.|\s|$)', _clean_heading_text(result[j])):
                end_idx = j
                break
        result = result[:end_idx]

    # Optional: erst ab einem bestimmten Unterkapitel beginnen (Titel bleibt erhalten).
    if from_section:
        fs_esc = re.escape(from_section)
        fs_idx = next((j for j, line in enumerate(result)
                       if re.match(r'^#{1,6}\s', line)
                       and re.match(rf'^{fs_esc}(\s|$)', _clean_heading_text(line))), None)
        if fs_idx is None:
            raise ValueError(f"--from '{from_section}' nicht innerhalb von Kapitel "
                             f"'{chapter_id}' gefunden (oder liegt außerhalb der Kapitelgrenzen).")
        head = []
        # Kapitel-Titel behalten – aber nur, wenn die erste Zeile wirklich die Kapitel-Wurzel ist
        # (nicht ein Unterkapitel, falls die OCR den Wurzel-Heading übersprungen hat).
        if result and re.match(r'^#{1,6}\s', result[0]) and \
                re.match(rf'^{escaped}(\s|$)', _clean_heading_text(result[0])):
            head = [result[0], '']
        result = head + result[fs_idx:]

    extracted = '\n'.join(result)
    _label = f"'{chapter_id}'" + (f" ab '{from_section}'" if from_section else "") \
                               + (f" bis '{to_section}'" if to_section else "")
    print(f"[KAPITEL] {_label} extrahiert: {len(result)} Zeilen, {len(extracted)} Zeichen")
    return extracted


def parse_qa_response(qa_text: str) -> list:
    """
    Parst den strukturierten QA-Output von verify_with_questions().
    Gibt eine Liste von Dicts zurück:
      [{'num': 1, 'antwort': '...', 'textgrundlage': '...', 'schluessel': '...', 'abdeckung': '...'}, ...]
    """
    items = []
    # Splitten an "Frage N" Zeilen
    blocks = re.split(r'(?m)^\*{0,2}(?:##\s*)?Frage\s+(\d+)\*{0,2}:?\s*$', qa_text)
    # Kein "Frage N"-Header gefunden (z.B. weil das LLM das Format einer kategorisierten,
    # unnummerierten Fragendatei gespiegelt hat) → Fallback über fett gesetzte Frage-Header.
    if len(blocks) <= 1:
        return _parse_qa_by_headers(qa_text)
    # blocks[0] = text vor Frage 1, dann abwechselnd: Fragenummer, Frageblock
    i = 1
    while i < len(blocks) - 1:
        num_str = blocks[i].strip()
        block = blocks[i + 1]
        block = re.sub(r'^\*\*([^*:]+:)\*\*', r'\1', block, flags=re.MULTILINE)
        i += 2
        try:
            num = int(num_str)
        except ValueError:
            continue

        item = {'num': num}
        item.update(_parse_qa_fields(block))
        # Fragetext (vom Fallback-Parser via serialize_qa_items mitgeführt) zurücklesen,
        # damit er nach der Nachbearbeitung im Lernfragen-Kapitel angezeigt werden kann.
        _ftm = re.search(r'(?m)^Fragetext:\s*(.+)', block)
        if _ftm:
            item['frage_text'] = _ftm.group(1).strip()

        # Unterpunkt-Modus: mehrere "**Label**"-Blöcke (fett, ohne Doppelpunkt), die jeweils
        # eine eigene Antwort tragen (z.B. eine Frage mit aufgelisteten Methoden).
        subs = []
        label_matches = list(re.finditer(r'(?m)^\*\*\s*([^*:]+?)\s*\*\*\s*$', block))
        for li, lm in enumerate(label_matches):
            seg_start = lm.end()
            seg_end = label_matches[li + 1].start() if li + 1 < len(label_matches) else len(block)
            segment = block[seg_start:seg_end]
            if re.search(r'^Antwort:', segment, re.MULTILINE):
                sub = {'label': lm.group(1).strip()}
                sub.update(_parse_qa_fields(segment))
                subs.append(sub)
        if len(subs) >= 2:
            item['subs'] = subs

        items.append(item)
    return items


def _parse_qa_by_headers(qa_text: str) -> list:
    """Fallback-Parser für QA-Output ohne 'Frage N'-Header: splittet an fett gesetzten
    Frage-Überschriften (**…**). Jeder Abschnitt mit einem 'Antwort:'-Feld wird ein Item;
    ###-Kategorie-Überschriften und ein evtl. LLM-Vorspann ('Absolut, als Ihr Lerncoach…')
    werden ignoriert. Der Header-Text wird als 'frage_text' mitgeführt (Anzeige im
    Lernfragen-Kapitel, falls die Fragendatei keine nummerierten Fragen liefert)."""
    items: list = []
    num = 0
    headers = list(re.finditer(r'(?m)^\*\*\s*(.+?)\s*\*\*\s*$', qa_text))
    for i, hm in enumerate(headers):
        seg_end = headers[i + 1].start() if i + 1 < len(headers) else len(qa_text)
        seg = qa_text[hm.end():seg_end]
        if not re.search(r'(?m)^Antwort:', seg):
            continue
        num += 1
        item = {'num': num, 'frage_text': hm.group(1).strip()}
        item.update(_parse_qa_fields(seg))
        items.append(item)
    return items


def _parse_qa_fields(text: str) -> dict:
    """Extrahiert die QA-Standardfelder (Antwort/Textgrundlage/Schlüsselbegriffe/Beleg/Abdeckung)
    aus einem Block- oder Unterpunkt-Segment."""
    def _extract(pattern, default='–', multiline=False):
        m = re.search(pattern, text, re.MULTILINE | re.DOTALL)
        if m:
            val = m.group(1).strip()
            # Einzeilige Felder auf ihre erste Zeile beschränken; das Antwortfeld darf
            # mehrzeilig bleiben (strukturierte Quiz-Antworten: Zuordnung/Reihenfolge/Lücken).
            return val if multiline else val.split('\n')[0].strip()
        return default

    # Antwort bis zum nächsten Feld-Schlüsselwort erfassen (NICHT an Leerzeilen abschneiden),
    # damit mehrzeilige Antworten vollständig erhalten bleiben. Fallback: bis Textende.
    antwort = _extract(
        r'^Antwort:\s*(.+?)(?=\n(?:Textgrundlage:|Schlüsselbegriffe:|Beleg:|Abdeckung:))',
        multiline=True,
    )
    if antwort == '–':
        antwort = _extract(r'^Antwort:\s*(.+)', multiline=True)

    return {
        'antwort':       antwort,
        'textgrundlage': _extract(r'^Textgrundlage:\s*(.+)'),
        'schluessel':    _extract(r'^Schlüsselbegriffe:\s*(.+)'),
        'beleg':         _extract(r'^Beleg:\s*(.+)', default=''),
        'abdeckung':     _extract(r'^Abdeckung:\s*(.+)'),
    }


def serialize_qa_items(items: list) -> str:
    """
    Serialisiert geparste QA-Items zurück ins Markdown-Format von verify_with_questions().
    Inverse zu parse_qa_response() – wird nach der Nachbearbeitung (Schritt 5b) genutzt,
    damit der Downstream-Pfad (Doc-Builder ruft parse_qa_response erneut auf) unverändert läuft.
    """
    def _fields(src: dict) -> list:
        ls = [
            f"**Antwort:** {src.get('antwort', '–')}",
            f"**Textgrundlage:** {src.get('textgrundlage', '–')}",
            f"**Schlüsselbegriffe:** {src.get('schluessel', '–')}",
        ]
        if src.get('beleg'):
            ls.append(f"**Beleg:** {src['beleg']}")
        ls.append(f"**Abdeckung:** {src.get('abdeckung', '–')}")
        return ls

    blocks = []
    for it in items:
        lines = [f"**Frage {it['num']}**"]
        if it.get('frage_text'):
            lines.append(f"**Fragetext:** {it['frage_text']}")
        if it.get('subs'):
            for sub in it['subs']:
                lines.append(f"**{sub.get('label', '–')}**")
                lines.extend(_fields(sub))
        else:
            lines.extend(_fields(it))
        blocks.append('\n'.join(lines))
    return '\n\n'.join(blocks)


def _add_comment_range_start(paragraph, comment_id: int):
    """Setzt w:commentRangeStart als erstes Kind eines Absatzes."""
    crs = OxmlElement('w:commentRangeStart')
    crs.set(qn('w:id'), str(comment_id))
    paragraph._p.insert(0, crs)


def _add_comment_range_end(paragraph, comment_id: int):
    """Setzt w:commentRangeEnd + w:commentReference ans Ende eines Absatzes."""
    p_elem = paragraph._p
    cre = OxmlElement('w:commentRangeEnd')
    cre.set(qn('w:id'), str(comment_id))
    p_elem.append(cre)
    run = OxmlElement('w:r')
    rpr = OxmlElement('w:rPr')
    rs = OxmlElement('w:rStyle')
    rs.set(qn('w:val'), 'CommentReference')
    rpr.append(rs)
    run.append(rpr)
    ref = OxmlElement('w:commentReference')
    ref.set(qn('w:id'), str(comment_id))
    run.append(ref)
    p_elem.append(run)


def _normalize_quote(text: str) -> str:
    """Normalisiert Text für Beleg-Matching: Markdown weg, Whitespace kollabiert, lowercase."""
    text = re.sub(r'[*_`"„“”»«\']', '', text)
    text = re.sub(r'\s+', ' ', text)
    return text.strip().lower()


def _find_para_by_quote(paras: list, quote: str):
    """
    Findet den Absatz in `paras`, dessen Text das (normalisierte) Beleg-Zitat enthält.
    Nutzt einen Präfix des Zitats (robust gegen leichte Abweichungen am Satzende).
    Gibt den Absatz zurück oder None.
    """
    nq = _normalize_quote(quote)
    if len(nq) < 8:
        return None
    needle = nq[:60]
    for p in paras:
        if not p.text.strip():
            continue
        if needle in _normalize_quote(p.text):
            return p
    return None


def _inject_comments_part(doc, comments_list: list):
    """
    Erstellt word/comments.xml und registriert es im OPC-Package.
    comments_list: [(id: int, text: str), ...]
    Muss VOR doc.save() aufgerufen werden.
    """
    WURI  = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    CT    = 'application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml'
    RT    = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments'
    WDATE = '2024-01-01T00:00:00Z'

    root = etree.Element(f'{{{WURI}}}comments',
                         nsmap={'w': WURI,
                                'wpc': 'http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas',
                                'r':   'http://schemas.openxmlformats.org/officeDocument/2006/relationships'})
    for cid, text in comments_list:
        c = etree.SubElement(root, f'{{{WURI}}}comment')
        c.set(f'{{{WURI}}}id',       str(cid))
        c.set(f'{{{WURI}}}author',   'Lernfragen')
        c.set(f'{{{WURI}}}date',     WDATE)
        c.set(f'{{{WURI}}}initials', 'LF')
        p = etree.SubElement(c, f'{{{WURI}}}p')
        r = etree.SubElement(p, f'{{{WURI}}}r')
        t = etree.SubElement(r, f'{{{WURI}}}t')
        t.text = text

    xml_bytes = etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True)
    try:
        part = Part(PackURI('/word/comments.xml'), CT, xml_bytes, doc.part.package)
        doc.part.relate_to(part, RT)
        print(f"   Word-Kommentare: {len(comments_list)} eingebettet")
    except Exception as e:
        print(f"   Kommentare konnten nicht eingebettet werden: {e}")


# Menschlich lesbare Namen für die gängigsten Sprachcodes (für Übersetzungs-Prompts).
LANGUAGE_NAMES = {
    "de": "Deutsch",
    "en": "Englisch",
    "fr": "Französisch",
    "es": "Spanisch",
    "it": "Italienisch",
}


def _language_name(code: str) -> str:
    """Liefert den deutschen Klarnamen zu einem Sprachcode (Fallback: Code selbst)."""
    return LANGUAGE_NAMES.get((code or "").lower(), code)


def detect_language(text: str) -> str:
    """Schritt 2a: Erkennt die Hauptsprache des Dokuments und liefert einen ISO-Code
    (z.B. 'de', 'en').

    Nimmt drei Stichproben (Anfang, Mitte, Ende), um Dokumente mit z.B. deutschem Titel
    aber englischem Haupttext korrekt zu erkennen. Bei endgültigem Fehler wird 'unknown'
    zurückgegeben – die aufrufende Logik übersetzt dann NICHT (sichere Annahme:
    Quellsprache = Zielsprache), statt blind auf Englisch/Übersetzung auszuweichen.
    """
    print("--- Schritt 2a: Erkenne Sprache des Dokuments ---")
    n = len(text)
    samples = [
        text[:1500],
        text[max(0, n // 2 - 750): n // 2 + 750],
        text[max(0, n - 1500):],
    ]
    leseprobe = "\n\n---\n\n".join(s for s in samples if s.strip())
    prompt = (
        "Bestimme die Hauptsprache des folgenden Textes. "
        "Antworte mit exakt einem ISO-639-1-Sprachcode in Kleinbuchstaben "
        "(z.B. 'de' für Deutsch, 'en' für Englisch). "
        "Ignoriere vereinzelte fremdsprachige Wörter (Eigennamen, Zitate) – "
        "frage nur nach der Hauptsprache. Gib nur den Code aus, sonst nichts.\n\n"
        f"Text:\n{leseprobe}"
    )
    try:
        response = call_gemini_with_retry(
            model_name='gemini-2.5-flash',
            contents=prompt,
            config=types.GenerateContentConfig(temperature=0.0, max_output_tokens=5)
        )
        code = re.sub(r'[^a-z]', '', response.text.strip().lower())[:2]
        if code:
            print(f"   Erkannte Sprache: {code}")
            return code
        print("   Sprachcode nicht erkennbar – behandle als 'unknown'.")
        return "unknown"
    except Exception as e:
        print(f"Sprachprüfung fehlgeschlagen ({e}); behandle als 'unknown' (keine Übersetzung).")
        return "unknown"


_ANY_HEADING_RE = re.compile(r'^#{1,6}\s')


def split_text_by_headings(text: str, max_chars: int = 15000) -> list:
    """Hilfsfunktion: Splittet Markdown-Text an Überschriften in logische Abschnitte.

    Erkennt Überschriften JEDER Tiefe (#-######) als möglichen Splitpunkt, nicht nur #/##:
    Dokumente mit uneinheitlicher Verschachtelung (z.B. Sammelbände, bei denen einzelne
    Kapitel nur ### statt # sind) könnten sonst ohne jeden Splitpunkt auf ein Vielfaches von
    max_chars anwachsen, bevor überhaupt ein Chunk-Ende möglich ist - mit dem Risiko, dass ein
    einzelner, zu großer Übersetzungs-Call Inhalte verliert.
    """
    chunks = []
    current_chunk = []
    current_length = 0

    for line in text.split('\n'):
        if _ANY_HEADING_RE.match(line) and current_length > max_chars:
            chunks.append('\n'.join(current_chunk))
            current_chunk = []
            current_length = 0

        current_chunk.append(line)
        current_length += len(line)

    if current_chunk:
        chunks.append('\n'.join(current_chunk))

    return chunks


def translate_text(text: str, source_lang: str = "en", target_lang: str = "de",
                   cache_dir: Path = None) -> str:
    """Schritt 2b: Übersetzt den Text abschnittsweise nach den strengen Regeln.

    source_lang/target_lang sind ISO-639-1-Codes; der Prompt wird daraus dynamisch gebaut,
    damit nie versehentlich in die falsche Richtung (z.B. Deutsch → Englisch) übersetzt wird.

    Wird cache_dir übergeben, wird jeder fertige Abschnitt einzeln zwischengespeichert
    (de_uebersetzung_chunk_NN.md). Ein Abbruch/Neustart übersetzt dann nur die noch
    fehlenden Abschnitte neu, statt von vorne zu beginnen.
    """
    src_name = _language_name(source_lang)
    tgt_name = _language_name(target_lang)
    print(f"--- Schritt 2b: Übersetze Text von {src_name} nach {tgt_name} (via Gemini 2.5 Pro) ---")

    system_prompt = (
        f"Übersetze den folgenden {src_name.lower()}en wissenschaftlichen Text "
        f"originalgetreu ins {tgt_name}.\n\n"
        "Ziel:\n"
        "Eine vollständige, sinntreue Übersetzung, keine Zusammenfassung.\n\n"
        "Strenge Regeln:\n"
        "- Nichts auslassen.\n"
        "- Nichts ergänzen.\n"
        "- Nichts interpretieren.\n"
        "- Keine Inhalte glätten, kürzen oder zusammenfassen.\n"
        "- Fachbegriffe konsistent übersetzen.\n"
        "- Überschriften, Absatzstruktur, Listen und Tabellenstruktur beibehalten.\n"
        "- Zitate, Autorennamen, Jahreszahlen, Variablennamen, Skalen, Hypothesen und statistische Angaben exakt erhalten.\n"
        "- Unklare oder beschädigte Stellen mit [UNKLAR: Originalstelle] markieren, nicht erraten.\n"
        "- Bildverweise, Tabellenverweise und Abbildungsbeschriftungen erhalten.\n"
        "- Markdown-Struktur beibehalten.\n\n"
        "Ausgabeformat:\n"
        "Nur die deutsche Übersetzung, kein Kommentar, keine Kontrollliste danach."
    )

    chunks = split_text_by_headings(text)
    translated_chunks = []

    for i, chunk in enumerate(chunks, 1):
        chunk_cache = cache_dir / f"de_uebersetzung_chunk_{i:02d}.md" if cache_dir else None
        if chunk_cache and chunk_cache.exists():
            print(f"   -> [SKIP] Abschnitt {i} von {len(chunks)} – bereits übersetzt.")
            translated_chunks.append(chunk_cache.read_text(encoding="utf-8"))
            continue
        print(f"   -> Übersetze Abschnitt {i} von {len(chunks)}...")
        try:
            response = call_gemini_with_retry(
                model_name='gemini-2.5-pro',
                contents=f"Text:\n{chunk}",
                config=types.GenerateContentConfig(system_instruction=system_prompt, temperature=0.1)
            )
            translated_chunks.append(response.text)
            if chunk_cache:
                chunk_cache.write_text(response.text, encoding="utf-8")
            time.sleep(2)
        except Exception as e:
            print(f"Fehler bei der Übersetzung von Abschnitt {i}: {e}")
            raise

    result = '\n\n'.join(translated_chunks)
    # Alle Abschnitte erfolgreich → Einzel-Caches aufräumen (die vollständige
    # de_uebersetzung.md übernimmt ab jetzt das Caching).
    if cache_dir:
        for f in cache_dir.glob("de_uebersetzung_chunk_*.md"):
            f.unlink()
    return result


def split_into_level1_chapters(text: str) -> list:
    """
    Splittet den (bereits normalisierten) Markdown-Text an nummerierten Kapitelüberschriften.
    Adaptiv: Bei einem einzigen Top-Level-Kapitel (z.B. extrahiertes Einzelkapitel) wird eine
    Ebene tiefer gesplittet, sodass Unterkapitel als eigene Chunks erkannt werden.
    Jeder Eintrag: {'heading': str, 'full_text': str}.
    """
    lines = text.split('\n')
    _html_re = re.compile(r'<[^>]+>')

    def _numbered_level(line: str):
        """Gibt den #-Level zurück wenn die Zeile eine nummerierte Überschrift ist (HTML-ignorant)."""
        stripped = _html_re.sub('', line).replace('**', '')
        m = re.match(r'^(#{1,6})\s+(\d+(?:\.\d+)*)\s', stripped)
        return len(m.group(1)) if m else None

    numbered_levels = [nl for line in lines if (nl := _numbered_level(line)) is not None]

    if not numbered_levels:
        return []

    min_level = min(numbered_levels)
    top_count = sum(1 for l in numbered_levels if l == min_level)
    # Einziges Top-Kapitel → eine Ebene tiefer splitten (z.B. extrahiertes Kapitel 4)
    split_level = min_level + 1 if top_count == 1 else min_level

    chapters = []
    current_heading = None
    current_lines = []
    pre_split_lines: list[str] = []

    for line in lines:
        if _numbered_level(line) == split_level:
            if current_heading is not None:
                chapters.append({'heading': current_heading, 'full_text': '\n'.join(current_lines), 'level': split_level})
            stripped = _html_re.sub('', line).replace('**', '')
            current_heading = re.sub(r'^#{1,6}\s+', '', stripped).strip()
            current_lines = [line]
        else:
            if current_heading is None:
                pre_split_lines.append(line)
            else:
                current_lines.append(line)

    if current_heading is not None:
        chapters.append({'heading': current_heading, 'full_text': '\n'.join(current_lines), 'level': split_level})

    # Im adaptiven Modus (split_level = min_level+1) landet Text vor dem ersten Unterkapitel
    # (z.B. Einleitungstext nach "# 5 Kompetenzen..." vor "## 5.1 ...") in pre_split_lines.
    # Diesen ans erste Kapitel hängen, damit er mitverdichtet wird.
    if pre_split_lines and chapters:
        preamble_str = '\n'.join(pre_split_lines).strip()
        if preamble_str:
            # Erste Überschriftenzeile im Vorspann = Kapitel-Root (z.B. "3 Analyse …").
            _root_idx = next((k for k, ln in enumerate(pre_split_lines)
                              if re.match(r'^#{1,6}\s', _html_re.sub('', ln).replace('**', ''))), None)
            _root_body = '\n'.join(pre_split_lines[_root_idx + 1:]).strip() if _root_idx is not None else ''
            if _root_idx is not None and len(_root_body) > 300:
                # Kapitel-Root mit eigenem Einleitungstext → eigenes Kapitel, damit seine
                # Zusammenfassung dem Root-Heading zugeordnet werden kann (sonst unsichtbar in Kap. 1).
                _root_heading = re.sub(r'^#{1,6}\s+', '',
                                       _html_re.sub('', pre_split_lines[_root_idx]).replace('**', '')).strip()
                chapters.insert(0, {'heading': _root_heading,
                                    'full_text': preamble_str,
                                    'level': split_level - 1})
            else:
                chapters[0]['full_text'] = preamble_str + '\n\n' + chapters[0]['full_text']

    return chapters


_ROMAN_LABEL_RE = re.compile(r'^\**\s*[IVXLCDM]+\.\s+\S')
_DIGIT_LABEL_RE = re.compile(r'^\**\s*\d+\.\s+\S')


def split_into_articles(text: str) -> list:
    """
    Splittet Markdown-Text in einzelne Artikel für Sammelbände/Festschriften. Primär anhand
    römisch nummerierter Überschriften-Labels ("I. Titel"), UNABHÄNGIG von der #-Tiefe - notwendig,
    weil einzelne Artikel eines Sammelbands oft uneinheitlich tief verschachtelt sind (z.B. manche
    Artikel als "#", andere als "###", je nach OCR/Layout der Originalseite),
    split_into_level1_chapters aber nur Ziffern erkennt und nur eine einzige Ebene splittet.
    Ziffern-Labels (z.B. "0. Vorwort") zählen nur als Artikelgrenze, wenn sie auf derselben
    #-Tiefe wie mindestens eine römische Überschrift liegen - sonst würden normale nummerierte
    Unterabschnitte innerhalb eines Artikels (z.B. "1. Einleitung", "2. Prozessschritte") fälschlich
    als eigene Artikel erkannt. Kein Kollisionsrisiko mit Fließtext (z.B. "I. Definition" als
    Aufzählungspunkt), da nur echte Markdown-Heading-Zeilen (beginnend mit #) geprüft werden.
    Fallback ohne jedes römische Label (z.B. Sammelband ganz ohne Nummerierung): flachste
    Heading-Ebene.
    Jeder Eintrag: {'heading': str, 'full_text': str, 'level': int}.
    """
    lines = text.split('\n')
    _html_re = re.compile(r'<[^>]+>')

    def _heading_level(line: str):
        m = re.match(r'^(#{1,6})\s+(\S.*)$', line)
        return len(m.group(1)) if m else None

    def _clean_heading_line(line: str, level: int) -> str:
        stripped = _html_re.sub('', line).replace('**', '')
        return re.sub(r'^#{1,6}\s+', '', stripped).strip()

    all_levels = [(i, lv) for i, line in enumerate(lines) if (lv := _heading_level(line)) is not None]
    roman_levels = {lv for i, lv in all_levels if _ROMAN_LABEL_RE.match(_clean_heading_line(lines[i], lv))}

    if roman_levels:
        def is_split_point(i, lv):
            cleaned = _clean_heading_line(lines[i], lv)
            if _ROMAN_LABEL_RE.match(cleaned):
                return True
            return lv in roman_levels and bool(_DIGIT_LABEL_RE.match(cleaned))
    else:
        if not all_levels:
            return []
        split_level = min(lv for _, lv in all_levels)
        is_split_point = lambda i, lv: lv == split_level

    chapters = []
    current_heading = None
    current_level = None
    current_lines: list = []
    pre_lines: list = []

    for i, line in enumerate(lines):
        lv = _heading_level(line)
        if lv is not None and is_split_point(i, lv):
            if current_heading is not None:
                chapters.append({'heading': current_heading, 'full_text': '\n'.join(current_lines), 'level': current_level})
            current_heading = _clean_heading_line(line, lv)
            current_level = lv
            current_lines = [line]
        else:
            if current_heading is None:
                pre_lines.append(line)
            else:
                current_lines.append(line)

    if current_heading is not None:
        chapters.append({'heading': current_heading, 'full_text': '\n'.join(current_lines), 'level': current_level})

    # Vorspann vor der ersten Artikelüberschrift (z.B. Titelseite/Vorwort) an den ersten
    # Artikel anhängen, statt ihn stillschweigend zu verwerfen.
    pre_str = '\n'.join(pre_lines).strip()
    if pre_str and chapters:
        chapters[0]['full_text'] = pre_str + '\n\n' + chapters[0]['full_text']

    return chapters


def _split_at_level2(chapter_text: str) -> tuple[str, list[dict]]:
    """
    Teilt einen Kapiteltext an ## Überschriften (genau 2 Rauten, nicht ###) auf.
    Gibt zurück: (preamble_text, [{'heading': str, 'text': str}, ...])
    Der preamble ist der Text vor der ersten ## Überschrift.
    """
    lines = chapter_text.split('\n')
    preamble_lines: list[str] = []
    sections: list[dict] = []
    current_heading: str | None = None
    current_lines: list[str] = []

    for line in lines:
        m = re.match(r'^##\s+(.+)$', line)
        if m:
            if current_heading is not None:
                sections.append({'heading': current_heading, 'text': '\n'.join(current_lines)})
            elif current_lines:
                preamble_lines = current_lines[:]
            current_heading = m.group(1).strip()
            current_lines = [line]
        else:
            current_lines.append(line)

    if current_heading is not None:
        sections.append({'heading': current_heading, 'text': '\n'.join(current_lines)})
    elif current_lines and not preamble_lines:
        preamble_lines = current_lines

    return '\n'.join(preamble_lines), sections


def _split_at_level(chapter_text: str, level: int) -> tuple[str, list[dict]]:
    """Teilt text an Überschriften der angegebenen #-Tiefe. Gibt (preamble, sections) zurück."""
    prefix = '#' * level + ' '
    avoid_prefix = '#' * level + '##'
    lines = chapter_text.split('\n')
    preamble_lines: list[str] = []
    sections: list[dict] = []
    current_heading: str | None = None
    current_lines: list[str] = []

    for line in lines:
        if line.startswith(prefix) and not line.startswith(avoid_prefix):
            if current_heading is not None:
                sections.append({'heading': current_heading, 'text': '\n'.join(current_lines)})
            elif current_lines:
                preamble_lines = current_lines[:]
            current_heading = line[len(prefix):].strip()
            current_lines = [line]
        else:
            current_lines.append(line)

    if current_heading is not None:
        sections.append({'heading': current_heading, 'text': '\n'.join(current_lines)})
    elif current_lines and not preamble_lines:
        preamble_lines = current_lines

    return '\n'.join(preamble_lines), sections


def _detect_chapter_level(text: str) -> int:
    """Gibt den #-Level der ersten Heading-Zeile zurück (Fallback: 2)."""
    for line in text.splitlines():
        m = re.match(r'^(#{1,6})\s', line)
        if m:
            return len(m.group(1))
    return 2


def _find_sublevel(text: str, chapter_heading_level: int) -> int | None:
    """Findet die nächste Heading-Ebene unterhalb chapter_heading_level, oder None."""
    for level in range(chapter_heading_level + 1, 7):
        prefix = '#' * level + ' '
        avoid = '#' * level + '##'
        if any(line.startswith(prefix) and not line.startswith(avoid) for line in text.splitlines()):
            return level
    return None


def _group_into_chunks(preamble: str, sections: list[dict], max_chars: int = 15_000) -> list[dict]:
    """
    Fasst sections greedy zu Chunks von max max_chars Zeichen zusammen.
    Sections werden NIE in der Mitte geteilt – immer vollständig in einem Chunk.
    Preamble geht immer in den ersten Chunk.
    """
    chunks: list[dict] = []
    current_secs: list[dict] = []
    current_len = len(preamble)

    for sec in sections:
        sec_len = len(sec['text'])
        if current_secs and current_len + sec_len > max_chars:
            chunks.append({'preamble': preamble if not chunks else '', 'sections': current_secs})
            current_secs = []
            current_len = 0
        current_secs.append(sec)
        current_len += sec_len

    if current_secs:
        chunks.append({'preamble': preamble if not chunks else '', 'sections': current_secs})
    elif not chunks:
        chunks.append({'preamble': preamble, 'sections': []})

    return chunks


def _summarize_chapter_by_sections(ch: dict, out_dir: Path) -> str:
    """
    Für große Kapitel (> 3 Level-2-Sections): jede Section einzeln zusammenfassen.
    Caching: je Section eine eigene Datei zusammenfassung_kap_XX_YY.md.
    Kombinierter Output wird immer frisch aus Sub-Caches gebaut (kein äußeres load_or_run).
    """
    i = ch['index']
    preamble, sections = _split_at_level2(ch['full_text'])
    parts: list[str] = []

    if preamble.strip():
        preamble_file = out_dir / f"zusammenfassung_kap_{i:02d}_00.md"
        needs_call = not preamble_file.exists()
        preamble_summary = load_or_run(
            preamble_file,
            lambda: _summarize_single_chapter(ch['heading'], preamble),
            f"Zusammenfassung Kap. {i} Einleitung: {ch['heading'][:50]}"
        )
        parts.append(preamble_summary)
        if needs_call:
            time.sleep(1)

    for j, sec in enumerate(sections, start=1):
        sec_file = out_dir / f"zusammenfassung_kap_{i:02d}_{j:02d}.md"
        needs_call = not sec_file.exists()
        sec_summary = load_or_run(
            sec_file,
            lambda s=sec: _summarize_single_chapter(s['heading'], s['text']),
            f"Zusammenfassung Kap. {i}.{j:02d}: {sec['heading'][:50]}"
        )
        if not re.match(r'^#+\s', sec_summary.strip()):
            sec_summary = f"## {sec['heading']}\n\n{sec_summary}"
        parts.append(sec_summary)
        if needs_call:
            time.sleep(1)

    return '\n\n'.join(parts)


def _summarize_single_chapter(heading: str, chapter_text: str, output_lang: str = "de") -> str:
    """Erstellt eine lernorientierte Zusammenfassung für ein einzelnes Kapitel.

    output_lang (ISO-Code) erzwingt die Ausgabesprache: Auch fremdsprachige Passagen im
    Original (z.B. ein englisches Abstract) werden in dieser Sprache zusammengefasst, statt
    in der Originalsprache zu verbleiben.
    """
    lang_name = _language_name(output_lang)
    prompt = (
        f"Erstelle eine lernorientierte Zusammenfassung für das folgende Kapitel: \"{heading}\"\n\n"
        f"AUSGABESPRACHE: Der gesamte FLIESSTEXT (Stichpunkte, Definitionen, Beschreibungen) MUSS "
        f"auf {lang_name} verfasst sein. Liegt ein Teil des Originals in einer anderen Sprache vor "
        f"(z.B. ein englisches Abstract), übersetze dessen INHALT in der Zusammenfassung nach {lang_name}.\n"
        f"WICHTIG: Die ÜBERSCHRIFTEN dagegen EXAKT und unverändert aus dem Original übernehmen "
        f"(gleicher Wortlaut, gleiche Sprache, gleiche Nummerierung) – NICHT übersetzen und NICHT "
        f"umformulieren. Sie dienen als Zuordnungsschlüssel und müssen 1:1 mit dem Original übereinstimmen.\n\n"
        "Pflichtanforderungen:\n"
        f"1. ALLE Unterkapitel müssen vorhanden sein – kein einziges Unterkapitel darf fehlen!\n"
        f"   Übernimm jede Überschrift exakt wie im Original (Wortlaut inkl. evtl. Nummerierung, "
        f"z.B. '4.1.1 Hedonisches Wohlbefinden').\n"
        "2. Pro Unterkapitel: mindestens 3–5 Stichpunkte mit den wichtigsten Inhalten.\n"
        "3. Studienergebnisse IMMER erhalten: Metaanalysen, Effektstärken, Befundrichtung, Autoren & Jahr.\n"
        "4. Definitionen: wörtlich oder sehr nah am Original übernehmen.\n"
        "5. Keine neuen Informationen ergänzen.\n"
        "6. Abbildungen und Tabellen kurz erwähnen und ihren Inhalt beschreiben.\n"
        "7. Stichpunkte statt Fließtext (Ausnahme: Definitionen).\n"
        "8. Länge: maximal 40–50 % des Originals – ABER vollständige Unterkapitelabdeckung hat Vorrang vor Kürze.\n"
        "9. Kästen (Fokus/Studie/Definition/Beispiel) sind eigenständige Lernobjekte – "
        "fasse jeden als abgeschlossene Einheit unter seiner eigenen Überschrift zusammen.\n\n"
        "Selbstprüfung (am Ende anhängen):\n"
        "- Liste alle Unterkapitel des Originals auf\n"
        "- Markiere fehlende Unterkapitel oder fehlende Studienergebnisse\n\n"
        f"Kapiteltext:\n{chapter_text}"
    )
    # Pflichtliste aller Unterkapitel: alle Headings außer dem Kapitel-Heading selbst.
    _all_headings = re.findall(r'^#{2,6}\s+\*?\*?(.+?)\*?\*?\s*$', chapter_text, re.MULTILINE)
    _heading_norm = normalize_heading(heading)
    _required = [h for h in _all_headings if normalize_heading(h) != _heading_norm]
    if len(_required) >= 1:
        _rlist = '\n'.join(f'- {_h}' for _h in _required[:80])
        prompt += (
            f"\n\nPFLICHT-VOLLSTÄNDIGKEIT – folgende {len(_required)} Unterabschnitte MÜSSEN ALLE "
            f"als eigene Überschrift erscheinen:\n{_rlist}\n"
            "Jeder Abschnitt benötigt mindestens 1 Stichpunkt. Keiner darf fehlen!\n"
        )
    # Sehr große Kapitel (>50 Abschnitte, z.B. 263 Kompetenzen): strenge Kompression nötig
    # damit alle Abschnitte in den Output passen.
    # 263 Abschnitte × 3 Bullets × 15 Wörter ≈ 12.000 Wörter → passt in 65.536 Output-Token.
    if len(_required) > 50:
        prompt += (
            f"\n\nWICHTIG FÜR DIESES SEHR GROSSE KAPITEL ({len(_required)} Abschnitte):\n"
            "- Maximal 3 Stichpunkte pro Unterabschnitt (Ausnahme: Definition immer vollständig)\n"
            "- Maximal 20 Wörter pro Stichpunkt\n"
            "- Vollständigkeit (alle Abschnitte vorhanden) hat ABSOLUTEN Vorrang vor Ausführlichkeit\n"
            "- Keinen Abschnitt auslassen, lieber sehr kurz als gar nicht!\n"
        )
    # Sehr große Kapitel brauchen 65536 Output-Token; normale Kapitel 32768 für ausführliche Summaries.
    _out_tokens = 65536 if len(_required) > 50 else 32768
    try:
        response = call_gemini_with_retry(
            model_name='gemini-2.5-pro',
            contents=prompt,
            config=types.GenerateContentConfig(
                temperature=0.2,
                max_output_tokens=_out_tokens
            )
        )
        return response.text
    except Exception as e:
        print(f"Fehler bei Zusammenfassung von '{heading}': {e}")
        raise


def _summarize_article_condensed(heading: str, article_text: str, output_lang: str = "de") -> str:
    """
    Erstellt eine VERDICHTETE Kurzübersicht für einen einzelnen Artikel eines Sammelbands
    (z.B. Festschrift-Beitrag): worum es geht, zentrale Erkenntnisse/Schlussfolgerungen, plus
    2-3 wörtliche Zitate. Im Gegensatz zu _summarize_single_chapter KEINE Pflicht zur
    vollständigen Unterkapitelabdeckung - der Artikel wird bewusst nicht vollständig
    aufbereitet, da er nicht prüfungsrelevant ist (verdichteter Modus, --condensed).
    """
    lang_name = _language_name(output_lang)
    prompt = (
        f"Erstelle eine VERDICHTETE Kurzübersicht für den folgenden Artikel: \"{heading}\"\n\n"
        f"AUSGABESPRACHE: Der Fließtext (Zusammenfassung) MUSS auf {lang_name} verfasst sein. "
        f"Die ÜBERSCHRIFT dagegen EXAKT und unverändert aus dem Original übernehmen (gleicher "
        f"Wortlaut, gleiche Sprache) - sie dient als Zuordnungsschlüssel und muss 1:1 mit dem "
        f"Original übereinstimmen.\n\n"
        "Anforderungen:\n"
        "1. 3-5 Stichpunkte: Worum geht es im Artikel (Thema, Fragestellung, Perspektive des "
        "Autors) und was wurde erkannt bzw. gefolgert (zentrale Ergebnisse, Thesen, "
        "Schlussfolgerungen).\n"
        "2. KEINE vollständige Unterkapitelabdeckung nötig - nur die wesentliche Kernaussage, "
        "so kurz wie möglich.\n"
        "3. Danach ein eigener Block '**Zitate:**' mit genau 2-3 kurzen, WÖRTLICHEN Zitaten aus "
        "dem Original (im Originalwortlaut, NICHT übersetzt, NICHT paraphrasiert), die die "
        "Kernaussagen des Artikels am besten belegen. Format exakt:\n"
        "**Zitate:**\n"
        "1. \"...\"\n"
        "2. \"...\"\n"
        "3. \"...\"\n"
        "4. Keine neuen Informationen ergänzen, keine Bewertung.\n\n"
        f"Artikeltext:\n{article_text}"
    )
    try:
        response = call_gemini_with_retry(
            model_name='gemini-2.5-pro',
            contents=prompt,
            config=types.GenerateContentConfig(temperature=0.2, max_output_tokens=8192)
        )
        return response.text
    except Exception as e:
        print(f"Fehler bei verdichteter Zusammenfassung von '{heading}': {e}")
        raise


def _extract_quotes_block(condensed_text: str) -> tuple[str, str]:
    """
    Trennt einen verdichteten Artikel-Block in (sichtbarer_Zusammenfassungstext, Zitate_Block).
    Erkennt den '**Zitate:**'-Marker (auch als '## Zitate' o.ä.); ohne Treffer ist der
    Zitate-Block leer und der gesamte Text gilt als sichtbare Zusammenfassung.
    """
    m = re.search(r'(?im)^\s*#{0,6}\s*\**\s*Zitate\s*:?\**\s*$', condensed_text)
    if not m:
        return condensed_text.strip(), ''
    summary_part = condensed_text[:m.start()].strip()
    quotes_part = condensed_text[m.end():].strip()
    return summary_part, quotes_part


def _ensure_summary_heading(summary: str, heading: str) -> str:
    """Stellt sicher, dass eine Kapitel-Zusammenfassung mit der zugehörigen Überschrift beginnt,
    damit der Docx-Builder sie per sum_lookup der Original-Sektion zuordnen kann.
    Hängt '## {heading}' nur voran, wenn die erste Überschrift fehlt oder eine andere ist –
    so erhält der Vorspann (Einleitung vor dem ersten Unterkapitel) eine sichtbare Summary."""
    norm = normalize_heading(_clean_heading_text(heading))
    for ln in summary.lstrip().splitlines():
        if re.match(r'^#{1,6}\s', ln):
            if normalize_heading(_clean_heading_text(ln)) == norm:
                return summary  # passendes Heading bereits vorhanden
            break  # erstes Heading gehört zu einem Unterabschnitt → Kapitel-Heading davorsetzen
        if ln.strip():
            break  # Fließtext vor erstem Heading → Heading davorsetzen
    return f"## {heading}\n\n{summary.lstrip()}"


def generate_summary_by_chapter(text: str, out_dir: Path, output_lang: str = "de") -> str:
    """
    Schritt 4: Erstellt die Zusammenfassung kapitelweise (je Level-1-Kapitel ein API-Call).
    Nutzt Caching pro Kapitel: zusammenfassung_kap_XX.md in out_dir.
    Kombiniert am Ende zu zusammenfassung.md.
    output_lang erzwingt die Ausgabesprache der Zusammenfassung.
    """
    print("--- Schritt 4: Erstelle kapitelweise Zusammenfassung (via Gemini 2.5 Pro) ---")

    text = normalize_heading_levels(text)

    # Preamble-Text (vor erstem nummerierten Heading) extrahieren.
    # split_into_level1_chapters() verwirft diesen Text sonst kommentarlos.
    # Bold-Marker für die Nummern-Erkennung ignorieren ('#### **3.1 …**' ist nummeriert).
    _lines_nb = text.split('\n')
    _first_num_idx = next(
        (k for k, ln in enumerate(_lines_nb)
         if re.match(r'^#{1,6}\s+\d', re.sub(r'<[^>]+>', '', ln).replace('**', ''))),
        None,
    )
    _preamble_text = ('\n'.join(_lines_nb[:_first_num_idx]).strip()
                      if _first_num_idx not in (None, 0) else '')

    chapters = split_into_level1_chapters(text)

    # Preamble nur dann an das erste Kapitel hängen, wenn split_into_level1_chapters den
    # Vorspann verworfen hat (erstes Kapitel ist nummeriert). Hat es bereits ein unnummeriertes
    # Wurzel-Kapitel (Kapiteltitel) erzeugt, steckt der Vorspann dort schon drin → kein Duplikat.
    _root_unnumbered = bool(
        chapters and not re.match(r'^\d', normalize_heading(chapters[0]['heading']))
    )
    if _preamble_text and chapters and not _root_unnumbered:
        chapters[0]['full_text'] = _preamble_text + '\n\n' + chapters[0]['full_text']

    _CHUNK_LIMIT = 15_000  # ~30 Seiten; Kapitel über diesem Limit werden chunk-weise verarbeitet

    def _summarize_chapter_smart(ch: dict, ci: int) -> str:
        """Fasst ein Kapitel zusammen: direkt wenn klein, chunk-weise wenn groß."""
        # 'level' wird von split_into_level1_chapters gesetzt; Fallback für manuell erstellte Dicts.
        chap_level = ch.get('level', _detect_chapter_level(ch['full_text']))
        sublevel = _find_sublevel(ch['full_text'], chap_level)

        if sublevel is None or len(ch['full_text']) <= _CHUNK_LIMIT:
            return _summarize_single_chapter(ch['heading'], ch['full_text'], output_lang)

        preamble, sections = _split_at_level(ch['full_text'], sublevel)
        chunks = _group_into_chunks(preamble, sections, _CHUNK_LIMIT)

        if len(chunks) <= 1:
            return _summarize_single_chapter(ch['heading'], ch['full_text'], output_lang)

        print(f"   Kapitel {ci} ({ch['heading'][:40]}) → {len(chunks)} Chunks à max {_CHUNK_LIMIT} Zeichen")
        parts: list[str] = []
        for j, chunk in enumerate(chunks, start=1):
            chunk_file = out_dir / f"zusammenfassung_kap_{ci:02d}_{j:02d}.md"
            chunk_body = '\n\n'.join(s['text'] for s in chunk['sections'])
            chunk_text = (chunk['preamble'] + '\n\n' + chunk_body).strip()
            is_first = (j == 1)
            continuation_hint = (
                '' if is_first else
                f'\n\nHINWEIS: Dies ist Teil {j} von {len(chunks)} dieses Kapitels. '
                'Beginne direkt mit den Unterkapitelüberschriften ohne das übergeordnete Kapitelheading zu wiederholen.'
            )

            def _chunk_gen(ct=chunk_text, hd=ch['heading'], hint=continuation_hint):
                return _summarize_single_chapter(hd, ct + hint, output_lang)

            chunk_label = f"Zusammenfassung Kap. {ci} Teil {j}/{len(chunks)}: {ch['heading'][:40]}"
            chunk_summary = load_or_run(chunk_file, _chunk_gen, chunk_label)

            # Doppeltes Kapitel-Heading aus Chunks 2+ entfernen (falls KI es wiederholt)
            if not is_first and chunk_summary.strip():
                first_line = chunk_summary.strip().splitlines()[0]
                if re.match(r'^#{1,3}\s', first_line) and ch['heading'][:15] in first_line:
                    chunk_summary = chunk_summary.strip()[len(first_line):].lstrip('\n')

            parts.append(chunk_summary)
            time.sleep(1)

        return '\n\n'.join(parts)

    if not chapters:
        print("   Keine Level-1-Kapitel gefunden, fasse Gesamttext zusammen...")
        fallback_path = out_dir / "zusammenfassung_kap_00.md"
        # Die Artikel-Einleitung erhält bereits in main() via insert_intro_heading() eine
        # eigene "## Einleitung"-Überschrift; der frühere Top-Prepend-Hack entfällt damit
        # (er platzierte das Heading fälschlich ganz oben vor dem Journal-Namen).
        fallback_ch = {'heading': 'Volltext', 'full_text': text, 'level': 1}
        result = load_or_run(
            fallback_path,
            lambda: _summarize_chapter_smart(fallback_ch, 0),
            "Zusammenfassung (Volltext)"
        )
        (out_dir / "zusammenfassung.md").write_text(result, encoding="utf-8")
        return result

    summaries = []
    for i, chapter in enumerate(chapters, 1):
        cache_path = out_dir / f"zusammenfassung_kap_{i:02d}.md"
        label = f"Zusammenfassung Kap. {i}: {chapter['heading'][:60]}"
        chapter_summary = load_or_run(
            cache_path,
            lambda ch=chapter, ci=i: _summarize_chapter_smart(ch, ci),
            label
        )
        # Unnummeriertes Wurzel-Kapitel (Kapiteltitel-Vorspann/Einleitung): Heading erzwingen,
        # damit die Einleitungs-Zusammenfassung sicher unter der Kapitelüberschrift erscheint
        # (Regel: Vorspanntext ⇒ sichtbare Zusammenfassung).
        if i == 1 and not re.match(r'^\d', normalize_heading(chapter['heading'])):
            chapter_summary = _ensure_summary_heading(chapter_summary, chapter['heading'])
        summaries.append(chapter_summary)
        time.sleep(1)

    combined = "\n\n".join(summaries)
    (out_dir / "zusammenfassung.md").write_text(combined, encoding="utf-8")
    print(f"   Zusammenfassung aus {len(chapters)} Kapiteln kombiniert.")
    return combined


def generate_condensed_summary(text: str, out_dir: Path, output_lang: str = "de") -> str:
    """
    Verdichteter Modus (--condensed): Erstellt für jeden Artikel eines Sammelbands
    (per split_into_articles erkannt, auch römisch nummerierte Überschriften) eine kurze
    Kernaussagen-Übersicht + 2-3 wörtliche Zitate statt einer vollständigen, kapitelweisen
    Zusammenfassung. Ein API-Call pro Artikel, Caching pro Artikel: verdichtung_artikel_XX.md.
    Kombiniert am Ende zu verdichtung.md.
    """
    print("--- Verdichteter Modus: Erstelle Kurzübersicht je Artikel (via Gemini 2.5 Pro) ---")

    text = normalize_heading_levels(text)
    articles = split_into_articles(text)

    if not articles:
        print("   Keine Artikelüberschriften gefunden, verdichte Gesamttext...")
        articles = [{'heading': 'Volltext', 'full_text': text, 'level': 1}]

    summaries = []
    for i, article in enumerate(articles, 1):
        cache_path = out_dir / f"verdichtung_artikel_{i:02d}.md"
        label = f"Verdichtung Artikel {i}: {article['heading'][:60]}"
        article_summary = load_or_run(
            cache_path,
            lambda a=article: _summarize_article_condensed(a['heading'], a['full_text'], output_lang),
            label
        )
        article_summary = _ensure_summary_heading(article_summary, article['heading'])
        summaries.append(article_summary)
        time.sleep(1)

    combined = "\n\n".join(summaries)
    (out_dir / "verdichtung.md").write_text(combined, encoding="utf-8")
    print(f"   Verdichtung aus {len(articles)} Artikeln kombiniert.")
    return combined


def verify_with_questions(summary_text: str, questions_path: str) -> str:
    """Schritt 5: Qualitätssicherung der Zusammenfassung anhand von Leitfragen."""
    print(f"--- Schritt 5: Qualitätssicherung via Leitfragen aus {questions_path} ---")

    with open(questions_path, "r", encoding="utf-8") as f:
        questions = f.read()

    prompt = (
        "Rolle:\nDu bist Lerncoach und Prüfer für Wirtschaftspsychologie.\n\n"
        "Aufgabe:\nBeantworte die leseleitenden Fragen ultrakompakt und mit exakt einem Unterkapitelverweis. "
        "Nutze ausschließlich den hochgeladenen Text als Wissensbasis.\n\n"
        "Bevor du antwortest:\n"
        "Schritt 1: Suche die relevanten Stellen im Dokument.\n"
        "Schritt 2: Liste die Textstellen stichpunktartig auf.\n"
        "Schritt 3: Erst danach beantworte die Frage.\n\n"
        "Wenn keine passende Stelle existiert:\n'Im Dokument nicht enthalten'. Nicht raten.\n\n"
        "Mehrteilige Fragen:\n"
        "- Enthält eine Frage eine Aufzählung von Unterpunkten (z.B. mehrere Methoden/Begriffe, "
        "oft mit 'o' oder '-' aufgelistet), beantworte JEDEN Unterpunkt EINZELN.\n"
        "- Schreibe pro Unterpunkt zuerst seinen Namen als **fette Überschrift** in einer eigenen "
        "Zeile, darunter das normale Antwortformat (Antwort/Textgrundlage/Schlüsselbegriffe/Beleg/"
        "Abdeckung). Kein Unterpunkt darf fehlen.\n\n"
        "Antwortregeln:\n"
        "- Maximal 3 Sätze pro Frage bzw. pro Unterpunkt.\n"
        "- Keine Einleitung.\n"
        "- Die Frage selbst NICHT in der Antwort wiederholen (sie steht bereits in der "
        "'Fragetext:'-Zeile).\n"
        "- Keine ausführlichen Erklärungen.\n"
        "- Nur prüfungsrelevante Kernaussage.\n"
        "- Wenn Zahlen/Studienwerte relevant sind: nennen.\n"
        "- Wenn die Antwort im Dokument nicht eindeutig steht: „Im Dokument nicht eindeutig beantwortbar.“\n\n"
        "Quellenregeln:\n"
        "- Verweise immer auf die genaueste vorhandene Überschrift.\n"
        "- Nicht nur „Kapitel 6.4“, sondern z. B. „6.4.1.2 Eine umfassende Übersicht“.\n"
        "- Wenn mehrere Unterkapitel nötig sind, maximal 3 nennen.\n"
        "- Zusätzlich 1–3 Schlüsselbegriffe aus die Textstelle nennen.\n"
        "- Keine groben Kapitelverweise, wenn Unterkapitel vorhanden sind.\n"
        "- Beleg: Ein wörtliches Zitat (5–15 Wörter), exakt aus der Wissensbasis kopiert "
        "– die Kernstelle, die die Frage belegt. Keine Paraphrase, keine Auslassungszeichen.\n\n"
        "Ausgabeformat pro Frage:\n"
        "Frage X\n"
        "Fragetext: [die Frage wörtlich aus der Fragenliste übernehmen]\n"
        "Antwort: [max. 3 Sätze]\n"
        "Textgrundlage: [genaues Unterkapitel]\n"
        "Schlüsselbegriffe: [1–3 Begriffe]\n"
        "Beleg: [wörtliches Zitat aus der Zusammenfassung, 5–15 Wörter]\n"
        "Abdeckung: vollständig / teilweise / nicht enthalten\n\n"
        "WICHTIG zum Format:\n"
        "- Nummeriere die Fragen fortlaufend als 'Frage 1', 'Frage 2', … in der Reihenfolge "
        "ihres Auftretens – unabhängig davon, wie die Eingabe gegliedert ist.\n"
        "- Die 'Fragetext:'-Zeile ist PFLICHT für jede Frage und enthält die Frage wörtlich.\n"
        "- Bei mehrteiligen Fragen trägt die 'Fragetext:'-Zeile die übergeordnete Frage; jeder "
        "Unterpunkt steht zusätzlich als **fette Überschrift** (sein Fragetext) darüber.\n"
        "- KEINE Kategorie-Überschriften, KEINE Einleitung, KEINE 'Schritt 1/2/3'-Zeilen im "
        "Ergebnis. Gib NUR die Frage-Blöcke im obigen Format aus.\n"
        "- Beantworte AUSSCHLIESSLICH die unten aufgeführten Fragen. Erfinde KEINE zusätzlichen "
        "Fragen und übernimm KEINE Fragen aus der Wissensbasis. Gib GENAU so viele Frage-Blöcke "
        "aus, wie es Eingabefragen gibt, in derselben Reihenfolge.\n\n"
        f"Wissensbasis (Zusammenfassung):\n{summary_text}\n\n"
        f"Fragen:\n{questions}"
    )
    # Gemini liefert gelegentlich eine leere Antwort (response.text is None – z.B. bei
    # Recitation-/Safety-Filtern oder transienten Fehlern). Das darf NICHT als None
    # zurückgegeben werden (sonst schreibt load_or_run None und der Lauf bricht ab):
    # daher mehrfach versuchen und bei anhaltender Leere klar fehlschlagen.
    last_reason = None
    for attempt in range(1, 4):
        try:
            response = call_gemini_with_retry(
                model_name='gemini-2.5-pro',
                contents=prompt,
                config=types.GenerateContentConfig(temperature=0.1)
            )
        except Exception as e:
            print(f"Fehler bei der Qualitätssicherung: {e}")
            raise
        text = getattr(response, "text", None)
        if text:
            return text
        cands = getattr(response, "candidates", None) or []
        last_reason = (getattr(cands[0], "finish_reason", None) if cands
                       else getattr(response, "prompt_feedback", None))
        print(f"      [QA leer] Versuch {attempt}/3 ohne Text "
              f"(finish_reason={last_reason}). Neuer Versuch …")
        time.sleep(3)
    raise ValueError(
        f"Qualitätssicherung lieferte dreimal keinen Text (letzter Grund: {last_reason})."
    )


_REWORK_LEVELS = {"teilweise", "nicht enthalten"}
_NO_SUPPLEMENT_MARKER = "KEINE ERGÄNZUNG MÖGLICH"


def rework_partial_answers(qa_text: str, working_text: str, questions_path: str = None):
    """
    Schritt 5b: Nacharbeit unvollständiger Antworten.

    Für jede Frage mit Abdeckung 'teilweise' oder 'nicht enthalten' wird der
    übersetzte Originaltext (working_text) der zugehörigen Textgrundlage geprüft
    und – falls dort vorhanden – die fehlende prüfungsrelevante Information ergänzt.
    Erfolgreiche Ergänzungen erhalten die Abdeckung 'vollständig durch Nachbearbeitung'.

    Rückgabe: (augmented_qa_text, supplement_map)
      augmented_qa_text: serialisierte QA (Markdown) mit aktualisierten Items
      supplement_map: {normalize_heading(textgrundlage): [ergaenzungstext, ...]}
    """
    items = parse_qa_response(qa_text)
    # Nacharbeits-Ziele: pro Unterpunkt eine Antwort-Einheit, sonst die Frage selbst.
    # Jedes Ziel referenziert das Dict (Item oder Sub) direkt → In-Place-Update.
    targets = []
    for it in items:
        if it.get('subs'):
            for sub in it['subs']:
                targets.append({'num': it['num'], 'label': sub.get('label', ''), 'd': sub})
        else:
            targets.append({'num': it['num'], 'label': '', 'd': it})
    todo = [t for t in targets if t['d'].get('abdeckung', '').strip().lower() in _REWORK_LEVELS]
    if not todo:
        return qa_text, {}

    print(f"--- Schritt 5b: Nacharbeit {len(todo)} unvollständige(r) Antwort(en) ---")

    # Fragetexte laden (optional) – verbessert den Nacharbeitungs-Prompt
    questions_map: dict = {}
    if questions_path and os.path.exists(questions_path):
        with open(questions_path, encoding='utf-8') as qf:
            for line in qf:
                qm = re.match(r'^(\d+)\.\s+(.+)', line.strip())
                if qm:
                    questions_map[int(qm.group(1))] = qm.group(2)

    # Originaltext nach Sektionen indizieren (für gezieltes Lookup via Textgrundlage)
    sections = parse_sections(working_text)
    sec_lookup: dict = {}
    for s in sections:
        if s['heading'] == '__preamble__':
            continue
        sec_lookup[normalize_heading(s['heading'])] = s['body']

    def _section_context(textgrundlage: str) -> str:
        key = normalize_heading(textgrundlage)
        if key in sec_lookup:
            return sec_lookup[key]
        last = normalize_heading(textgrundlage.split('.')[-1])
        if last in sec_lookup:
            return sec_lookup[last]
        return working_text  # Fallback: gesamtes Kapitel

    supplement_map: dict = {}
    for t in todo:
        d = t['d']
        num, label = t['num'], t['label']
        tag = f"Frage {num}" + (f" – {label}" if label else "")
        context = _section_context(d.get('textgrundlage', ''))
        q_text = questions_map.get(num, '')
        if label:
            q_text = f"{q_text}\nUnterpunkt: {label}".strip()
        prompt = (
            "Rolle:\nDu bist Lerncoach und Prüfer für Wirtschaftspsychologie.\n\n"
            "Situation:\nEine leseleitende Frage wurde anhand einer Zusammenfassung nur "
            f"'{d.get('abdeckung')}' beantwortet. Prüfe den nachstehenden Originaltext und "
            "ergänze ausschließlich die fehlende, prüfungsrelevante Kerninformation.\n\n"
            "Regeln:\n"
            "- Maximal 3 Sätze.\n"
            "- Nur Inhalte, die der Originaltext belegt – nicht raten, nichts erfinden.\n"
            "- Keine Wiederholung der bereits vorhandenen Antwort.\n"
            f"- Wenn der Originaltext die Lücke nicht schließt, antworte exakt: {_NO_SUPPLEMENT_MARKER}\n\n"
            f"Frage:\n{q_text or '(Fragetext nicht verfügbar – siehe bisherige Antwort)'}\n\n"
            f"Bisherige Antwort:\n{d.get('antwort', '')}\n\n"
            f"Originaltext (Textgrundlage '{d.get('textgrundlage', '')}'):\n{context}"
        )
        try:
            response = call_gemini_with_retry(
                model_name='gemini-2.5-pro',
                contents=prompt,
                config=types.GenerateContentConfig(temperature=0.1)
            )
            supplement = (response.text or '').strip()
        except Exception as e:
            print(f"  {tag}: Nacharbeit fehlgeschlagen ({e}) – unverändert.")
            continue

        if not supplement or _NO_SUPPLEMENT_MARKER in supplement:
            print(f"  {tag}: keine Ergänzung im Originaltext gefunden.")
            continue

        d['antwort'] = f"{d.get('antwort', '').rstrip()} {supplement}".strip()
        d['abdeckung'] = "vollständig durch Nachbearbeitung"
        tg = d.get('textgrundlage', '')
        for key in {normalize_heading(tg), normalize_heading(tg.split('.')[-1])}:
            supplement_map.setdefault(key, []).append(supplement)
        print(f"  {tag}: ergänzt → vollständig durch Nachbearbeitung.")

    return serialize_qa_items(items), supplement_map


# ---------------------------------------------------------------------------
# Markdown-Parsing & Word-Dokument-Aufbau
# ---------------------------------------------------------------------------

def parse_sections(text: str) -> list:
    """
    Zerlegt Markdown-Text in eine flache Liste von Abschnitten.
    Jeder Eintrag: {'level': int, 'heading': str, 'body': str}
    Abschnitte vor der ersten Überschrift haben level=0, heading='__preamble__'.
    """
    sections = []
    current = {'level': 0, 'heading': '__preamble__', 'lines': []}

    for line in text.split('\n'):
        m = re.match(r'^(#{1,6})\s+(.+)$', line)
        if m:
            body = '\n'.join(current['lines']).strip()
            if body or current['heading'] != '__preamble__':
                sections.append({
                    'level': current['level'],
                    'heading': current['heading'],
                    'body': body
                })
            current = {'level': len(m.group(1)), 'heading': m.group(2).strip(), 'lines': []}
        else:
            current['lines'].append(line)

    body = '\n'.join(current['lines']).strip()
    if body or current['heading'] != '__preamble__':
        sections.append({
            'level': current['level'],
            'heading': current['heading'],
            'body': body
        })

    return sections


# Kasten-Labels für eingeschobene Lehrbuch-Boxen (Fokus/Studie/Definition/Beispiel/Exkurs).
_BOX_LABEL_RE = re.compile(r'^(fokus|studie|beispiel|definition|exkurs)\b', re.IGNORECASE)


def _is_box_heading(heading: str) -> bool:
    """True wenn die Überschrift ein eingeschobener Lehrbuch-Kasten ist (Fokus:/Studie: …)."""
    return bool(_BOX_LABEL_RE.match(_clean_heading_text(heading)))


def _split_paragraphs(body: str) -> list:
    """Body in Absätze splitten (Trenner: Leerzeile). Leere Absätze werden entfernt."""
    return [p.strip() for p in re.split(r'\n\s*\n', body) if p.strip()]


def _classify_box_boundaries(boxes: list) -> list:
    """
    LLM-Klassifikation: Für jeden Kasten die Anzahl der FÜHRENDEN Absätze, die wirklich
    zum Kasten gehören (Rest = irrtümlich von der OCR angehängter Fließtext des
    Elternabschnitts). boxes: [{'title': str, 'paragraphs': [str, ...]}, ...].
    Rückgabe: list[int] gleicher Länge, jeweils geclampt auf [1, len(paragraphs)].
    """
    lines = [
        "Du erhältst Textkästen aus einem Lehrbuch. Durch die OCR wurde an jeden Kasten "
        "fälschlich nachfolgender Fließtext angehängt, der zum umgebenden Kapitel gehört.\n",
        "Aufgabe: Bestimme für jeden Kasten, wie viele der ERSTEN Absätze tatsächlich zum "
        "Kasten gehören (thematisch zum Kastentitel passen). Die restlichen Absätze sind "
        "irrtümlich angehängter Fließtext.\n",
        "Regeln:\n- Ändere keinen Text.\n- Mindestens 1 Absatz gehört zum Kasten.\n"
        '- Gib NUR JSON zurück, Format {"0": n0, "1": n1, ...} (Kastenindex → Anzahl).\n',
        "Kästen:",
    ]
    for bi, box in enumerate(boxes):
        lines.append(f'\n=== Kasten {bi}: "{box["title"]}" ===')
        for pi, para in enumerate(box['paragraphs'], 1):
            lines.append(f'[Absatz {pi}] {para.replace(chr(10), " ")}')
    prompt = '\n'.join(lines)

    response = call_gemini_with_retry(
        model_name='gemini-2.5-pro',
        contents=prompt,
        config=types.GenerateContentConfig(temperature=0, response_mime_type='application/json'),
    )
    raw = response.text.strip()
    raw = re.sub(r'^```(?:json)?|```$', '', raw, flags=re.MULTILINE).strip()
    data = json.loads(raw)

    result = []
    for bi, box in enumerate(boxes):
        n = int(data.get(str(bi), 1))
        result.append(max(1, min(n, len(box['paragraphs']))))
    return result


def repair_box_structure(text: str) -> str:
    """
    Trennt eingeschobene Lehrbuch-Kästen (Fokus/Studie/Definition/Beispiel) von dem
    Fließtext, den die OCR fälschlich in den Kasten-Body gezogen hat, und führt diesen
    Fließtext zum Elternabschnitt zurück. Jeder Kasten wird dadurch ein eigenständiges,
    vollständiges Lernobjekt; der Elterntext bleibt zusammenhängend.

    Ohne Kästen → Originaltext (kein API-Call). Bei API-/Parse-Fehler → Originaltext
    unverändert (kein Crash, kein Datenverlust).
    """
    sections = parse_sections(text)
    box_indices = [
        i for i, s in enumerate(sections)
        if s['heading'] != '__preamble__' and _is_box_heading(s['heading'])
    ]
    if not box_indices:
        return text

    boxes = [
        {'title': _clean_heading_text(sections[i]['heading']),
         'paragraphs': _split_paragraphs(sections[i]['body'])}
        for i in box_indices
    ]
    try:
        counts = _classify_box_boundaries(boxes)
    except Exception as e:
        print(f"   [Box-Reparatur übersprungen] Klassifikation fehlgeschlagen: {e}")
        return text

    count_map = dict(zip(box_indices, counts))

    out_blocks: list[str] = []
    parent: dict | None = None  # {'head': str|None, 'paras': list[str], 'boxes': list[str]}

    def _flush(p):
        if p is None:
            return
        parts = []
        if p['head'] is not None:
            parts.append(p['head'])
        if p['paras']:
            parts.append('\n\n'.join(p['paras']))
        block = '\n\n'.join(parts).strip()
        if block:
            out_blocks.append(block)
        out_blocks.extend(p['boxes'])

    for i, s in enumerate(sections):
        head_line = None if s['heading'] == '__preamble__' else '#' * s['level'] + ' ' + s['heading']

        if i in count_map:
            paras = _split_paragraphs(s['body'])
            n = count_map[i]
            box_content, running = paras[:n], paras[n:]
            box_parts = [head_line] if head_line else []
            if box_content:
                box_parts.append('\n\n'.join(box_content))
            box_block = '\n\n'.join(box_parts).strip()
            if parent is not None:
                parent['paras'].extend(running)
                if box_block:
                    parent['boxes'].append(box_block)
            else:
                # Kasten ohne vorangehenden Elternabschnitt: eigenständig belassen.
                if box_block:
                    out_blocks.append(box_block)
                if running:
                    out_blocks.append('\n\n'.join(running))
        else:
            _flush(parent)
            parent = {'head': head_line, 'paras': _split_paragraphs(s['body']), 'boxes': []}

    _flush(parent)
    return '\n\n'.join(out_blocks).strip() + '\n'


def normalize_heading(h: str) -> str:
    """Für Matching: Bold-/Italic-Marker entfernen, lowercase. Nummern bleiben für eindeutige Keys."""
    h = re.sub(r'\*+', '', h)   # strip ** und *
    return h.lower().strip()


def _textgrundlage_keys(textgrundlage: str) -> set:
    """Kandidat-Schlüssel für die Zuordnung einer QA-Textgrundlage zu einer Dokument-Sektion.

    Eine Textgrundlage kann zusätzliche Details tragen (z.B.
    '4.3.1 Simulationsbasiertes Training, • Wirkungsweise der Methode und Befunde').
    Neben dem Volltext werden daher auch das erste Komma-/Bullet-Segment (die eigentliche
    Überschrift) – mit und ohne führende Kapitelnummer – sowie der letzte Punkt-Teil als
    Schlüssel erzeugt, damit das Matching zur Überschrift gelingt.
    """
    # Markdown-Links/HTML entfernen und übrig gebliebene eckige Klammern (OCR-Artefakt
    # '[Blended Learning]' ohne URL) tilgen, damit der Schlüssel zur bereinigten Überschrift passt.
    tg = _clean_heading_text(textgrundlage or '').replace('[', '').replace(']', '')
    keys = {normalize_heading(tg), normalize_heading(tg.split('.')[-1])}
    first = re.split(r'[,;•]', tg)[0].strip()
    if first:
        keys.add(normalize_heading(first))
        keys.add(normalize_heading(re.sub(r'^[\d.]+\s*', '', first)))
    return {k for k in keys if k}


def _remap_textgrundlage(textgrundlage: str, renumber: dict) -> str:
    """Ersetzt die (Original-)Kapitelnummer einer QA-Textgrundlage durch die finale
    Dokumentnummer aus `renumber` (Original-Schlüssel → neue Nummer), damit die in den
    Lernfragen angezeigte 'Quelle' im Dokument navigierbar ist. Ohne Treffer unverändert.
    """
    if not renumber or not textgrundlage:
        return textgrundlage
    cand = _textgrundlage_keys(textgrundlage)
    # Nummern-tragende Schlüssel zuerst (eindeutiger), dann Titel-Schlüssel.
    ordered = sorted(cand, key=lambda k: (0 if re.search(r'\d', k) else 1, k))
    for k in ordered:
        if k in renumber:
            new_num = renumber[k]
            if re.match(r'^\s*[\d.]+', textgrundlage):
                return re.sub(r'^\s*[\d.]+', new_num, textgrundlage, count=1)
            return f"{new_num} {textgrundlage}".strip()
    return textgrundlage


def normalize_heading_levels(text: str) -> str:
    """
    Normalisiert inkonsistente Markdown-Überschriftenebenen.
    Nummerierte Kapitelüberschriften erhalten konsistente Ebenen:
      "1 Titel"     → # (H1)
      "4.1 Titel"   → ## (H2)
      "4.1.1 Titel" → ### (H3)
    Nicht-nummerierte Überschriften werden nicht verändert.
    Tail-nummerierte Überschriften ("Titel 7.2.2") werden umgeordnet und
    ebenfalls normalisiert – mindestens 2 Punkte erforderlich (schließt
    einfache Abbildungsnummern wie "18.4" aus).
    """
    result = []
    for line in text.split('\n'):
        m = re.match(r'^(#{1,6})\s+(.+)$', line)
        if m:
            content = m.group(2).strip()
            clean = content.replace('**', '').strip()
            clean_nohtml = re.sub(r'<[^>]+>', '', clean)
            num_m = re.match(r'^(\d+(?:\.\d+)*)(\s|$)', clean_nohtml)
            if num_m:
                dots = num_m.group(1).count('.')
                new_level = '#' * min(dots + 1, 6)
                result.append(f'{new_level} {content}')
            else:
                # OCR-Artefakt: Kapitelnummer am Ende ("Titel 7.2.2" → "7.2.2 Titel").
                # Mindestens 2 Punkte (N.M.P) erforderlich – schließt Jahre und einfache
                # Abbildungsnummern aus ("Studie 2023", "Abbildung 18.4" bleiben unverändert).
                tail_m = re.match(r'^(.+?)\s+(\d+(?:\.\d+){2,})\s*$', clean_nohtml)
                if tail_m:
                    dots = tail_m.group(2).count('.')
                    new_level = '#' * min(dots + 1, 6)
                    reordered = f"{tail_m.group(2)} {tail_m.group(1).rstrip()}"
                    result.append(f'{new_level} {reordered}')
                else:
                    result.append(line)
        else:
            result.append(line)
    return '\n'.join(result)


def add_formatted_text(paragraph, text, default_color=None):
    """Parse Markdown: ***bold+italic***, **bold**, *italic*."""
    parts = re.split(r'(\*{3}[^*]+?\*{3}|\*{2}[^*]+?\*{2}|\*[^*]+?\*)', text)
    for part in parts:
        if part.startswith('***') and part.endswith('***') and len(part) > 6:
            run = paragraph.add_run(part[3:-3])
            run.bold = True
            run.italic = True
        elif part.startswith('**') and part.endswith('**') and len(part) > 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            run = paragraph.add_run(part)

        if default_color:
            run.font.color.rgb = default_color


def _clean_unklar_cell(text: str) -> str:
    """
    Verbessert unleserliche OCR-Zellen: [UNKLAR: N<br>fo<br>k<br>t...]
    Fragmente werden zusammengefügt. Kurze Fragmente (≤2 Zeichen) werden
    direkt verbunden (Zeichen-OCR), längere mit Leerzeichen getrennt.
    Ergebnis: kompakter [OCR: ...]-Hinweis statt seitenlanger Garbage-Spalte.
    """
    def rebuild(m):
        inner = m.group(1).strip()
        parts = [p.strip() for p in re.split(r'<br\s*/?>', inner) if p.strip()]
        if not parts:
            return '[unleserlich]'
        avg_len = sum(len(p) for p in parts) / len(parts)
        joined = ''.join(parts) if avg_len <= 2.5 else ' '.join(parts)
        truncated = joined[:80] + ('…' if len(joined) > 80 else '')
        return f'[OCR: {truncated}]'
    return re.sub(r'\[UNKLAR:\s*(.*?)\]', rebuild, text, flags=re.DOTALL)


def _html_entities(text: str) -> str:
    """Wandelt einfache HTML-Entities in Plaintext um."""
    return (text
            .replace('&amp;', '&')
            .replace('&lt;', '<')
            .replace('&gt;', '>')
            .replace('&nbsp;', ' ')
            .replace('&#x27;', "'")
            .replace('&quot;', '"'))


def _set_cell_text(cell, raw_text: str, bold: bool = False, font_size: Pt = None):
    """Schreibt Text in eine Word-Tabellenzelle.
    Kopfzeilen (bold=True): <br> → Leerzeichen (kein Umbruch).
    Datenzellen: <br> → eigener Absatz je Teil.
    Bullet-Concatenation (•A•B•C) → separate Absätze je Eintrag.
    """
    font_size = font_size or Pt(9.5)
    raw_text = _clean_unklar_cell(raw_text)
    raw_text = _html_entities(raw_text)

    if not bold and raw_text.count('•') > 1:
        bullet_parts = [p.strip() for p in raw_text.split('•') if p.strip()]
        raw_text = '\n'.join(bullet_parts)

    if bold:
        # Kopfzeile: alle Umbrüche als Leerzeichen, einzeiliger Text
        text = raw_text.replace('<br>', ' ').replace('\n', ' ').strip()
        para = cell.paragraphs[0]
        para.clear()
        run = para.add_run(text)
        run.bold = True
        run.font.size = font_size
    else:
        raw_text = raw_text.replace('<br>', '\n')
        parts = raw_text.split('\n')
        para = cell.paragraphs[0]
        para.clear()
        run = para.add_run(parts[0].strip())
        run.font.size = font_size
        for part in parts[1:]:
            para = cell.add_paragraph()
            run = para.add_run(part.strip())
            run.font.size = font_size


def _calc_col_widths(rows_data: list, num_cols: int) -> list:
    """
    Berechnet Spaltenbreiten in cm proportional zum längsten Zellinhalt je Spalte.
    Gesamtbreite: 16 cm (A4 minus Standardränder). Mindestbreite: 1,2 cm.
    """
    PAGE_CM = 16.0
    MIN_CM  = 1.2
    max_lens = []
    for c in range(num_cols):
        ml = max(
            (len(row[c].replace('<br>', ' ').replace('\n', ' ')) for row in rows_data if c < len(row)),
            default=1
        )
        max_lens.append(max(ml, 1))
    total = sum(max_lens)
    raw = [max(PAGE_CM * ml / total, MIN_CM) for ml in max_lens]
    scale = PAGE_CM / sum(raw)
    return [w * scale for w in raw]


def _shade_cell(cell, fill_hex: str):
    """Setzt Hintergrundfarbe einer Tabellenzelle (z.B. 'DDEEFF')."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


_STATS_RE = re.compile(r'^(\d+\.\d+)\s+(\d+\.\d+)\s+(\.?\d+\.?\d*)$')

def _fix_concatenated_stats_cells(rows_data: list) -> list:
    """
    Korrigiert OCR-Artefakt: M, SD und ritc werden manchmal in einer Zelle
    zusammengefasst (z.B. '3.44 1.00 .49'). Wenn links und rechts der
    betreffenden Zelle leere Nachbarzellen existieren, werden die drei Werte
    auf die drei Zellen verteilt.
    """
    for row in rows_data:
        for j in range(len(row)):
            m = _STATS_RE.match(row[j].strip())
            if not m:
                continue
            # Prüfen ob Nachbarzellen leer sind (links+rechts oder nur rechts+rechts)
            if j >= 1 and j + 1 < len(row) and row[j - 1] == '' and row[j + 1] == '':
                row[j - 1], row[j], row[j + 1] = m.group(1), m.group(2), m.group(3)
            elif j + 2 < len(row) and row[j + 1] == '' and row[j + 2] == '':
                row[j], row[j + 1], row[j + 2] = m.group(1), m.group(2), m.group(3)
    return rows_data


def _drop_empty_columns(rows_data: list) -> list:
    """Entfernt Spalten, die in allen Zeilen leer sind (OCR-Artefakt)."""
    if not rows_data:
        return rows_data
    num_cols = max(len(r) for r in rows_data)
    keep = [
        c for c in range(num_cols)
        if any(c < len(r) and r[c].strip() for r in rows_data)
    ]
    return [[r[c] if c < len(r) else '' for c in keep] for r in rows_data]


def _clean_table_cell(cell: str) -> str:
    """Bereinigt eine einzelne Tabellenzelle von OCR-Artefakten:
    - normalisiert geschützte Leerzeichen und kollabiert lange Whitespace-Läufe
      (OCR hängt teils Tausende Leerzeichen an → sonst riesige Spaltenbreiten),
    - leert Zellen, die nur aus Trennstrichen/Unterstrichen bestehen (als Tabellenlinie
      fehlerkannte horizontale Rule → erscheint sonst als '------'-Spalte)."""
    cell = cell.replace('\xa0', ' ')
    cell = re.sub(r'\s{2,}', ' ', cell).strip()
    if cell and not re.search(r'[^\s\-–—_]', cell):
        return ''
    return cell


def add_markdown_table_to_doc(doc, table_lines: list):
    """
    Konvertiert eine Liste von Markdown-Pipe-Zeilen in eine Word-Tabelle.
    Leere Zellen in derselben Spalte werden vertikal mit der Zelle darüber zusammengeführt.
    Tabellen werden nie ausgeblendet.
    """
    SEP_RE = re.compile(r'^\|[-|: ]+\|$')

    rows_data = []
    for raw_line in table_lines:
        line = raw_line.strip()
        if SEP_RE.match(line):
            continue
        cells = [_clean_table_cell(c) for c in line.split('|')]
        if cells and cells[0] == '':
            cells = cells[1:]
        if cells and cells[-1] == '':
            cells = cells[:-1]
        if not any(c for c in cells):
            continue  # komplett leere Zeile (z.B. reine Strich-/Linienzeile) überspringen
        rows_data.append(cells)

    if not rows_data:
        return

    num_cols = max(len(r) for r in rows_data)

    # Zellinhalt normieren (alle Zeilen auf gleiche Spaltenanzahl)
    for r in rows_data:
        while len(r) < num_cols:
            r.append('')

    rows_data = _fix_concatenated_stats_cells(rows_data)
    rows_data = _drop_empty_columns(rows_data)
    num_cols = max(len(r) for r in rows_data)
    num_rows = len(rows_data)

    # Titel-Zeile erkennen: Zeile 0 hat ≤1 Non-Empty-Cell bei mehreren Spalten – oder sie
    # beginnt mit einer Tabellen-Caption ("TABELLE 1 …" / "TABLE 1 …"), die die OCR als erste
    # Tabellenzeile statt als Überschrift erfasst hat. Sie wird dann zu einer zusammengeführten
    # Titelzeile, sodass der eigentliche Spaltenkopf in Zeile 1 erhalten bleibt.
    non_empty_row0 = sum(1 for c in rows_data[0] if c.strip())
    row0_is_caption = bool(re.match(r'(?i)^\**\s*tab(?:elle|le)\s+\d', rows_data[0][0].strip()))
    has_title_row = (num_rows > 1 and num_cols > 1 and (non_empty_row0 <= 1 or row0_is_caption))
    header_row_idx = 1 if has_title_row else 0

    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid'

    # Spaltenbreiten proportional zum Inhalt setzen
    col_widths_cm = _calc_col_widths(rows_data, num_cols)
    font_size = Pt(8) if num_cols > 5 else Pt(9.5)
    for ci, w_cm in enumerate(col_widths_cm):
        for cell in table.column_cells(ci):
            cell.width = Cm(w_cm)

    # Zellen befüllen
    for r_idx, row in enumerate(rows_data):
        for c_idx, cell_text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            is_header = (r_idx == header_row_idx)
            is_title  = (has_title_row and r_idx == 0)
            _set_cell_text(cell, cell_text, bold=(is_header or is_title), font_size=font_size)
            if is_header:
                _shade_cell(cell, 'D9E2F3')

    # Titelzeile: alle Zellen in Zeile 0 zusammenführen
    if has_title_row:
        anchor = table.cell(0, 0)
        for c in range(1, num_cols):
            anchor.merge(table.cell(0, c))

    # Leere Zellen vertikal zusammenführen (Spanning-Simulation)
    for c in range(num_cols):
        r = 1
        while r < num_rows:
            if rows_data[r][c] == '':
                try:
                    table.cell(r - 1, c).merge(table.cell(r, c))
                except Exception:
                    pass
            r += 1

    doc.add_paragraph()  # Abstand nach Tabelle


def _clean_for_hidden(text: str) -> str:
    """Bereinigt Markdown-Text für den Hidden-Originaltext-Block.
    Entfernt Links, HTML-Tags, Seitenreferenzen und andere OCR-Artefakte.
    Bilder (![]()) bleiben erhalten."""
    text = re.sub(r'(?<!!)\[([^\]]*)\]\([^)]*\)', r'\1', text)  # [Text](url) → Text (nicht Bilder)
    text = re.sub(r'<[^>]+>', '', text)                          # HTML-Tags entfernen
    text = re.sub(r'^\[\d+\]\s*$', '', text, flags=re.MULTILINE)  # [1] allein → weg
    return text


def _strip_kontrollliste(text: str) -> str:
    """Entfernt Kontrolllisten-Blöcke aus der übersetzten Markdown-Datei.
    Muster: (optionales ---) gefolgt von **Kontrollliste** und Bullet-Zeilen."""
    text = re.sub(
        r'(?m)^---\s*\n\*\*Kontrollliste\*\*.*?(?=\n#{1,6}\s|\Z)',
        '',
        text,
        flags=re.DOTALL
    )
    # Kontrollliste ohne vorangehendes ---
    text = re.sub(
        r'(?m)^\*\*Kontrollliste\*\*.*?(?=\n#{1,6}\s|\Z)',
        '',
        text,
        flags=re.DOTALL
    )
    return text


def _strip_selbstpruefung(text: str) -> str:
    """Entfernt KI-interne Selbstprüfungs-Blöcke aus der Zusammenfassung.
    Diese Abschnitte sind KI-Validierungsartefakte und kein Lerninhalt.
    Unterstützt Varianten: '### **Selbstprüfung**', '**Selbstprüfung**', 'Selbstprüfung'."""
    # Vorangehendes --- wegstreifen (erscheint typischerweise direkt vor dem Block)
    text = re.sub(
        r'(?m)^---\s*\n(?=(?:#{1,6}\s+)?\*{0,3}\*?\*?[Ss]elbstpr)',
        '',
        text
    )
    text = re.sub(
        r'(?m)^(?:#{1,6}\s+)?\*{0,3}\*?\*?[Ss]elbstpr[uü]fung\*?\*?\*?.*?(?=\n#{1,6}\s|\Z)',
        '',
        text,
        flags=re.DOTALL
    )
    return text


def _compress_heading_levels(text: str) -> str:
    """
    Korrigiert inkonsistente OCR-Heading-Ebenen.
    Verhindert Sprünge > 1 Ebene; Geschwister-Überschriften (gleiche OCR-Ebene)
    erhalten denselben tatsächlichen Level – kein Kaskadeneffekt.
    """
    ocr_to_actual: dict = {}
    prev_actual = 0
    result = []
    for line in text.split('\n'):
        m = re.match(r'^(#{1,9})\s', line)
        if m:
            ocr_level = len(m.group(1))
            # Nummerierte Headings wurden bereits von normalize_heading_levels korrekt gesetzt
            if re.match(r'^#{1,9}\s+\*{0,2}\d', re.sub(r'<[^>]+>', '', line)):
                result.append(line)
                prev_actual = ocr_level
                continue
            if ocr_level in ocr_to_actual:
                actual = ocr_to_actual[ocr_level]
            elif ocr_level > prev_actual + 1:
                actual = prev_actual + 1
                ocr_to_actual[ocr_level] = actual
            else:
                actual = ocr_level
                ocr_to_actual[ocr_level] = actual
            prev_actual = actual
            result.append('#' * actual + line[ocr_level:])
        else:
            result.append(line)
    return '\n'.join(result)


def _strip_ocr_y_prefix(text: str) -> str:
    """Entfernt isoliertes 'y '-Präfix aus Headings und Bullet-Items.
    OCR-Artefakt: Sonderzeichen (Bullet-Pfeil) einer Sonderzeichenschrift wird als 'y' gelesen.
    Bereinigt außerdem trailing \\_ Sequenzen in Heading-Zeilen (OCR-Artefakt für Unterstriche)."""
    lines = []
    for line in text.split('\n'):
        if re.match(r'^#{1,6}\s', line):
            # Trailing \_ Sequenzen aus Heading-Text entfernen (z.B. "# Titel: \_ \_ \_")
            line = re.sub(r'(\s*\\_)+\s*$', '', line).rstrip()
        # Heading: ## y Titel  ODER  ## **y Titel** (y innerhalb Bold-Marker, wie KI es erzeugt)
        line = re.sub(r'^(#{1,6}\s+)(\**)y\s+', r'\1\2', line)
        # Bullet: beliebige Einrückung + Bullet-Marker + y (inkl. eingerückte Sublisten)
        line = re.sub(r'^(\s*[-*]\s+)y\s+', r'\1', line)
        lines.append(line)
    return '\n'.join(lines)


def _image_display_width(img_path: str, max_in: float = 5.5, assumed_dpi: int = 150):
    """Berechnet Anzeigebreite: min(natürliche Bildbreite, max_in).
    Nutzt PIL/Pillow falls verfügbar; andernfalls Dateigröße als Heuristik."""
    try:
        from PIL import Image as _PILImage
        with _PILImage.open(img_path) as im:
            natural_in = im.width / assumed_dpi
            return Inches(min(natural_in, max_in))
    except Exception:
        pass
    try:
        kb = os.path.getsize(img_path) / 1024
        if kb < 10:
            return Inches(1.5)
        if kb < 50:
            return Inches(3.0)
    except Exception:
        pass
    return Inches(max_in)


def _is_decorative_image(img_path: str, min_px: int = 100) -> bool:
    """True wenn das Bild wahrscheinlich dekorativ ist (Icon, Randsymbol).
    Kriterium: kleinste Seite < min_px (alle Kronen-Icons haben min ≤ 100 px;
    Inhaltsbilder wie 'Was ist das?'-Grafiken haben min > 100 px).
    Fallback ohne PIL: Dateigröße < 10 KB."""
    try:
        from PIL import Image as _PILImage
        with _PILImage.open(img_path) as im:
            w, h = im.width, im.height
            return min(w, h) < min_px
    except Exception:
        pass
    try:
        return os.path.getsize(img_path) < 10 * 1024
    except Exception:
        return False


_CAPTION_RE = re.compile(
    r'^\**\s*(?:[-•*]\s+)?\**\s*'
    r'(?:Tabelle|Abbildung|Tab\.|Abb\.|Table|Figure|Fig\.)\s*(?:\d|:)',
    re.IGNORECASE
)
_HRULE_RE = re.compile(r'^-{3,}$|^\*{3,}$|^_{3,}$')


_CAPTION_TITLE_RE = re.compile(
    r'^\**\s*(?:[-•*]\s+)?\**\s*'
    r'(?:Tabelle|Abbildung|Tab\.|Abb\.|Table|Figure|Fig\.)\s*\d*\s*[:.\-–)]?\s*(.*)',
    re.IGNORECASE
)

# Caption-Schlüsselwort ohne Anker (für Präfix-Normalisierung).
_CAPTION_KEYWORD = r'(?:Tabelle|Abbildung|Tab\.|Abb\.|Table|Figure|Fig\.)'

# Ein einzelnes OCR-Streuzeichen (verlorener Icon-/Bullet-/Symbolfont-Glyph): bloße Ziffer
# OHNE folgenden Punkt/Ziffer (echte Sektionsnummern wie '3.1' bleiben unberührt), einzelner
# Buchstabe (j, y …) oder ein Symbol. Bewusst eng gehalten, um echten Inhalt nicht zu treffen.
_OCR_GLYPH = r'(?:\d(?![\d.])|[A-Za-zÄÖÜäöüß](?![A-Za-zÄÖÜäöüß])|[.■▪●•·∎◾◼])'

_CAPTION_GLYPH_PREFIX_RE = re.compile(
    rf'^\s*(?:#{{1,6}}\s+)?{_OCR_GLYPH}\s+(\**\s*{_CAPTION_KEYWORD}.*)$',
    re.IGNORECASE,
)
_HEADING_GLYPH_ARTIFACT_RE = re.compile(rf'^#{{1,6}}\s+{_OCR_GLYPH}\s+(\*\*.+)$')
_BULLET_GLYPH_ARTIFACT_RE = re.compile(rf'^(\s*[-*]\s+){_OCR_GLYPH}\s+(\*\*.+)$')


def normalize_ocr_glyph_artifacts(text: str) -> str:
    """Bereinigt verlorene OCR-Icon-/Bullet-Glyphen am Zeilenanfang (Marker liest Symbolfont-
    Zeichen als Streuzeichen wie '.', '■', 'j' oder bloße Ziffern '4'/'5').

    Drei zeilenweise Transformationen (Reihenfolge wichtig):
      1. Caption-Präfix: '. **Abb. 3.1** …' / '■ Abb. 3.2 …' / '#### . **Tab. 3.3** …'
         → reine Caption ('**Abb. 3.1** …'), damit _CAPTION_RE greift und sie sichtbar bleibt.
      2. Heading-Artefakt: '#### 4 **Modelllernen:**' → Bullet '- **Modelllernen:**'
         (bloße Glyph-Headings sind Icon-Listenpunkte, keine Unterkapitel → keine Nummer).
      3. Bullet-Artefakt: '- 4 **Positive Verstärkung:**' → '- **Positive Verstärkung:**'.

    Echte Sektionsnummern ('#### **3.1 …**', '#### 3.2.3 …') und echte nummerierte Listen
    ('- 1. **…**') bleiben unberührt (Glyph-Regex schließt 'Ziffer + Punkt' aus)."""
    out = []
    for line in text.split('\n'):
        m_cap = _CAPTION_GLYPH_PREFIX_RE.match(line)
        if m_cap:
            out.append(m_cap.group(1).strip())
            continue
        m_h = _HEADING_GLYPH_ARTIFACT_RE.match(line)
        if m_h:
            out.append(f"- {m_h.group(1)}")
            continue
        m_b = _BULLET_GLYPH_ARTIFACT_RE.match(line)
        if m_b:
            out.append(f"{m_b.group(1)}{m_b.group(2)}")
            continue
        out.append(line)
    return '\n'.join(out)


_TOC_SEP_RE = re.compile(r'^\s*\|[\s:|-]+\|?\s*$')
# Datenzeile eines Kapitel-Inhaltsverzeichnisses: erste Zelle = Sektionsnummer (3.1, 3.2.1),
# irgendwo eine Seitenangabe als '– 40' / '<br>– 40' (Bindestrich/Gedankenstrich + Zahl).
_TOC_PAGE_RE = re.compile(r'[-–—]\s*\d{1,4}\b')


def strip_toc_tables(text: str) -> str:
    """Entfernt Kapitel-Mini-Inhaltsverzeichnisse, die als Markdown-Tabelle vorliegen
    ('| 3.1 | Behavioristische Ansätze<br>– 40 |' …). Solche TOC-Tabellen stehen oft am
    Kapitelanfang und sind kein Inhalt – sie stören Einleitung, Zusammenfassung und Layout.

    Erkennung je zusammenhängendem Tabellenblock: Mehrheit der Datenzeilen hat als erste Zelle
    eine reine Sektionsnummer UND eine Seitenangabe ('– NN'); Schwelle ≥3 solcher Zeilen.
    Echte Datentabellen (ohne dieses Muster) bleiben unberührt."""
    lines = text.split('\n')
    out: list[str] = []
    i = 0
    n = len(lines)
    while i < n:
        if not lines[i].lstrip().startswith('|'):
            out.append(lines[i])
            i += 1
            continue
        # Zusammenhängenden Tabellenblock einsammeln.
        j = i
        block = []
        while j < n and lines[j].lstrip().startswith('|'):
            block.append(lines[j])
            j += 1
        data_rows = [b for b in block if not _TOC_SEP_RE.match(b)]
        toc_rows = 0
        for b in data_rows:
            cells = [c.strip() for c in b.strip().strip('|').split('|')]
            first = next((c for c in cells if c), '')
            if re.match(r'^\d+(\.\d+)*$', first) and _TOC_PAGE_RE.search(b):
                toc_rows += 1
        is_toc = toc_rows >= 3 and toc_rows >= len(data_rows) / 2
        if not is_toc:
            out.extend(block)
        i = j
    return '\n'.join(out)


def _textfigure_visible_lines(lines: list) -> set:
    """Zeilen-Indizes, die zu einer *Text-Figur* gehören und im Embed-Modus (hide_text=True)
    sichtbar bleiben sollen – analog zu Bild-/Tabellen-Figuren.

    Hintergrund: Marker klassifiziert manche Abbildungen (z.B. Leitfragen-Kästen) als Text
    statt als Bild. Solche Figuren haben eine Caption ('Abb. N: TITEL'), aber kein '![](…)'.
    Ohne Sonderbehandlung verschwindet ihr Inhalt im ausgeblendeten Originaltext.

    Vorgehen je Caption: Der sichtbare Block reicht von der zugehörigen Intro-Überschrift bis
    zur Caption. Die Intro-Überschrift wird über Titel-Abgleich gefunden (Heading-Text ≈
    Caption-Titel) – das ist das robuste Signal für genau den Figur-Kasten. Die Rückwärtssuche
    stoppt an einem Bild oder einer nummerierten Kapitelüberschrift. Ohne passende Intro-
    Überschrift bleibt nur die Caption-Zeile selbst sichtbar (kein Aufblähen der Kapitel-Prosa).
    """
    n = len(lines)

    def is_heading(k):
        return bool(re.match(r'^#{1,6}\s', lines[k]))

    def is_image(k):
        return bool(re.match(r'^\s*!\[[^\]]*\]\([^)]+\)\s*$', lines[k]))

    visible = set()
    for i, line in enumerate(lines):
        if not _CAPTION_RE.match(line.strip()):
            continue
        m = _CAPTION_TITLE_RE.match(re.sub(r'<[^>]+>', '', line.strip()))
        title = (m.group(1) if m else '').strip().strip('*').strip().lower()
        start = i  # Fallback: nur die Caption-Zeile
        j = i - 1
        while j >= 0:
            if is_image(j):
                break
            if is_heading(j):
                hc = _clean_heading_text(lines[j]).lower()
                if re.match(r'^\d', hc):
                    break  # nummerierte Kapitelüberschrift = Grenze
                if len(title) >= 8 and (hc in title or title in hc):
                    start = j  # passende Intro-Überschrift gefunden
                    break
            j -= 1
        visible.update(range(start, i + 1))
    return visible


def audit_figure_captions(text: str, base_path: str = None) -> dict:
    """Prüft, ob zu jeder Abbildungs-/Tabellen-Caption eine sichtbare Abbildung gehört, und
    gibt eine Warnung aus, falls Captions nur als Text (ohne Grafik) vorliegen oder ein
    referenziertes Bild fehlt. Verhindert, dass unbemerkt Abbildungen fehlen.

    Zuordnung je Caption über ihre Heading-Sektion: enthält die Sektion ein vorhandenes,
    nicht-dekoratives Bild → 'image'; eine Tabelle → 'table'; ein referenziertes Bild, dessen
    Datei fehlt → 'missing'; sonst 'text' (Text-Figur – Inhalt sichtbar, aber keine Grafik).
    """
    lines = text.split('\n')
    head_idx = [i for i, l in enumerate(lines) if re.match(r'^#{1,6}\s', l)]
    bounds = head_idx + [len(lines)]

    def section_of(line_no):
        # Body der Heading-Sektion, die line_no enthält (oder Gesamttext, falls kein Heading)
        prev = [h for h in head_idx if h <= line_no]
        if not prev:
            return lines
        s = prev[-1]
        e = bounds[head_idx.index(s) + 1]
        return lines[s:e]

    counts = {'image': 0, 'table': 0, 'text': 0, 'missing': 0}
    text_only, missing = [], []
    for i, line in enumerate(lines):
        if not _CAPTION_RE.match(line.strip()):
            continue
        body = section_of(i)
        status = 'text'
        for b in body:
            im = re.match(r'^\s*!\[[^\]]*\]\(([^)]+)\)\s*$', b)
            if im:
                rel = im.group(1)
                path = os.path.join(base_path, rel) if base_path else rel
                if os.path.exists(path):
                    # Vorhandenes Bild in einer Caption-Sektion wird eingebettet – auch kleine
                    # Panels (Fix C schützt sie vor dem Dekorativ-Filter).
                    status = 'image'
                    break
                status = 'missing'
            elif b.strip().startswith('|') and status != 'image':
                status = 'table'
        counts[status] += 1
        cap = re.sub(r'<[^>]+>', '', line.strip())[:70]
        if status == 'text':
            text_only.append(cap)
        elif status == 'missing':
            missing.append(cap)

    total = sum(counts.values())
    if total:
        msg = (f"[ABBILDUNG] {total} Caption(s): {counts['image']} mit Bild, "
               f"{counts['table']} mit Tabelle, {counts['text']} nur Text")
        if counts['missing']:
            msg += f", {counts['missing']} Bild FEHLT"
        print(msg)
        for cap in missing:
            print(f"    [!] Bilddatei fehlt: {cap}")
        for cap in text_only:
            print(f"    [i] nur als Text vorhanden (keine Grafik): {cap}")
    return {'counts': counts, 'text_only': text_only, 'missing': missing}


_IMG_LINE_RE = re.compile(r'^\s*!\[[^\]]*\]\(([^)]+)\)\s*$')


def render_pdf_figure_region(pdf_path: str, page_index: int, out_path: str,
                             scale: float = 2.5) -> bool:
    """Rendert die Abbildungs-Region einer PDF-Seite als ein Bild (für Figuren, die Marker in
    Text + Icons zerlegt hat). Region = Bounding-Box aller Pfad-Objekte (Diagrammlinien/Boxen),
    erweitert um Textobjekte im selben vertikalen Band (Box-Beschriftungen/Erklärtext); die
    Caption darunter und der Seitenkopf darüber bleiben außen vor.

    Gibt True zurück, wenn ein Bild geschrieben wurde, sonst False (kein pypdfium2, keine Seite,
    keine Pfad-Objekte). Bestehender Cache wird wiederverwendet."""
    if os.path.exists(out_path):
        return True
    try:
        import pypdfium2 as pdfium
    except Exception:
        return False
    try:
        pdf = pdfium.PdfDocument(pdf_path)
    except Exception:
        return False
    try:
        if page_index < 0 or page_index >= len(pdf):
            return False
        page = pdf[page_index]
        W, H = page.get_size()
        paths, texts = [], []
        for obj in page.get_objects():
            try:
                pos = obj.get_pos()  # (left, bottom, right, top), PDF-Koordinaten (unten=0)
            except Exception:
                continue
            if obj.type == 2:        # Pfad (Diagrammlinien/Boxen/Pfeile)
                paths.append(pos)
            elif obj.type in (1, 3): # Text / eingebettetes Bild
                texts.append(pos)
        if not paths:
            return False
        pl = min(b[0] for b in paths); pb = min(b[1] for b in paths)
        pr = max(b[2] for b in paths); pt = max(b[3] for b in paths)
        # Textobjekte im vertikalen Band der Pfade (Box-Texte) einbeziehen – Caption (unterhalb
        # der Pfade) und Seitenkopf (oberhalb) bleiben außen vor.
        for b in texts:
            cy = (b[1] + b[3]) / 2
            if pb - 6 <= cy <= pt + 6:
                pl = min(pl, b[0]); pr = max(pr, b[2])
                pb = min(pb, b[1]); pt = max(pt, b[3])
        pad = 8.0
        L = max(0.0, pl - pad); R = min(W, pr + pad)
        B = max(0.0, pb - pad); T = min(H, pt + pad)
        bitmap = page.render(scale=scale).to_pil()
        # PDF-Koordinaten (unten=0) → Pixel (oben=0)
        crop = bitmap.crop((int(L * scale), int((H - T) * scale),
                            int(R * scale), int((H - B) * scale)))
        if crop.width < 20 or crop.height < 20:
            return False
        crop.convert('RGB').save(out_path, 'JPEG', quality=85)
        return True
    except Exception as e:
        print(f"   [ABBILDUNG] Rendern von Seite {page_index} fehlgeschlagen: {e}")
        return False
    finally:
        try:
            pdf.close()
        except Exception:
            pass


def inline_decomposed_figures(text: str, pdf_path: str, base_path: str) -> str:
    """Ersetzt zerstückelte Abbildungen (Marker hat eine Grafik in Text + winzige Icons zerlegt)
    durch ein einziges, aus der Quell-PDF gerendertes Bild.

    Vorgehen je Caption-Sektion (Heading-Segment mit 'Abb./Tab. N'-Caption): enthält die Sektion
    KEINE Tabelle und NUR dekorative/winzige Bilder, gilt sie als zerstückelt. Die Seite wird aus
    den '_page_N_'-Icon-Namen bestimmt, die Figur-Region gerendert und die erste Icon-Bildzeile
    durch das gerenderte Bild ersetzt; weitere Icon-Bildzeilen der Sektion werden entfernt.

    No-Op, wenn pdf_path/base_path fehlen oder das Rendern scheitert (Icons bleiben dann)."""
    if not pdf_path or not base_path or not os.path.exists(pdf_path):
        return text
    lines = text.split('\n')
    head_idx = [i for i, l in enumerate(lines) if re.match(r'^#{1,6}\s', l)]
    bounds = head_idx + [len(lines)]

    def seg_range(i):
        prev = [h for h in head_idx if h <= i]
        if not prev:
            return 0, (head_idx[0] if head_idx else len(lines))
        s = prev[-1]
        return s, bounds[head_idx.index(s) + 1]

    drop: set[int] = set()           # zu entfernende Icon-Zeilen
    replace: dict[int, str] = {}     # Zeilenindex → neue Bildzeile
    handled_segments: set = set()
    rendered = 0
    for i, line in enumerate(lines):
        if not _CAPTION_RE.match(line.strip()):
            continue
        s, e = seg_range(i)
        if s in handled_segments:
            continue
        handled_segments.add(s)
        img_lines, has_table, has_real_img, pages = [], False, False, []
        for k in range(s, e):
            m = _IMG_LINE_RE.match(lines[k])
            if m:
                rel = m.group(1)
                path = os.path.join(base_path, rel)
                if os.path.exists(path) and not _is_decorative_image(path):
                    has_real_img = True
                else:
                    img_lines.append(k)
                pm = re.search(r'_page_(\d+)_', rel)
                if pm:
                    pages.append(int(pm.group(1)))
            elif lines[k].strip().startswith('|'):
                has_table = True
        # Nur zerstückelte Figuren behandeln: kein echtes Bild, keine Tabelle, aber Icon-Bilder.
        if has_real_img or has_table or not img_lines or not pages:
            continue
        page_index = max(set(pages), key=pages.count)
        out_name = f"_rendered_fig_p{page_index}.jpeg"
        out_path = os.path.join(base_path, out_name)
        if render_pdf_figure_region(pdf_path, page_index, out_path):
            replace[img_lines[0]] = f"![]({out_name})"
            for k in img_lines[1:]:
                drop.add(k)
            rendered += 1
    if not replace:
        return text
    out = []
    for idx, line in enumerate(lines):
        if idx in drop:
            continue
        out.append(replace.get(idx, line))
    print(f"   [ABBILDUNG] {rendered} zerstückelte Abbildung(en) aus PDF gerendert.")
    return '\n'.join(out)


def inline_flattened_tables(text: str, pdf_path: str, base_path: str) -> str:
    """Fügt für 'flachgewalzte' Tabellen ein aus der Quell-PDF gerendertes Bild ein.

    Hintergrund: Manche Tabellen erfasst die OCR nicht als Pipe-Tabelle, sondern walzt sie in
    einen Riesen-Caption-Absatz ('Tabelle N …' mit dem kompletten Tabelleninhalt) plus
    Pseudo-Überschriften flach – sie sind dann nicht mehr als Tabelle erkennbar.

    Erkennung: Caption-Zeile 'Tabelle N …'/'Table N …' mit auffälliger Länge (> 300 Zeichen)
    und ohne Pipe ('|'). Seitenindex aus dem page-N-Span der Caption (Fallback: nächster Span
    darüber). Das gerenderte Bild wird direkt NACH der Caption eingefügt; der zerlegte Text
    bleibt als Fallback erhalten. No-op, wenn pdf_path/base_path fehlen oder das Rendern
    scheitert (kein pypdfium2 / keine Pfad-Objekte)."""
    if not pdf_path or not base_path or not os.path.exists(pdf_path):
        return text
    lines = text.split('\n')
    # Echte (flachgewalzte) Tabellen-Caption beginnt mit einem page-Span, den Marker beim
    # Tabellenobjekt setzt – Fließtext-Sätze wie "Tabelle 2 zeigt …" tun das nicht. Genau diese
    # Span-am-Zeilenanfang-Bedingung trennt Caption von Tabellen-Verweis im Text.
    LEAD_PAGE_SPAN_RE = re.compile(r'^<span[^>]*id="page-(\d+)-\d+"[^>]*>')
    LEAD_SPAN_RE = re.compile(r'^(?:<span[^>]*>\s*|</span>\s*)+')
    CAP_RE = re.compile(r'^\*{0,2}\s*(?:Tabelle|Table)\s+\d+\b')
    inserts: dict[int, str] = {}     # Zeilenindex → Bildzeile (danach einfügen)
    rendered = 0
    for i, line in enumerate(lines):
        s = line.strip()
        m_span = LEAD_PAGE_SPAN_RE.match(s)
        if not m_span or len(s) <= 300 or '|' in s:
            continue
        core = LEAD_SPAN_RE.sub('', s)
        if not CAP_RE.match(core):
            continue
        page_index = int(m_span.group(1))
        out_name = f"_rendered_table_p{page_index}.jpeg"
        out_path = os.path.join(base_path, out_name)
        if render_pdf_figure_region(pdf_path, page_index, out_path):
            inserts[i] = f"![]({out_name})"
            rendered += 1
    if not inserts:
        return text
    out = []
    for idx, line in enumerate(lines):
        out.append(line)
        if idx in inserts:
            out.append('')
            out.append(inserts[idx])
    print(f"   [TABELLE] {rendered} flachgewalzte Tabelle(n) aus PDF gerendert.")
    return '\n'.join(out)


def _caption_protected_image_lines(lines: list) -> set:
    """Zeilen-Indizes von Bildern, die im selben Heading-Segment wie eine nummerierte
    Abbildungs-/Tabellen-Caption ('Abb./Tab. N') liegen. Solche Bilder werden NIE als
    dekorativ verworfen – z.B. die vier kleinen Panels von Abb. 3.1, die einzeln < 100 px
    sind, zusammen aber die prüfungsrelevante Abbildung bilden."""
    head_idx = [k for k, l in enumerate(lines) if re.match(r'^#{1,6}\s', l)]
    bounds = head_idx + [len(lines)]

    def segment(k):
        prev = [h for h in head_idx if h <= k]
        if not prev:
            return 0, (bounds[0] if head_idx else len(lines))
        s = prev[-1]
        return s, bounds[head_idx.index(s) + 1]

    img_idx = [k for k, l in enumerate(lines)
               if re.match(r'^\s*!\[[^\]]*\]\([^)]+\)\s*$', l)]
    protected = set()
    for k in img_idx:
        s, e = segment(k)
        if any(_CAPTION_RE.match(lines[m].strip()) for m in range(s, e)):
            protected.add(k)
    return protected


def _hide_paragraph(p, indent_cm: float = 1.5):
    """Setzt alle Runs eines Absatzes auf hidden + fügt Einrückung hinzu.
    Entfernt w:numPr (Bullet-Label) und setzt w:vanish in w:pPr/w:rPr."""
    p.paragraph_format.left_indent = Cm(indent_cm)
    for run in p.runs:
        run.font.hidden = True
    pPr = p._p.get_or_add_pPr()
    numPr = pPr.find(qn('w:numPr'))
    if numPr is not None:
        pPr.remove(numPr)
    rPr = pPr.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        pPr.append(rPr)
    vanish = OxmlElement('w:vanish')
    rPr.append(vanish)


def process_markdown_to_docx(doc, block_text, hide_text=False, base_path=None,
                             skip_images=False, headings_as_bold=False):
    """
    Interpretiert Markdown und fügt Inhalte dem Word-Dokument hinzu.
    - Tabellen und Bilder: IMMER sichtbar (nie ausgeblendet).
    - Abbildungs-/Tabellenbeschriftungen: sichtbar, auch wenn hide_text=True.
    - hide_text=True: Text wird hidden (nicht druckbar) + eingerückt; grau für Sichtbarkeit.
    - headings_as_bold=True: Markdown-Headings als fetter Normal-Text statt Word-Heading-Style
      (verhindert Einträge in der Navigationsleiste bei Summary-Body-Text).
    - Horizontale Linien (---) werden übersprungen.
    """
    color_map = {
        'heading1': MM_HEADING_COLORS[1],
        'heading2': MM_HEADING_COLORS[2],
        'heading3': MM_HEADING_COLORS[3],
        'heading4': MM_HEADING_COLORS[4],
        'hidden_text':    RGBColor(0xBB, 0xBB, 0xBB),
        'hidden_heading': RGBColor(0x99, 0x99, 0x99),
    }

    # Im Hidden-Block: Links, HTML und OCR-Artefakte bereinigen
    if hide_text:
        block_text = _clean_for_hidden(block_text)

    lines = block_text.split('\n')
    # Text-Figuren (Caption ohne Bild) bleiben auch im Hidden-Block sichtbar – wie Bild-Figuren.
    textfig_visible = _textfigure_visible_lines(lines) if hide_text else set()
    # Bilder, die zu einer nummerierten Abbildungs-/Tabellen-Caption gehören, NIE als dekorativ
    # verwerfen (kleine Panels einer Abb., z.B. Abb. 3.1, sind prüfungsrelevant).
    caption_protected_imgs = _caption_protected_image_lines(lines)
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # ── HTML-Tags aus Fließtext entfernen (außer Bildzeilen) ──
        if not stripped.startswith('!'):
            stripped = re.sub(r'<[^>]+>', '', stripped)

        # ── Horizontale Linie überspringen (--- / *** / ___) ──
        if _HRULE_RE.match(stripped):
            i += 1
            continue

        # ── Tabelle ──
        if stripped.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            add_markdown_table_to_doc(doc, table_lines)
            continue

        # ── Bild: immer sichtbar (außer wenn skip_images=True oder dekorativ) ──
        img_m = re.match(r'^!\[([^\]]*)\]\(([^)]+)\)$', stripped)
        if img_m:
            if skip_images:
                i += 1
                continue
            img_rel = img_m.group(2)
            img_path = os.path.join(base_path, img_rel) if base_path else img_rel
            if os.path.exists(img_path):
                if _is_decorative_image(img_path) and i not in caption_protected_imgs:
                    # Icons, Logos, Randsymbole überspringen (außer Caption-gebundene Panels)
                    i += 1
                    continue
                try:
                    doc.add_picture(img_path, width=_image_display_width(img_path))
                    doc.add_paragraph()
                except Exception:
                    doc.add_paragraph(f"[Bild nicht eingebettet: {img_rel}]")
            else:
                doc.add_paragraph(f"[Bild nicht gefunden: {img_rel}]")
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        # Abbildungs-/Tabellenbeschriftungen nie ausblenden; Text-Figuren-Inhalt sichtbar lassen
        is_caption = bool(_CAPTION_RE.match(stripped))
        do_hide = hide_text and not is_caption and i not in textfig_visible

        # ── Sub-Bullets (2+ führende Leerzeichen, Strich/Stern) ──
        sub_m = re.match(r'^(\s{2,})([-*])\s+(.+)', line)
        if sub_m:
            depth = len(sub_m.group(1)) // 2
            style = 'List Bullet 3' if depth >= 2 else 'List Bullet 2'
            p = doc.add_paragraph(style=style)
            color = color_map['hidden_text'] if do_hide else None
            add_formatted_text(p, sub_m.group(3), default_color=color)
            p.paragraph_format.space_after = Pt(0)
            if do_hide:
                _hide_paragraph(p)
            i += 1
            continue

        # ── Nummerierte Sub-Items (2+ führende Leerzeichen + Zahl) ──
        num_sub_m = re.match(r'^(\s{2,})(\d+)[.)]\s+(.+)', line)
        if num_sub_m:
            depth = len(num_sub_m.group(1)) // 2
            style = 'List Bullet 3' if depth >= 2 else 'List Bullet 2'
            p = doc.add_paragraph(style=style)
            color = color_map['hidden_text'] if do_hide else None
            add_formatted_text(p, num_sub_m.group(3), default_color=color)
            p.paragraph_format.space_after = Pt(0)
            if do_hide:
                _hide_paragraph(p)
            i += 1
            continue

        # ── Überschriften ──
        if stripped.startswith('#### '):
            hdg_text = stripped[5:]
            if headings_as_bold and not do_hide:
                p = doc.add_paragraph(style='Normal')
                p.add_run(hdg_text).bold = True
            else:
                p = doc.add_heading(level=4)
                color = color_map['hidden_heading'] if do_hide else color_map['heading4']
                add_formatted_text(p, hdg_text, default_color=color)
                if do_hide:
                    _hide_paragraph(p)
        elif stripped.startswith('### '):
            hdg_text = stripped[4:]
            if headings_as_bold and not do_hide:
                p = doc.add_paragraph(style='Normal')
                p.add_run(hdg_text).bold = True
            else:
                p = doc.add_heading(level=3)
                color = color_map['hidden_heading'] if do_hide else color_map['heading3']
                add_formatted_text(p, hdg_text, default_color=color)
                if do_hide:
                    _hide_paragraph(p)
        elif stripped.startswith('## '):
            hdg_text = stripped[3:]
            if headings_as_bold and not do_hide:
                p = doc.add_paragraph(style='Normal')
                p.add_run(hdg_text).bold = True
            else:
                p = doc.add_heading(level=2)
                color = color_map['hidden_heading'] if do_hide else color_map['heading2']
                add_formatted_text(p, hdg_text, default_color=color)
                if do_hide:
                    _hide_paragraph(p)
        elif stripped.startswith('# '):
            hdg_text = stripped[2:]
            if headings_as_bold and not do_hide:
                p = doc.add_paragraph(style='Normal')
                p.add_run(hdg_text).bold = True
            else:
                p = doc.add_heading(level=1)
                color = color_map['hidden_heading'] if do_hide else color_map['heading1']
                add_formatted_text(p, hdg_text, default_color=color)
                if do_hide:
                    _hide_paragraph(p)
        # ── Aufzählung ──
        elif stripped.startswith('* ') or stripped.startswith('- '):
            bullet_m = re.match(r'^[*-]\s+(.*)', stripped)
            bullet_content = bullet_m.group(1) if bullet_m else stripped[2:]
            if not bullet_content.strip():
                i += 1
                continue
            p = doc.add_paragraph(style='List Bullet')
            color = color_map['hidden_text'] if do_hide else None
            add_formatted_text(p, bullet_content, default_color=color)
            p.paragraph_format.space_after = Pt(0)
            if do_hide:
                _hide_paragraph(p)
        # ── Normaler Text ──
        else:
            p = doc.add_paragraph()
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            color = color_map['hidden_text'] if do_hide else None
            add_formatted_text(p, stripped, default_color=color)
            if do_hide:
                _hide_paragraph(p)
        i += 1


def _set_heading_color(heading_paragraph, color: RGBColor):
    """Setzt die Schriftfarbe aller Runs einer Überschrift."""
    for run in heading_paragraph.runs:
        run.font.color.rgb = color


# Überschriften, die standardmäßig übersprungen werden (Referenzen etc.)
_SKIP_HEADINGS = frozenset([
    'referenzen', 'literaturverzeichnis', 'references', 'bibliography',
    'literatur', 'quellen', 'quellenverzeichnis', 'endnoten', 'endnotes',
    'orcid', 'interessenkonflikt', 'conflict of interest',
    'erklärung zur datenverfügbarkeit', 'data availability statement',
    'danksagung', 'acknowledgments', 'kontrollliste',
    'history', 'historie',
])

_SKIP_HEADING_PATTERNS = ('m.sc', 'prof. dr', '@')

def _is_skip_heading(heading: str) -> bool:
    """Gibt True zurück, wenn diese Überschrift standardmäßig übersprungen werden soll."""
    key = normalize_heading(heading)
    if key in _SKIP_HEADINGS:
        return True
    return any(pat in key for pat in _SKIP_HEADING_PATTERNS)


# Front-Matter-Überschriften am Dokumentanfang (Impressum, Inhaltsverzeichnis),
# die standardmäßig übersprungen werden. Titel + Titelbild davor bleiben erhalten.
_FRONT_MATTER_HEADINGS = frozenset([
    'impressum', 'imprint',
    'inhalt', 'inhaltsverzeichnis', 'table of contents', 'contents',
])

# Front-Matter-Überschriften, die nur über einen Präfix erkennbar sind (Titelvarianten).
# 'Checklisten im Buch und online' (Springer-Vorspann) listet Download-Checklisten je
# 'Kapitel N' auf – diese Einträge dürfen extract_chapter NICHT als echtes Kapitel sehen.
_FRONT_MATTER_PREFIXES = ('checklisten im buch',)


def _is_front_matter_heading(key: str) -> bool:
    """True, wenn die (normalisierte) Überschrift ein Front-Matter-Block-Anfang ist."""
    return key in _FRONT_MATTER_HEADINGS or key.startswith(_FRONT_MATTER_PREFIXES)


def strip_front_matter(text: str) -> str:
    """Entfernt Front-Matter-Blöcke (Impressum, Inhaltsverzeichnis) am Dokumentanfang.

    Titel (erste Überschrift) und ggf. Titelbild davor bleiben erhalten. Ein Block läuft
    von seiner Überschrift bis zur nächsten gleich-/höherrangigen Überschrift, die selbst
    kein Front-Matter ist; tiefere Unterabschnitte des Blocks werden mit entfernt.
    Arbeitet zeilenbasiert, um HTML-Spans und Bild-Referenzen unverändert zu lassen.
    """
    out: list[str] = []
    skipping = False
    skip_level = 0
    for line in text.split('\n'):
        m = re.match(r'^(#{1,6})\s+(.*)$', line)
        if m:
            level = len(m.group(1))
            key = normalize_heading(_clean_heading_text(line))
            if skipping and level <= skip_level:
                # Block endet bei gleich-/höherrangiger Überschrift.
                skipping = False
            if _is_front_matter_heading(key):
                skipping = True
                skip_level = level
                continue
            if skipping and level > skip_level:
                continue  # tieferes Front-Matter-Unterkapitel
        if skipping:
            continue
        out.append(line)
    return '\n'.join(out)


_PUBMETA_PATTERNS = [
    r'urheberrechtlich geschützt durch die American Psychological',
    r'copyrighted by the American Psychological',
    r'ausschließlich für (die persönliche|den persönlichen Gebrauch)',
    r'intended solely for the personal',
    r'(Nutzung durch den einzelnen Anwender|use of the individual user)',
    r'(darf nicht (weiter|weit) verbreitet|disseminated broadly)',
    r'©\s*\d{4}\s+American Psychological Association',
    r'\bISSN:',
    r'^\s*\d{4},\s*(Bd\.|Vol\.)\s*\d+',
    r'(online erstveröffentlicht|published online first)',
    r'(fungierte als Action Editor|served as action editor)',
]
_PUBMETA_RE = re.compile('|'.join(_PUBMETA_PATTERNS), re.IGNORECASE)


def strip_publication_metadata(text: str) -> str:
    """Entfernt verstreute Journal-Publikations-Metadaten (Copyright-/Personal-Use-Hinweis,
    ISSN/Heft/DOI-Kopf, 'Online First'/Action-Editor), die die OCR zwischen den eigentlichen
    Inhalt streut und die den Lesefluss stören. Arbeitet zeilenweise; Bild- und
    Überschriftenzeilen bleiben unangetastet. Behält bewusst Autorenbeiträge, Förderung,
    ORCID und Korrespondenz. No-op bei Dokumenten ohne diese Muster."""
    out: list[str] = []
    removed = False
    for line in text.split('\n'):
        stripped = line.strip()
        if stripped and not stripped.startswith(('!', '#')) and _PUBMETA_RE.search(stripped):
            removed = True
            continue
        out.append(line)
    if not removed:
        return text
    # durch entfernte Zeilen entstandene Mehrfach-Leerzeilen auf eine reduzieren
    return re.sub(r'\n{3,}', '\n\n', '\n'.join(out))


def collapse_duplicate_title(text: str) -> str:
    """Entfernt einen doppelten Titel-/Journal-Banner am Dokumentanfang.

    Erscheint der Dokumenttitel im Front-Bereich (vor der ersten inhaltlichen Überschrift wie
    ZUSAMMENFASSUNG/ABSTRACT/SCHLÜSSELWÖRTER/EINLEITUNG, einer nummerierten Überschrift, oder
    einer römisch/ziffern-gelabelten Artikelüberschrift wie in Sammelbänden, z.B. "I. Titel")
    mehrfach, wird alles vor der letzten Titel-Wiederholung verworfen – Journal-Banner,
    Zitationsdaten und die erste Titel-Dublette fallen weg, der saubere Artikelkopf
    (Titel + Autoren + Affiliationen + Abstract) bleibt. No-op ohne wiederholten Titel.
    Ohne die Artikel-Label-Erkennung würde die Funktion bei Sammelbänden mit fett formatierten
    Artikelüberschriften (z.B. "### **0. Vorwort**") nie beim echten Front-Ende abbrechen und
    stattdessen bis zu einer beliebigen, in mehreren Artikeln wiederkehrenden Abschnittsüberschrift
    (z.B. "Literaturverzeichnis") weiterscannen – mit dem Risiko, ganze Artikel fälschlich als
    "Front-Matter-Dublette" zu verwerfen."""
    lines = text.split('\n')
    # Gegen den BEREINIGTEN Text prüfen (ohne HTML-Anker wie <span id="page-…">, die Marker-OCR
    # praktisch vor jede Überschrift setzt) – sonst matcht '^\d' nie und der Front-Matter-Abbruch
    # feuert nicht, wodurch die Schleife bis weit in den Kapitelinhalt hinein weiterläuft.
    content_re = re.compile(
        r'(?i)^(ZUSAMMENFASSUNG|ABSTRACT|SCHLÜSSELWÖRTER|KEYWORDS|EINLEITUNG)\b'
        r'|^\d')
    seen: dict = {}
    last_dup = None
    for idx, ln in enumerate(lines):
        if not re.match(r'^#{1,6}\s+\S', ln):
            continue
        norm_text = _clean_heading_text(ln)
        if content_re.match(norm_text):
            break                                   # Front-Bereich endet bei erstem Inhalt
        if _ROMAN_LABEL_RE.match(norm_text) or _DIGIT_LABEL_RE.match(norm_text):
            break                                   # Artikel-/Kapitelüberschrift (Sammelband) = Inhalt
        norm = normalize_heading(norm_text)
        if norm and norm in seen:
            last_dup = idx                          # spätere (Artikelkopf-)Variante behalten
        else:
            seen[norm] = idx
    return '\n'.join(lines[last_dup:]).strip() if last_dup is not None else text


def insert_intro_heading(text: str) -> str:
    """Gibt dem unbeschrifteten Artikel-Intro eine eigene '## Einleitung'-Überschrift.

    In vielen Artikeln folgt der Einleitungstext direkt der Schlüsselwörter-Liste, ohne eigene
    Überschrift – er hängt dann als Body an der 'SCHLÜSSELWÖRTER'-Sektion und wird vom
    Summarizer (der nach Überschriften gliedert) auf die Stichwörter reduziert. Diese Funktion
    trennt den Intro-Text als eigene Sektion ab, damit er in der Pflichtliste landet und
    zwingend zusammengefasst wird. No-op ohne Keyword-Heading oder ohne substanziellen
    Intro-Text (z.B. Lehrbuch-Kapitel); idempotent (bricht ab, wenn 'Einleitung' schon existiert)."""
    m = re.search(r'(?im)^#{1,6}\s*\**\s*(SCHLÜSSELWÖRTER|SCHLAGWÖRTER|KEYWORDS)\b.*$', text)
    if not m or re.search(r'(?im)^#{1,6}\s*\**\s*Einleitung\b', text):
        return text
    body_start = m.end()
    nxt = re.search(r'(?m)^#{1,6}\s+\S', text[body_start:])
    body_end = body_start + nxt.start() if nxt else len(text)
    paras = re.split(r'\n\s*\n', text[body_start:body_end].strip())
    if len(paras) < 2:
        return text                                 # nur Stichwörter, kein Intro
    intro = '\n\n'.join(paras[1:]).strip()
    if len(intro) < 600:
        return text
    new_body = f"\n\n{paras[0]}\n\n## Einleitung\n\n{intro}\n\n"
    return text[:body_start] + new_body + text[body_end:]


def _advance_counter(counters: list, level: int) -> str:
    """
    Erhöht den Zähler für `level` (1-basiert) und setzt tiefere Ebenen zurück.
    Gibt die kompakte Nummerierung zurück – Null-Ebenen werden übersprungen,
    damit Ebenensprünge (z.B. level 2 → level 4) keine '0.0'-Segmente erzeugen.
    Beispiel: counters=[4,0,0,1] → '4.1' statt '4.0.0.1'
    """
    counters[level - 1] += 1
    for i in range(level, len(counters)):
        counters[i] = 0
    return '.'.join(str(counters[i]) for i in range(level) if counters[i] > 0)


def _parent_level_from_chapter(parent_chapter: str) -> int:
    """
    Leitet den Heading-Level des Elternkapitels aus der Kapitelnummer ab.
    '4.2.1.2' → 4 Punkte-getrennte Teile → Heading 5
    Annahme: Top-Level-Kapitel (z.B. '1') sind Heading 2, da Heading 1 der Dokumenttitel ist.
    """
    return len(parent_chapter.split('.')) + 1


def _normalize_chapter_num(num: str) -> str:
    """Säubert eine (evtl. OCR-verschmutzte) Kapitelnummer.
    Kollabiert Mehrfachpunkte ('3...1' → '3.1') und entfernt führende/abschließende Punkte
    ('1.' → '1'). Liefert reine Punkt-getrennte Ziffernfolgen.
    """
    num = re.sub(r'\.{2,}', '.', num)   # 3...1 → 3.1
    return num.strip('.')               # 1.  → 1 ,  .1 → 1


def _join_number_and_rest(num: str, rest: str) -> str:
    """Fügt eine Kapitelnummer und den Resttext sauber zusammen (genau ein Trenn-Whitespace
    bei vorhandenem Titel, kein Trailing-Whitespace bei reiner Nummer)."""
    rest = rest or ""
    if rest and not rest[:1].isspace():
        rest = f" {rest}"
    return f"{num}{rest}"


def _prefix_chapter_number(heading_text: str, prefix: str) -> str:
    """
    Setzt den Elternpräfix vor die Kapitelnummer einer Überschrift.
    '1 EINLEITUNG'  + '4.2.1.2' → '4.2.1.2.1 EINLEITUNG'
    '4.1 Abschnitt' + '4.2.1.2' → '4.2.1.2.4.1 Abschnitt'
    '1. EINLEITUNG' + '6.2.4.1' → '6.2.4.1.1 EINLEITUNG'  (kein Trailing-Punkt)
    Überschriften ohne führende Zahl bleiben unverändert.
    """
    m = re.match(r'^(\d[\d\.]*)(.*)', heading_text.strip())
    if m:
        num = _normalize_chapter_num(m.group(1))
        return _join_number_and_rest(f"{prefix}.{num}", m.group(2))
    return heading_text


def _rebase_chapter_number(heading_text: str, chapter_root: str, parent_chapter: str) -> str:
    """
    Ersetzt die Nummer eines extrahierten Kapitels durch das Elternkapitel.
    chapter_root='1.3', parent='2.4.1.1':
      '1.3 Titel'   → '2.4.1.1 Titel'
      '1.3.2 Titel' → '2.4.1.1.2 Titel'
      '1.3'         → '2.4.1.1'        (auch reine Nummer ohne Titel)
    Nicht-Nachfahren werden regulär präfixiert (Fallback).
    """
    m = re.match(r'^(\d[\d.]*?)(\s.*|$)', heading_text.strip())
    if not m:
        return heading_text
    num, rest = _normalize_chapter_num(m.group(1)), m.group(2)
    if num == chapter_root:
        return _join_number_and_rest(parent_chapter, rest)
    if num.startswith(chapter_root + '.'):
        return _join_number_and_rest(f"{parent_chapter}.{num[len(chapter_root) + 1:]}", rest)
    return _join_number_and_rest(f"{parent_chapter}.{num}", rest)   # Fallback


def build_translation_word_document(translated_text: str, output_path: str, base_path: str = None):
    """Erstellt ein Word-Dokument aus dem übersetzten Text.
    Preamble (Abstract) eingerückt; Skip-Headings (Referenzen, Historie, Autor) gefiltert.
    """
    print(f"--- Erstelle Übersetzungs-Word-Dokument -> {output_path} ---")

    text = _strip_kontrollliste(translated_text)
    text = normalize_heading_levels(text)
    text = _strip_ocr_y_prefix(text)
    text = _compress_heading_levels(text)
    sections = parse_sections(text)

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for section in sections:
        heading = section['heading']
        body    = section['body']
        level   = section['level']

        if heading == '__preamble__':
            if body.strip():
                before = len(doc.paragraphs)
                process_markdown_to_docx(doc, body, hide_text=False, base_path=base_path, skip_images=True)
                for p in doc.paragraphs[before:]:
                    p.paragraph_format.left_indent = Cm(1.5)
            continue

        clean = _clean_heading_text(heading)
        if _is_skip_heading(clean):
            continue

        h = doc.add_heading(clean, level=level)
        _set_heading_color(h, MM_HEADING_COLORS.get(level, MM_HEADING_COLORS[9]))

        if body.strip():
            process_markdown_to_docx(doc, body, hide_text=False, base_path=base_path)

    doc.save(output_path)
    print("Übersetzungs-Word-Dokument erstellt.")


def build_interleaved_word_document(translated_text: str, summary_text: str, qa_text: str,
                                    output_path: str, base_path: str = None,
                                    parent_chapter: str = None, parent_level: int = None,
                                    skip_references: bool = True,
                                    questions_path: str = None,
                                    doc_title: str = "Lernskript",
                                    extracted_chapter: str = None,
                                    supplement_map: dict = None,
                                    title_as_parent: bool = False):
    """
    Erstellt ein Word-Dokument, bei dem Zusammenfassung und Originaltext
    kapitelweise verschränkt sind.

    parent_chapter: Elternkapitel im Zieldokument, z.B. '4.2.1.2'.
      Wenn angegeben, werden alle Kapitelnummern mit diesem Präfix versehen
      ('1 EINLEITUNG' → '4.2.1.2.1 EINLEITUNG') und Heading-Level entsprechend
      verschoben. 'Lernskript'-Titel und QA-Block werden dann weggelassen.
    parent_level:   Heading-Level des Elternkapitels (Standard: aus parent_chapter
      automatisch berechnet, z.B. '4.2.1.2' → 5).
    """
    print(f"--- Erstelle interleaved Word-Dokument -> {output_path} ---")

    # Audit: Abbildungen/Tabellen-Captions auf sichtbare Abbildung prüfen (Warnung bei Lücken).
    audit_figure_captions(translated_text, base_path=base_path)

    # Elternkapitel-Logik
    if parent_chapter:
        lvl_shift = (parent_level if parent_level else _parent_level_from_chapter(parent_chapter))
        print(f"    Einfügemodus: Präfix '{parent_chapter}', Heading-Shift +{lvl_shift}")
    else:
        lvl_shift = 0

    # Strukturelle Neu-Nummerierung: Im reinen Einbettungs-Modus (parent_chapter, KEIN
    # extrahiertes Einzelkapitel) wird die gesamte Artikel-Gliederung allein aus der
    # Überschriften-Hierarchie neu nummeriert. Die OCR-Eigennummern (z.B. "1.", "3.") sind
    # nur teilweise vorhanden und inkonsistent → sie werden verworfen.
    structural_renumber = bool(parent_chapter and not extracted_chapter)

    counters = [0] * 9  # Zähler je Heading-Ebene für Auto-Nummerierung

    # --- QA vorbereiten: Textgrundlage-Map für Kommentare ---
    has_qa = qa_text and qa_text.strip() not in ("", "Keine Leitfragen zur Prüfung übergeben.")
    qa_items = parse_qa_response(qa_text) if has_qa else []

    # Antwort-Einheiten flach machen: pro Unterpunkt eine Einheit, sonst eine pro Frage.
    # Jede Einheit trägt eigene Textgrundlage + Beleg → eigene Markierung im Originaltext.
    qa_units: list = []
    for item in qa_items:
        if item.get('subs'):
            for sub in item['subs']:
                qa_units.append({'num': item['num'], 'label': sub.get('label', ''),
                                 'textgrundlage': sub.get('textgrundlage', ''),
                                 'beleg': sub.get('beleg', '')})
        else:
            qa_units.append({'num': item['num'], 'label': '',
                             'textgrundlage': item.get('textgrundlage', ''),
                             'beleg': item.get('beleg', '')})

    # Original-Kapitelnummer → neue (rebasierte) Nummer im Dokument. Wird im Sektions-Loop
    # gefüllt und remappt die in der QA angezeigte "Quelle" auf eine im Dokument navigierbare Nummer.
    tg_renumber: dict = {}

    textgrundlage_map: dict = {}  # normalize_heading(textgrundlage) → [qa_unit, ...]
    for unit in qa_units:
        for key in _textgrundlage_keys(unit['textgrundlage']):
            textgrundlage_map.setdefault(key, [])
            if unit not in textgrundlage_map[key]:
                textgrundlage_map[key].append(unit)

    comment_list: list = []   # [(comment_id, comment_text)]
    comment_id: list  = [0]  # mutable int-wrapper

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # --- Lernskript-Titel (Standalone-Modus) ---
    if not parent_chapter:
        h = doc.add_heading(doc_title, level=1)
        _set_heading_color(h, MM_HEADING_COLORS[1])

    # --- Heading-Level normalisieren & Sections parsen ---
    translated_text = _strip_kontrollliste(translated_text)
    translated_text = normalize_heading_levels(translated_text)
    translated_text = _strip_ocr_y_prefix(translated_text)
    summary_text    = _strip_kontrollliste(summary_text)
    summary_text    = _strip_selbstpruefung(summary_text)
    summary_text    = normalize_heading_levels(summary_text)
    summary_text    = _strip_ocr_y_prefix(summary_text)
    orig_sections   = parse_sections(translated_text)
    sum_sections    = parse_sections(summary_text)

    # Lookup: normalisierter Heading → Zusammenfassungstext
    sum_lookup = {}
    for s in sum_sections:
        if s['heading'] == '__preamble__':
            continue
        sum_lookup[normalize_heading(_clean_heading_text(s['heading']))] = s['body']

    # Scoped lookup: parent_key → {child_key → body}
    # Verhindert Conflation gleichnamiger Sub-Headings (z.B. "Off the job" unter jeder Kompetenz).
    # Ohne Scoping liefert sum_lookup["off the job"] immer den letzten Eintrag (letzte Kompetenz).
    sum_scoped: dict[str, dict[str, str]] = {}
    _scp_parent: str | None = None
    for _ss in sum_sections:
        if _ss['heading'] == '__preamble__':
            continue
        _ssk = normalize_heading(_clean_heading_text(_ss['heading']))
        if _ss['level'] <= 2:
            _scp_parent = _ssk
            sum_scoped[_ssk] = {'__self__': _ss['body']}
        elif _scp_parent is not None:
            sum_scoped[_scp_parent][_ssk] = _ss['body']

    # Dokument-Typ bestimmen: hat es nummerierte Kapitel?
    has_numbered_chapters = any(
        re.match(r'^\d', _clean_heading_text(s['heading']).strip())
        for s in orig_sections
        if s['heading'] != '__preamble__'
    )

    # Häufigkeit unnummerierter Überschriften zählen.
    # Einzigartige (freq=1) sind echte Kapitelthemen (z.B. "Agilität", "Analytisches Denken").
    # Wiederkehrende (freq>1) sind Frage-Muster unter jedem Unterkapitel ("Was ist das?").
    _unnumbered_freq: dict[str, int] = {}
    for _s in orig_sections:
        if _s['heading'] == '__preamble__':
            continue
        _k = normalize_heading(_clean_heading_text(_s['heading']))
        if not re.match(r'^\d', _k):
            _unnumbered_freq[_k] = _unnumbered_freq.get(_k, 0) + 1
    _auto_parent: str | None = None  # z. B. "5.3" – aktuelles nummeriertes Elternkapitel
    _auto_counter: int = 0
    _max_child: int = 0  # höchstes unmittelbares Kind-Segment unter parent_chapter (für Lernfragen-Nr.)

    # Vorberechnung: hat jede Section mindestens einen Nachfolger mit Summary-Inhalt?
    # Verhindert sichtbare Leer-Überschriften bei Elternkapiteln, deren Kinder alle
    # kein Summary haben (z.B. "5.3.27 Rückmeldung zu Konfliktverhalten" → nur Off/On the job).
    # Spiegelt dieselbe Kinder-Logik wie has_children:
    #   a) tiefere Ebene gilt immer als Nachfolger
    #   b) nummerierter Parent + unnummerierter Geschwister → konzeptuelles Kind
    #      (z.B. "5.3 Überblick" → "Agilität" auf gleicher Ebene)
    _any_visible_desc: list[bool] = [False] * len(orig_sections)
    for _pi in range(len(orig_sections)):
        _ps = orig_sections[_pi]
        if _ps['heading'] == '__preamble__':
            continue
        _plevel = _ps['level']
        _ps_numbered = bool(re.match(r'^\d', normalize_heading(_clean_heading_text(_ps['heading']))))
        for _ci in range(_pi + 1, len(orig_sections)):
            _cs = orig_sections[_ci]
            _cs_level = _cs['level']
            if _cs_level < _plevel:
                break  # Aufgestiegen = außerhalb des Scopes
            if _cs_level == _plevel:
                # Gleiche Ebene: nur als konzeptuelles Kind wenn parent nummeriert und
                # dieses Section unnummeriert ist (nummerierter → unnummerierter Folger).
                _cs_k = normalize_heading(_clean_heading_text(_cs['heading']))
                _cs_numbered = bool(re.match(r'^\d', _cs_k))
                if _cs_numbered or not _ps_numbered:
                    break  # Nummeriertes Geschwister oder keine konzeptuelle Eltern-Kind-Relation
                # Unnummeriert nach nummeriertem Parent → zählt als Kind, weiter suchen
            _clk = normalize_heading(_clean_heading_text(_cs['heading']))
            if sum_lookup.get(_clk, '').strip():
                _any_visible_desc[_pi] = True
                break

    # Pre-pass: welche Sections werden sichtbar? Nur diese erhalten sequentielle Auto-Nummern.
    # Ohne diesen Pre-pass entstehen Lücken wie 5.3.1, 5.3.5, 5.3.8 weil unsichtbare Sections
    # trotzdem nummeriert werden.
    _visible_section_indices: set[int] = set()
    for _vi, _vs in enumerate(orig_sections):
        if _vs['heading'] == '__preamble__':
            continue
        if re.search(r'^(?:\\_)+\s*$', _vs['body'], re.MULTILINE):
            continue  # Platzhalter-Section
        _vk = normalize_heading(_clean_heading_text(_vs['heading']))
        if sum_lookup.get(_vk, '').strip() or _any_visible_desc[_vi]:
            _visible_section_indices.add(_vi)

    # title_as_parent: Die Wurzel-Sektion (Artikeltitel, Tiefe 0) ist der Anker des
    # Dokuments und trägt direkt das Elternkapitel (z.B. 6.2.4.3). Sie immer als sichtbar
    # markieren, auch ohne eigene Zusammenfassung – sonst bekäme der Titel keine Nummer.
    if structural_renumber and title_as_parent:
        for _ti, _ts in enumerate(orig_sections):
            if _ts['heading'] != '__preamble__':
                _visible_section_indices.add(_ti)
                break

    # Tiefen-Prepass für die strukturelle Neu-Nummerierung: bildet die (inkonsistenten,
    # lückenhaften) OCR-Überschriftenebenen auf konsistente Tiefen 1,2,3… ab – analog zu
    # _compress_heading_levels, aber über die Sektions-Levels statt über Markdown-Zeilen und
    # ohne Sonderbehandlung nummerierter Headings (deren Eigennummern werden ohnehin verworfen).
    depth_map: dict[int, int] = {}
    if structural_renumber and title_as_parent:
        # Variante "Titel = Elternkapitel": Die erste Überschrift (Artikeltitel) ist die
        # permanente Wurzel (Tiefe 0 → trägt direkt das Elternkapitel, z.B. 6.2.4.1). Alle
        # weiteren Überschriften werden per Outline-Stack relativ verschachtelt, wobei die
        # Wurzel NIE vom Stack genommen wird → auch gleichrangige OCR-H1 (Studie 1/2/3) werden
        # zu Kindern des Titels (Tiefe 1), ihre Unterabschnitte zu Tiefe 2 usw.
        _stack: list[int] = []
        _first = True
        for _di, _ds in enumerate(orig_sections):
            if _ds['heading'] == '__preamble__':
                continue
            _lvl = _ds['level']
            if _first:
                depth_map[_di] = 0
                _first = False
                continue
            while _stack and _stack[-1] >= _lvl:
                _stack.pop()
            depth_map[_di] = len(_stack) + 1
            _stack.append(_lvl)
    elif structural_renumber:
        # Standard-Einbettung: bildet die (inkonsistenten, lückenhaften) OCR-Überschriftenebenen
        # auf konsistente Tiefen 1,2,3… ab – analog zu _compress_heading_levels, aber über die
        # Sektions-Levels statt über Markdown-Zeilen und ohne Sonderbehandlung nummerierter
        # Headings (deren Eigennummern werden ohnehin verworfen).
        _ocr_to_actual: dict[int, int] = {}
        _prev_actual = 0
        for _di, _ds in enumerate(orig_sections):
            if _ds['heading'] == '__preamble__':
                continue
            _lvl = _ds['level']
            if _lvl in _ocr_to_actual:
                _actual = _ocr_to_actual[_lvl]
            elif _lvl > _prev_actual + 1:
                _actual = _prev_actual + 1
                _ocr_to_actual[_lvl] = _actual
            else:
                _actual = _lvl
                _ocr_to_actual[_lvl] = _actual
            _prev_actual = _actual
            depth_map[_di] = _actual

    _current_competency_key: str | None = None  # aktuell verarbeitete Kompetenz für Scoped Lookup
    _struct_intro_attached = False  # Einleitungs-Summary im structural_renumber-Modus nur einmal anhängen

    # --- Interleaved Aufbau ---
    for idx, section in enumerate(orig_sections):
        if section['heading'] == '__preamble__':
            sum_preamble = next(
                (s['body'] for s in sum_sections if s['heading'] == '__preamble__'), ''
            )
            # KI-Metatext ("Absolut! Hier ist...") aus Summary-Preamble entfernen
            sum_preamble = re.sub(
                r'^[^\n]*(?:absolut|hier ist|lernorientierte)[^\n]*\n?',
                '', sum_preamble, flags=re.IGNORECASE
            ).strip()
            if sum_preamble:
                process_markdown_to_docx(doc, sum_preamble, hide_text=False, base_path=base_path)
            if section['body']:
                process_markdown_to_docx(doc, section['body'], hide_text=True, base_path=base_path)
            continue

        level     = section['level']
        heading   = section['heading']
        orig_body = section['body']

        clean_heading = _clean_heading_text(heading)
        lookup_key = normalize_heading(clean_heading)  # vor Präfix-Addition für Lookups
        orig_clean_heading = clean_heading  # Originalüberschrift (vor Rebase) für Quellen-Remapping

        if skip_references and _is_skip_heading(clean_heading):
            continue

        # Platzhalter-Sektionen (OCR-Notizlinien \_\_\_ im Body) komplett überspringen
        if re.search(r'^(?:\\_)+\s*$', orig_body, re.MULTILINE):
            continue

        # Normalisierung: "Titel 7.2.2" → "7.2.2 Titel" (OCR-Artefakt: Kapitelnummer am Zeilenende)
        if extracted_chapter:
            _ec_esc = re.escape(extracted_chapter)
            _m_tail = re.match(r'^(.+?)\s+(\d[\d.]*)\s*$', clean_heading)
            if _m_tail and re.match(rf'^{_ec_esc}\.', _m_tail.group(2)):
                clean_heading = f"{_m_tail.group(2)} {_m_tail.group(1).rstrip()}"

        # In eingebettetem Modus (parent_chapter gesetzt, keine nummerierten Kapitel):
        # Level-1-Headings sind OCR-Artefakte (Artikeltitel, Journal-Metadaten).
        # Originaltext als versteckten Text erhalten, aber KEINE Nav-Überschrift erzeugen,
        # damit der Counter für level-2+ sauber bei 1 beginnt (5.1.1, 5.1.2 ...).
        if parent_chapter and not structural_renumber and not has_numbered_chapters and level == 1:
            if len(orig_body.strip()) > 800:
                # Substanzieller Einleitungstext (Abstract/Intro): Summary anzeigen wenn vorhanden.
                _skipped_sum = sum_lookup.get(lookup_key, '') or sum_lookup.get('einleitung', '')
                if _skipped_sum.strip():
                    process_markdown_to_docx(doc, _skipped_sum, hide_text=False, base_path=base_path)
            if orig_body.strip():
                process_markdown_to_docx(doc, orig_body, hide_text=True, base_path=base_path)
            continue

        if structural_renumber:
            # OCR-Eigennummer verwerfen und allein aus der Hierarchie-Tiefe neu nummerieren.
            _title = re.sub(r'^\d[\d.]*\s+', '', clean_heading).strip()
            _depth = depth_map.get(idx, level)
            if idx in _visible_section_indices:
                if _depth <= 0:
                    # Wurzel/Artikeltitel (nur bei title_as_parent): trägt direkt das Elternkapitel.
                    clean_heading = _join_number_and_rest(parent_chapter, _title)
                else:
                    local_num = _advance_counter(counters, _depth)
                    number = f"{parent_chapter}.{local_num}" if local_num else parent_chapter
                    clean_heading = _join_number_and_rest(number, _title)
            else:
                # Unsichtbare Sektion (wird ausgeblendet): keine Nummer vergeben.
                clean_heading = _title
        elif parent_chapter:
            if re.match(r'^\d', clean_heading):
                if extracted_chapter:
                    clean_heading = _rebase_chapter_number(clean_heading, extracted_chapter, parent_chapter)
                else:
                    clean_heading = _prefix_chapter_number(clean_heading, parent_chapter)
            elif not has_numbered_chapters:
                # Unnummeriert ohne nummerierte Kapitel: hier auto-nummerieren.
                # Bei has_numbered_chapters übernimmt der Auto-Block unten
                # (verhindert Doppel-Nummerierung).
                local_num = _advance_counter(counters, level)
                number = f"{parent_chapter}.{local_num}" if local_num else parent_chapter
                clean_heading = _join_number_and_rest(number, clean_heading)

        if parent_chapter:
            num_m = re.match(r'^([\d.]+)\b', clean_heading)
            if num_m:
                display_level = min(len(num_m.group(1).split('.')) + 1, 9)
            elif _is_box_heading(clean_heading) or (extracted_chapter and level > 1 and not re.match(r'^\d', clean_heading)):
                # Kästen + nicht-nummerierte Level-2+-Headings im Einbette-Modus: 2 Ebenen unter Elternkapitel.
                display_level = min(lvl_shift + 2, 9)
            else:
                display_level = min(level + lvl_shift, 9)
        else:
            display_level = level

        sum_body = sum_lookup.get(lookup_key, '')
        # Structural-Renumber: Die Intro-Zusammenfassung (synthetische "Einleitung"-Sektion) matcht
        # keinen Original-Heading. Sie an die erste Tiefe-1-Sektion (Artikeltitel, z.B. 6.2.4.1.1)
        # anhängen, wenn diese selbst keine eigene Zusammenfassung hat.
        if (structural_renumber and not _struct_intro_attached
                and depth_map.get(idx) == (0 if title_as_parent else 1)
                and not sum_body.strip()):
            sum_body = sum_lookup.get('einleitung', '')
            _struct_intro_attached = True
        # clean_heading kann nach Rebase bereits eine Zahl vorne haben (tail-normalisierte Headings).
        originally_numbered = bool(re.match(r'^\d', lookup_key)) or bool(re.match(r'^\d', clean_heading))

        # Auto-Nummerierung: einzigartige unnummerierte Überschriften (freq=1) nach einem
        # nummerierten Kapitel erhalten automatisch eine Unterkapitelnummer (z.B. 5.3.1).
        # Wiederkehrende Überschriften ("Was ist das?" etc.) bleiben unnummeriert.
        # Nur sichtbare Sections (_visible_section_indices) erhalten eine Nummer → keine Lücken.
        # Im strukturellen Neu-Nummerierungs-Modus ist die Nummerierung oben abschließend erfolgt.
        if has_numbered_chapters and not structural_renumber:
            if originally_numbered:
                # clean_heading wurde oben bereits präfixiert/rebased (parent_chapter-Pfad),
                # daher die Elternnummer direkt daraus übernehmen – NICHT aus dem rohen
                # lookup_key. Sonst fehlt das Elternpräfix und ein Trailing-Punkt der
                # OCR-Nummer ("3.") erzeugt Doppelpunkte wie "3..1".
                _m = re.match(r'^(\d[\d.]*)', clean_heading.strip())
                _auto_parent = _normalize_chapter_num(_m.group(1)) if _m else None
                _auto_counter = 0
                _current_competency_key = None
            elif (_auto_parent is not None and _unnumbered_freq.get(lookup_key, 0) == 1
                  and not (extracted_chapter and _is_box_heading(clean_heading))
                  and not (extracted_chapter and level > 1)):
                # Unique unnummerierte Unterüberschrift: auto-nummerieren.
                # Kästen (Fokus/Studie/…) und Level-2+-Headings im Einbette-Modus ausgenommen
                # – sie erscheinen als eigenständige Lernobjekte ohne Kapitelnummer.
                originally_numbered = True
                _current_competency_key = lookup_key
                if idx in _visible_section_indices:
                    _auto_counter += 1
                    clean_heading = f"{_auto_parent}.{_auto_counter} {clean_heading}"
                _dots = clean_heading.split()[0].count('.')
                display_level = min(_dots + 2, 9)

        # Quellen-Remapping: Original-Überschrift (Schlüssel) → finale Dokumentnummer.
        # clean_heading trägt hier bereits die rebasierte/aufgebaute Nummer.
        _newnum_m = re.match(r'^([\d.]+)', clean_heading)
        if _newnum_m:
            for _k in _textgrundlage_keys(orig_clean_heading):
                tg_renumber.setdefault(_k, _newnum_m.group(1))
            # Höchstes unmittelbares Kind-Segment unter parent_chapter mitführen, damit die
            # Lernfragen die nächste freie Geschwister-Nummer erhalten (z.B. nach 3.2.3.8 → 3.2.3.9).
            if parent_chapter:
                _full = _newnum_m.group(1).rstrip('.')
                if _full.startswith(parent_chapter + '.'):
                    _seg = _full[len(parent_chapter) + 1:].split('.')[0]
                    if _seg.isdigit():
                        _max_child = max(_max_child, int(_seg))

        # Scoped lookup: Sub-Sections (freq>1, wiederkehrend) unter der aktuellen Kompetenz.
        # Verhindert dass sum_lookup["off the job"] immer den letzten Summary-Eintrag liefert.
        if not originally_numbered and _current_competency_key:
            _scoped = sum_scoped.get(_current_competency_key, {}).get(lookup_key, '')
            if _scoped.strip():
                sum_body = _scoped

        _next_sec = orig_sections[idx + 1] if idx + 1 < len(orig_sections) else None
        _next_is_unnumbered = (
            _next_sec is not None and
            _next_sec['heading'] != '__preamble__' and
            not re.match(r'^\d', _clean_heading_text(_next_sec['heading']).strip())
        )
        has_children = (
            _next_sec is not None and
            _next_sec['heading'] != '__preamble__' and
            # Echte Kinder: tiefere Ebene ODER nummerierte Section mit unnummeriertem Folger.
            # Geschwister (unnummeriert → unnummeriert gleicher Ebene) zählen NICHT als Kinder.
            (_next_sec['level'] > level or (originally_numbered and _next_is_unnumbered)) and
            # Nur als Elternknoten sichtbar wenn min. ein Nachfolger Summary-Inhalt hat —
            # verhindert goldene Leer-Überschriften ohne sichtbaren Folgeinhalt.
            _any_visible_desc[idx]
        )
        # Strukturwurzel im title_as_parent-Modus (Artikeltitel, Tiefe 0): trägt das
        # Elternkapitel und ist der Anker des Dokuments → immer als sichtbare Überschrift,
        # auch ohne eigene Zusammenfassung und ohne von der flachen OCR-Ebene erkannte Kinder.
        is_struct_root = structural_renumber and title_as_parent and depth_map.get(idx) == 0

        # Nur wirklich leere Sektionen überspringen (kein Original, kein Summary, keine Kinder).
        if not sum_body.strip() and not orig_body.strip() and not has_children and not is_struct_root:
            continue

        # Überschrift sichtbar: wenn Zusammenfassung vorhanden ODER Elternkapitel mit Kindern.
        # Sonst ausgeblendet — Originalinhalt bleibt im Dokument, ist im Summary-View aber
        # unsichtbar → kein leerer Gliederungspunkt, kein verlorener Inhalt.
        show_heading_visible = bool(sum_body.strip()) or has_children or is_struct_root
        if show_heading_visible:
            if has_numbered_chapters:
                nav_worthy = originally_numbered or (extracted_chapter and _is_box_heading(clean_heading))
            elif parent_chapter:
                nav_worthy = True
            else:
                nav_worthy = display_level <= 2
            if nav_worthy:
                h = doc.add_heading(clean_heading, level=display_level)
                _set_heading_color(h, MM_HEADING_COLORS.get(display_level, MM_HEADING_COLORS[9]))
            else:
                h = doc.add_paragraph(style='Normal')
                h.add_run(clean_heading).bold = True
        else:
            # Überschrift ausblenden wie den Originaltext darunter
            h = doc.add_paragraph(style='Normal')
            h.add_run(clean_heading).bold = True
            _hide_paragraph(h, indent_cm=1.5)

        # Zusammenfassung + Kommentar-Erkennung
        if sum_body.strip():
            before = len(doc.paragraphs)
            process_markdown_to_docx(doc, sum_body, hide_text=False, base_path=base_path,
                                     headings_as_bold=True)

            # Schlüssel zur Zuordnung Sektion → QA (vor-Präfix-Heading + Last-Part-Fallback)
            match_keys = [lookup_key]
            last_part = re.sub(r'^[\d.]+\s*', '', lookup_key).strip()
            if last_part:
                match_keys.append(last_part)

            # --- Ergänzungen aus der Nachbearbeitung (Schritt 5b) anfügen ---
            if supplement_map:
                seen_sup: set = set()
                for k in match_keys:
                    for sup in supplement_map.get(k, []):
                        if not sup or sup in seen_sup:
                            continue
                        seen_sup.add(sup)
                        p_sup = doc.add_paragraph()
                        r_tag = p_sup.add_run("[Ergänzt durch Nachbearbeitung] ")
                        r_tag.bold = True
                        r_tag.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
                        p_sup.add_run(sup)

            new_paras = doc.paragraphs[before:]
            first_para = next((p for p in new_paras if p.text.strip()), None)
            last_para  = next((p for p in reversed(new_paras) if p.text.strip()), first_para)

            # Word-Kommentar pro Antwort-Einheit (Frage bzw. Unterpunkt) – an der Beleg-Stelle
            # verankert (Fallback: ganze Sektion). Mehrteilige Fragen erhalten so je Unterpunkt
            # eine eigene Markierung.
            if first_para and textgrundlage_map:
                units = []
                seen_u: set = set()
                for k in match_keys:
                    for unit in textgrundlage_map.get(k, []):
                        uid = id(unit)
                        if uid not in seen_u:
                            seen_u.add(uid)
                            units.append(unit)
                for unit in sorted(units, key=lambda u: (u['num'], u['label'])):
                    beleg = unit.get('beleg', '')
                    target = _find_para_by_quote(new_paras, beleg) if beleg else None
                    start_p = target or first_para
                    end_p   = target or last_para
                    cid = comment_id[0]
                    comment_id[0] += 1
                    _add_comment_range_start(start_p, cid)
                    _add_comment_range_end(end_p, cid)
                    label = f" – {unit['label']}" if unit.get('label') else ""
                    comment_list.append((cid, f"Frage {unit['num']}{label}"))

        if orig_body.strip():
            process_markdown_to_docx(doc, orig_body, hide_text=True, base_path=base_path)

    # --- Fragentext laden (optional) ---
    # Mehrzeilige Fragen (mit aufgelisteten Unterpunkten) komplett erfassen: Folgezeilen
    # bis zur nächsten nummerierten Frage bzw. Leerzeile gehören zur aktuellen Frage.
    questions_map: dict = {}
    if questions_path and os.path.exists(questions_path):
        cur_q = None
        with open(questions_path, encoding='utf-8') as qf:
            for raw in qf:
                stripped = raw.strip()
                qm = re.match(r'^(\d+)\.\s+(.+)', stripped)
                if qm:
                    cur_q = int(qm.group(1))
                    questions_map[cur_q] = qm.group(2).strip()
                elif cur_q is not None:
                    if stripped:
                        questions_map[cur_q] += '\n' + stripped
                    else:
                        cur_q = None  # Leerzeile beendet die Frage

    # --- QA-Unterkapitel am Dokumentende ---
    if has_qa and qa_items:
        if parent_chapter:
            # Nach Fix 4 werden Level-1-Headings in nicht-nummerierten Dokumenten übersprungen
            # (OCR-Artefakte), sodass counters[0] = 0 bleibt. Die Top-Level-Inhalte liegen
            # bei Level 2 (counters[1]). Für nummerierte Dokumente zählt weiterhin counters[0].
            if structural_renumber:
                # Strukturelle Neu-Nummerierung: Top-Level-Sektionen zählen counters[0];
                # Lernfragen ist das nächste Geschwister auf Tiefe 1.
                qa_top_num = counters[0] + 1
            elif extracted_chapter and has_numbered_chapters:
                # Rebase-Modus: Die direkten Kinder unter parent_chapter tragen rebasierte
                # Nummern (3.2.3.1 … 3.2.3.8). Lernfragen ist das nächste freie Geschwister
                # (→ 3.2.3.9). _max_child trackt das höchste vergebene Kind-Segment; Fallback
                # _auto_counter für rein auto-nummerierte Unterkapitel ohne Eigennummer.
                qa_top_num = max(_max_child, _auto_counter) + 1
            elif not has_numbered_chapters:
                qa_top_num = counters[1] + 1
            else:
                qa_top_num = counters[0] + 1
            qa_hdg_text = f"{parent_chapter}.{qa_top_num} Lernfragen"
            qa_level    = min(lvl_shift + 1, 9)
            q_sub_level = min(lvl_shift + 2, 9)
        else:
            qa_hdg_text = "Lernfragen"
            qa_level    = 1
            q_sub_level = 2

        h = doc.add_heading(qa_hdg_text, level=qa_level)
        _set_heading_color(h, MM_HEADING_COLORS.get(qa_level, MM_HEADING_COLORS[9]))

        for q_idx, item in enumerate(qa_items, start=1):
            if parent_chapter:
                fq_hdg = f"{parent_chapter}.{qa_top_num}.{q_idx} Frage {item['num']}"
            else:
                fq_hdg = f"Frage {item['num']}"
            h_f = doc.add_heading(fq_hdg, level=q_sub_level)
            _set_heading_color(h_f, MM_HEADING_COLORS.get(q_sub_level, MM_HEADING_COLORS[9]))

            # Fragetext anzeigen wenn verfügbar (mehrzeilig inkl. Unterpunkte)
            # Fragetext bevorzugt aus der (nummerierten) Fragendatei; sonst den vom
            # Fallback-Parser mitgeführten Header-Text der Frage nutzen.
            q_text = questions_map.get(item['num']) or item.get('frage_text')
            if q_text:
                q_lines = q_text.split('\n')
                p_q = doc.add_paragraph()
                for li, ql in enumerate(q_lines):
                    r_q = p_q.add_run(ql)
                    r_q.bold = True
                    r_q.italic = True
                    if li < len(q_lines) - 1:
                        r_q.add_break()

            def _qa_meta(src: dict):
                quelle = _remap_textgrundlage(src.get('textgrundlage', '–'), tg_renumber)
                meta = doc.add_paragraph()
                r = meta.add_run(
                    f"Quelle: {quelle}  |  "
                    f"Schlüsselbegriffe: {src.get('schluessel', '–')}  |  "
                    f"Abdeckung: {src.get('abdeckung', '–')}"
                )
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

            if item.get('subs'):
                # Mehrteilige Frage: jeden Unterpunkt mit eigener Überschrift, Antwort und Meta.
                for sub in item['subs']:
                    p_lbl = doc.add_paragraph()
                    p_lbl.add_run(sub.get('label', '')).bold = True
                    doc.add_paragraph(sub.get('antwort', ''))
                    _qa_meta(sub)
            else:
                doc.add_paragraph(item['antwort'])
                _qa_meta(item)

    # --- Kommentare einbetten und speichern ---
    if comment_list:
        _inject_comments_part(doc, comment_list)
    doc.save(output_path)
    print(f"Word-Dokument erstellt ({len(comment_list)} Kommentare).")


def build_condensed_word_document(condensed_text: str, output_path: str, base_path: str = None,
                                  parent_chapter: str = None, parent_level: int = None,
                                  doc_title: str = "Vertiefende Literatur") -> None:
    """
    Verdichteter Modus (--condensed): Baut ein Word-Dokument mit einer Elternüberschrift
    (Sammelband-Titel) und je einem Unterkapitel pro Artikel. Sichtbar ist nur die
    Kurzübersicht; ausgeblendet sind ausschließlich die 2-3 vom Modell gewählten Zitate
    (NICHT der vollständige Originaltext - anders als build_interleaved_word_document).
    Kein QA-Overlay, keine Vollständigkeitsprüfung, keine strukturelle Neu-Nummerierung -
    die Dokumentstruktur ist bewusst einfach: Elternkapitel → Artikel 1..N.
    """
    print(f"--- Erstelle verdichtetes Word-Dokument -> {output_path} ---")

    condensed_text = normalize_heading_levels(condensed_text)
    condensed_text = _strip_selbstpruefung(condensed_text)
    sections = [s for s in parse_sections(condensed_text) if s['heading'] != '__preamble__']

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    if parent_chapter:
        p_level = min(parent_level or _parent_level_from_chapter(parent_chapter), 9)
        h = doc.add_heading(level=p_level)
        add_formatted_text(h, f"{parent_chapter} {doc_title}",
                           default_color=MM_HEADING_COLORS.get(p_level, MM_HEADING_COLORS[9]))
        child_level = min(p_level + 1, 9)
    else:
        h = doc.add_heading(doc_title, level=1)
        _set_heading_color(h, MM_HEADING_COLORS[1])
        child_level = 2

    for i, sec in enumerate(sections, 1):
        heading_text = _clean_heading_text(sec['heading'])
        if parent_chapter:
            heading_text = f"{parent_chapter}.{i} {heading_text}"
        h = doc.add_heading(level=child_level)
        add_formatted_text(h, heading_text,
                           default_color=MM_HEADING_COLORS.get(child_level, MM_HEADING_COLORS[9]))

        summary_body, quotes_body = _extract_quotes_block(sec['body'])
        if summary_body:
            process_markdown_to_docx(doc, summary_body, hide_text=False,
                                     base_path=base_path, headings_as_bold=True)
        if quotes_body:
            process_markdown_to_docx(doc, quotes_body, hide_text=True,
                                     base_path=base_path, headings_as_bold=True)

    doc.save(output_path)
    print(f"Word-Dokument erstellt ({len(sections)} Artikel, verdichteter Modus).")


# ---------------------------------------------------------------------------
# Leitfragen-Quiz-Prüfung (bestehende .docx-Lernunterlage bearbeiten)
# ---------------------------------------------------------------------------
#
# Neuer Verarbeitungspfad, getrennt vom PDF→Lernmittel-Flow: Eingabe ist eine
# fertige .docx-Lernunterlage mit einem Kapitel "Leitfragen-Quiz", dessen Fragen
# nur als Screenshots (Bilder) vorliegen. Das Dokument wird ausschließlich
# ERGÄNZT (nie umgeschrieben/gelöscht):
#   1. Fragebilder → Text (Gemini-Vision)
#   2. Prüfung gegen den restlichen Dokumenttext (bestehendes QA-Feature)
#   3. Antwort + Abdeckungsgrad unter jede Frage in Kapitel 10 einfügen
#   4. Kommentar-Marker ("Quizfrage N") an den relevanten Quellkapiteln setzen
#      bzw. bestehende Kommentare erweitern
# Ausgabe in eine neue Kopie; das Original bleibt unangetastet.
# Das Quiz-Kapitel wird IMMER explizit über --quiz-chapter benannt (kein Standard):
# ein fester Default wäre falsch, sobald sich Nummer/Wortlaut ändert.


def _iter_paragraph_image_blobs(doc, paragraph):
    """Liefert (blob, mime) für jedes eingebettete Bild (w:blip r:embed) eines Absatzes."""
    out = []
    for blip in paragraph._p.findall('.//' + qn('a:blip')):
        rId = blip.get(qn('r:embed'))
        if not rId:
            continue
        part = doc.part.related_parts.get(rId)
        if part is not None:
            out.append((part.blob, part.content_type))
    return out


def extract_quiz_questions(doc, quiz_chapter_heading: str):
    """Findet das Quiz-Kapitel (Heading 2) und sammelt pro 'Frage N' (Heading 3) die
    folgenden Bild-Absätze. Gibt (questions, quiz_range) zurück:
      questions: [{'num', 'heading_para', 'anchor', 'images': [(blob, mime), ...]}]
      quiz_range: (lo, hi) Absatz-Indizes des Kapitels (für KB-Ausschluss)
    'anchor' ist der Absatz DIREKT nach dem Bildblock der Frage (Einfügepunkt der Antwort).
    """
    paras = doc.paragraphs

    def heading_level(p):
        """Heading-Ebene (1..9) oder None, wenn der Absatz keine Überschrift ist."""
        st = p.style.name if p.style else ""
        m = re.match(r'Heading (\d)', st)
        return int(m.group(1)) if m else None

    def norm(text):
        """Zeichensetzungs-tolerante Normalisierung (entfernt Anführungszeichen/Markdown,
        kollabiert Whitespace, lowercase) – siehe _normalize_quote."""
        return _normalize_quote(text)

    def core(text):
        """Wie norm(), zusätzlich ohne führende Kapitelnummer ('10 ', '3.4.1 '),
        damit der Abgleich unabhängig von Nummerierung UND Zeichensetzung ist."""
        return re.sub(r'^\s*\d+(?:\.\d+)*\.?\s*', '', norm(text)).strip()

    target_full = norm(quiz_chapter_heading)
    target_core = core(quiz_chapter_heading)

    # 1) Exakter (nummern-/zeichensetzungs-toleranter) Treffer, beliebige Ebene
    start = start_level = None
    for i, p in enumerate(paras):
        lvl = heading_level(p)
        if lvl is None:
            continue
        if norm(p.text) == target_full or (target_core and core(p.text) == target_core):
            start, start_level = i, lvl
            break
    # 2) Fallback: Teiltext-Treffer ('enthält')
    if start is None and target_core:
        for i, p in enumerate(paras):
            lvl = heading_level(p)
            if lvl is None:
                continue
            if target_core in norm(p.text):
                start, start_level = i, lvl
                break
    if start is None:
        raise ValueError(
            f"Quiz-Kapitel '{quiz_chapter_heading}' nicht gefunden. "
            "Tipp: den stabilen Textteil ohne Kapitelnummer angeben, z.B. "
            "--quiz-chapter \"Leitfragen-Quiz zum Kurs Personalauswahl\"."
        )

    # Kapitelende: nächste Überschrift gleicher oder höherer Ebene (<= start_level)
    hi = len(paras)
    for i in range(start + 1, len(paras)):
        lvl = heading_level(paras[i])
        if lvl is not None and lvl <= start_level:
            hi = i
            break

    # Fragen: 'Frage N'-Überschriften (beliebige Ebene) + folgende Bild-Absätze
    questions = []
    frage_re = re.compile(r'^Frage\s+(\d+)\b')

    def is_frage_heading(p):
        return heading_level(p) is not None and frage_re.match(p.text.strip())

    i = start + 1
    while i < hi:
        m = frage_re.match(paras[i].text.strip())
        if heading_level(paras[i]) is not None and m:
            num = int(m.group(1))
            images = []
            j = i + 1
            while j < hi and not is_frage_heading(paras[j]):
                images.extend(_iter_paragraph_image_blobs(doc, paras[j]))
                j += 1
            anchor = paras[j] if j < len(paras) else None
            questions.append({'num': num, 'heading_para': paras[i], 'anchor': anchor, 'images': images})
            i = j
        else:
            i += 1
    if not questions:
        raise ValueError(
            f"Im Kapitel '{paras[start].text.strip()}' wurden keine 'Frage N'-Überschriften "
            "gefunden. Ist es wirklich das Quiz-Kapitel?"
        )
    return questions, (start, hi)


def _ocr_quiz_image(images) -> str:
    """Transkribiert eine Quiz-Frage aus ihren Screenshot(s) via Gemini-Vision."""
    supported = {'image/png', 'image/jpeg', 'image/webp', 'image/heic', 'image/heif'}
    parts = []
    for blob, mime in images:
        if mime in supported:
            parts.append(types.Part.from_bytes(data=blob, mime_type=mime))
    if not parts:
        return ""
    prompt = (
        "Dies ist ein Screenshot einer einzelnen Quiz-/Prüfungsfrage aus einer "
        "Lernunterlage. Transkribiere den vollständigen Fragetext wörtlich, inklusive "
        "aller Antwortoptionen (z.B. a), b), c)) falls vorhanden. Gib NUR den reinen "
        "Fragetext aus – keine Einleitung, keine Erklärung, keine Bildbeschreibung."
    )
    parts.append(prompt)
    resp = call_gemini_with_retry(
        model_name='gemini-2.5-pro',
        contents=parts,
        config=types.GenerateContentConfig(temperature=0.0),
    )
    return (resp.text or "").strip()


def ocr_quiz_questions(questions, cache_path: Path, force: bool = False) -> dict:
    """Vision-OCR aller Fragebilder mit JSON-Cache. Gibt {num: fragetext} zurück."""
    cache = {}
    if cache_path.exists() and not force:
        cache = {int(k): v for k, v in json.loads(cache_path.read_text(encoding="utf-8")).items()}
    updated = False
    for q in questions:
        num = q['num']
        if num in cache and cache[num].strip():
            continue
        print(f"   [Vision] Frage {num} ({len(q['images'])} Bild(er)) …")
        cache[num] = _ocr_quiz_image(q['images']) or "(Bild nicht lesbar)"
        updated = True
        cache_path.parent.mkdir(parents=True, exist_ok=True)
        cache_path.write_text(json.dumps(cache, ensure_ascii=False, indent=2), encoding="utf-8")
    if updated:
        print(f"   Vision-OCR abgeschlossen: {len(cache)} Fragen → {cache_path}")
    else:
        print(f"[SKIP] Vision-OCR – bereits vorhanden: {cache_path}")
    return cache


def _heading_match_keys(text: str) -> set:
    """Präzise Matching-Schlüssel einer Überschrift/Textgrundlage: volle Überschrift,
    führende Kapitelnummer (z.B. '3.4.1') und der nummernfreie Titel. Keine bloßen
    Einzelziffern (Mehrdeutigkeit)."""
    t = _clean_heading_text(text or '').replace('[', '').replace(']', '')
    t = re.split(r'[,;•]', t)[0].strip()
    keys = set()
    full = normalize_heading(t)
    if full:
        keys.add(full)
    numm = re.match(r'^(\d+(?:\.\d+)*)\b', t)
    if numm:
        keys.add(numm.group(1))
    title = normalize_heading(re.sub(r'^\d+(?:\.\d+)*\.?\s*', '', t))
    if len(title) > 2:
        keys.add(title)
    return {k for k in keys if k}


def _enclosing_h1_range(doc, para_idx: int):
    """Absatz-Bereich (start, end) der Heading-1-Kurseinheit, die para_idx enthält.
    Fallback: gesamtes Dokument, falls keine H1 vorhanden."""
    paras = doc.paragraphs
    h1s = [i for i, p in enumerate(paras) if p.style and p.style.name == "Heading 1"]
    start, end = 0, len(paras)
    for i in h1s:
        if i <= para_idx:
            start = i
        elif i > para_idx:
            end = i
            break
    return start, end


def build_quiz_knowledge_base(doc, include_range, exclude_range):
    """Baut aus dem Dokument-Bereich `include_range` (ohne das Quiz-Kapitel `exclude_range`)
    eine gegliederte Markdown-Wissensbasis und eine Map {match_key: Absatz} zur
    Marker-Verankerung."""
    inc_lo, inc_hi = include_range
    exc_lo, exc_hi = exclude_range
    lines = []
    heading_map = {}
    for i, p in enumerate(doc.paragraphs):
        if not (inc_lo <= i < inc_hi):
            continue
        if exc_lo <= i < exc_hi:
            continue
        txt = p.text.strip()
        if not txt:
            continue
        st = p.style.name if p.style else ""
        m = re.match(r'Heading (\d)', st)
        if m:
            lvl = int(m.group(1))
            lines.append('#' * min(lvl, 6) + ' ' + txt)
            for key in _heading_match_keys(txt):
                heading_map.setdefault(key, p)
        else:
            lines.append(txt)
    return '\n\n'.join(lines), heading_map


def _resolve_tg_para(textgrundlage: str, heading_map: dict):
    """Löst eine QA-Textgrundlage auf den zugehörigen Überschrift-Absatz auf (oder None)."""
    if not textgrundlage or textgrundlage.strip() in ('–', '-', ''):
        return None
    keys = sorted(_heading_match_keys(textgrundlage),
                  key=lambda k: (0 if re.search(r'\d', k) else 1, -len(k)))
    for key in keys:
        if key in heading_map:
            return heading_map[key]
    return None


def _existing_comment_id(paragraph):
    """Kommentar-ID, falls der Absatz bereits in einem Kommentarbereich liegt, sonst None."""
    el = paragraph._p.find(qn('w:commentRangeStart'))
    if el is None:
        el = paragraph._p.find('.//' + qn('w:commentReference'))
    if el is not None:
        try:
            return int(el.get(qn('w:id')))
        except (TypeError, ValueError):
            return None
    return None


def _render_multiline_answer(paragraph, text):
    """Rendert eine (ggf. mehrzeilige, strukturierte) Antwort in einen Absatz:
    Zeilenumbrüche via Break, ganze Zeilen in **…** als fettes Label (z.B. Kategorie/
    Lücke bei Zuordnungs-/Reihenfolge-/Lückenfragen)."""
    lines = (text or "–").split('\n')
    for i, ln in enumerate(lines):
        if i > 0:
            paragraph.add_run().add_break()
        s = ln.strip()
        if not s:
            continue
        m = re.fullmatch(r'\*\*\s*(.+?)\s*\*\*', s)
        if m:
            paragraph.add_run(m.group(1)).bold = True
        else:
            paragraph.add_run(s)


_QA_META_RE = re.compile(r'^Quelle:.*\bAbdeckung:')

# Eindeutiger Absatzstil, mit dem NUR das Tool seine eingefügten Antwort-Absätze markiert.
# Wichtig: manche Nutzer schreiben ihre EIGENEN Antworten selbst als "Antwort: …" – diese
# dürfen bei einem Re-Lauf NIEMALS entfernt werden. Deshalb erkennt die Idempotenz-Logik
# Tool-Absätze am Stil (unzweideutig), nicht am Textpräfix.
_QUIZ_ANSWER_STYLE = "Quiz Antwort"


def _ensure_quiz_answer_style(doc):
    """Stellt den (auf 'Normal' basierenden, optisch unauffälligen) Markierungsstil bereit."""
    from docx.enum.style import WD_STYLE_TYPE
    styles = doc.styles
    if any(s.name == _QUIZ_ANSWER_STYLE for s in styles):
        return _QUIZ_ANSWER_STYLE
    st = styles.add_style(_QUIZ_ANSWER_STYLE, WD_STYLE_TYPE.PARAGRAPH)
    try:
        st.base_style = styles['Normal']
    except Exception:
        pass
    try:
        st.hidden = True
        st.quick_style = False
    except Exception:
        pass
    return _QUIZ_ANSWER_STYLE


def _strip_previous_quiz_answers(doc, quiz_range):
    """Entfernt frühere, VOM TOOL eingefügte Antwort-/Meta-Absätze im Quiz-Kapitel und macht
    den Lauf idempotent. Erkennung primär am eindeutigen Stil `_QUIZ_ANSWER_STYLE`; als
    Fallback für alte, ungetaggte Einfügungen eine Meta-Zeile ('Quelle: … | Abdeckung:',
    tool-spezifisch) samt der unmittelbar davor stehenden 'Antwort:'-Zeile. Dadurch werden
    vom NUTZER selbst geschriebene 'Antwort:'-Zeilen (ohne folgende Meta-Zeile) NIE entfernt.
    Auf einer sauberen Basis ein No-op; Fragen/Bilder/sonstiger Inhalt bleiben unberührt."""
    lo, hi = quiz_range
    block = doc.paragraphs[lo:hi]
    to_remove = []
    seen = set()

    def mark(p):
        if id(p._p) not in seen:
            seen.add(id(p._p))
            to_remove.append(p)

    for i, p in enumerate(block):
        styled = p.style is not None and p.style.name == _QUIZ_ANSWER_STYLE
        if styled:
            mark(p)
        if _QA_META_RE.match(p.text.strip()):
            mark(p)
            if i > 0 and block[i - 1].text.strip().startswith("Antwort:"):
                mark(block[i - 1])

    for p in to_remove:
        parent = p._p.getparent()
        if parent is not None:
            parent.remove(p._p)
    return len(to_remove)


def _insert_qa_answer_before(anchor, item):
    """Fügt Antwort- und Meta-Absätze VOR dem Anker-Absatz ein (nur Insert, nichts löschen).
    Alle eingefügten Absätze erhalten den Stil `_QUIZ_ANSWER_STYLE`, damit ein Re-Lauf sie
    eindeutig (ohne Textheuristik) wiederfinden und ersetzen kann."""
    def _tag(p):
        try:
            p.style = _QUIZ_ANSWER_STYLE
        except Exception:
            pass
        return p

    def emit(antwort, quelle, schluessel, abdeckung, label=None):
        if label:
            pl = _tag(anchor.insert_paragraph_before())
            pl.add_run(label).bold = True
        pa = _tag(anchor.insert_paragraph_before())
        pa.add_run("Antwort: ").bold = True
        # Mehrzeilige, strukturierte Antworten (Zuordnung/Reihenfolge/Lücken) auf eigener
        # Zeile beginnen; einzeilige Antworten bleiben inline hinter "Antwort: ".
        if antwort and '\n' in antwort:
            pa.add_run().add_break()
        _render_multiline_answer(pa, antwort or "–")
        pm = _tag(anchor.insert_paragraph_before())
        r = pm.add_run(
            f"Quelle: {quelle or '–'}  |  "
            f"Schlüsselbegriffe: {schluessel or '–'}  |  "
            f"Abdeckung: {abdeckung or '–'}"
        )
        r.font.size = Pt(9)
        low = (abdeckung or '').lower()
        if 'nicht enthalten' in low or 'teilweise' in low:
            r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        else:
            r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    if item.get('subs'):
        for sub in item['subs']:
            emit(sub.get('antwort'), sub.get('textgrundlage'),
                 sub.get('schluessel'), sub.get('abdeckung'), label=sub.get('label'))
    else:
        emit(item.get('antwort'), item.get('textgrundlage'),
             item.get('schluessel'), item.get('abdeckung'))


def _quiz_short_name(heading: str) -> str:
    """Kompakter Quizname für Kommentar-Zeilen, z.B. „Quiz Kapitel 1" oder „Leitfragen-Lernquiz".
    Schneidet Untertitel (' - …'), ' zum Kurs …' und eine führende Kapitelnummer ab."""
    h = (heading or '').strip().strip('"„“”»«')
    h = re.split(r'\s+[-–—]\s+', h)[0].strip()
    h = re.split(r'\s+zum\s+Kurs\b', h, flags=re.IGNORECASE)[0].strip()
    h = re.sub(r'^\d+(?:\.\d+)*\.?\s+', '', h).strip()
    return h or (heading or '').strip()


def _remove_quiz_comments(doc) -> int:
    """Entfernt ALLE vom Tool (Autor „Quiz") gesetzten Kommentare samt ihrer Range-Marker
    (`w:commentRangeStart/End` + Referenz-Run) und lässt Nutzer-/Lernfragen-Kommentare
    unberührt. Macht das Neusetzen der Marker idempotent (kein Mix aus alt und neu)."""
    try:
        comments_elm = doc.comments._comments_elm
    except Exception:
        return 0
    quiz_ids = set()
    for c_elm in list(comments_elm.comment_lst):
        if (c_elm.author or "") == "Quiz":
            quiz_ids.add(str(c_elm.id))
            c_elm.getparent().remove(c_elm)
    if not quiz_ids:
        return 0
    body = doc.element.body
    for tag in ('w:commentRangeStart', 'w:commentRangeEnd'):
        for el in list(body.iter(qn(tag))):
            if el.get(qn('w:id')) in quiz_ids:
                parent = el.getparent()
                if parent is not None:
                    parent.remove(el)
    # Referenz-Runs (der Run, der w:commentReference enthält) komplett entfernen
    for ref in list(body.iter(qn('w:commentReference'))):
        if ref.get(qn('w:id')) in quiz_ids:
            run = ref.getparent()
            if run is not None and run.getparent() is not None:
                run.getparent().remove(run)
    return len(quiz_ids)


def _apply_quiz_markers(doc, analyses) -> int:
    """Setzt Kommentar-Marker an den KONKRETEN Textstellen (Beleg-verankert), NIEMALS innerhalb
    eines Quiz-Kapitels. Kommt eine Stelle in mehreren Quizfragen vor, sammeln sich die Quellen
    in EINEM „Quiz"-Kommentar (Mehrfach-Vorkommen = interessantes Signal). Bestehende Nutzer-/
    Lernfragen-Kommentare bleiben unberührt (separater Quiz-Kommentar koexistiert).
    MUSS VOR der Mutations-Phase laufen (pristine Absatz-Indizes)."""
    paras = doc.paragraphs
    n = len(paras)

    # Vereinigung ALLER Quiz-Kapitel-Bereiche → dort wird nie markiert (auch nicht auf
    # Quizfragen eines anderen Quiz).
    excluded = set()
    for a in analyses:
        lo, hi = a['quiz_range']
        excluded.update(range(lo, hi))

    def _markable(i):
        if i in excluded:
            return False
        p = paras[i]
        return not (p.style is not None and p.style.name == _QUIZ_ANSWER_STYLE)

    markable_idx = [i for i in range(n) if _markable(i)]
    # Normalisierten Absatztext EINMAL vorberechnen (statt je Frage erneut) – sonst wird die
    # Beleg-Suche über zehntausende Absätze × hunderte Fragen extrem langsam.
    entries = [(i, paras[i], _normalize_quote(paras[i].text)) for i in markable_idx]

    def _find_beleg(cands, beleg):
        """Zielabsatz per Beleg-Zitat aus vorbereiteten (idx, para, norm)-Einträgen."""
        nq = _normalize_quote(beleg)
        if len(nq) < 8:
            return None
        needle = nq[:60]
        for _i, p, nt in cands:
            if nt and needle in nt:
                return p
        return None

    # Heading-Map NUR aus markierbarem Korpus (für den Überschrift-Fallback).
    heading_map = {}
    for i, p, _nt in entries:
        st = p.style.name if p.style else ""
        if re.match(r'Heading (\d)', st):
            for key in _heading_match_keys(p.text):
                heading_map.setdefault(key, p)

    groups = {}   # id(p._p) -> {'para': p, 'refs': [str, ...]}
    order = []

    def _add(target, ref):
        k = id(target._p)
        if k not in groups:
            groups[k] = {'para': target, 'refs': []}
            order.append(k)
        if ref not in groups[k]['refs']:
            groups[k]['refs'].append(ref)

    beleg_hits = heading_fallback = unresolved = 0
    for a in analyses:
        qname = _quiz_short_name(a['heading'])
        lo, hi = a['kb_scope']
        section = [e for e in entries if lo <= e[0] < hi]
        for num in sorted(a['items_by_num']):
            item = a['items_by_num'][num]
            units = item['subs'] if item.get('subs') else [item]
            for u in units:
                beleg = (u.get('beleg') or '').strip()
                target = _find_beleg(section, beleg) if beleg else None
                if target is None and beleg:
                    target = _find_beleg(entries, beleg)
                via_beleg = target is not None
                if target is None:
                    target = _resolve_tg_para(u.get('textgrundlage', ''), heading_map)
                if target is None:
                    unresolved += 1
                    continue
                if via_beleg:
                    beleg_hits += 1
                else:
                    heading_fallback += 1
                label = u.get('label')
                ref = f"»{qname}« – Frage {num}" + (f" ({label})" if label else "")
                _add(target, ref)

    multi = 0
    for k in order:
        para = groups[k]['para']
        refs = groups[k]['refs']
        if len(refs) > 1:
            multi += 1
        runs = para.runs or [para.add_run("")]
        if len(refs) == 1:
            body = f"In einem Quiz behandelt: {refs[0]}"
        else:
            body = "In Quizzes behandelt:\n" + "\n".join(f"• {r}" for r in refs)
        doc.add_comment(runs, text=body, author="Quiz", initials="QZ")

    print(f"   Marker (Beleg-verankert): {len(order)} Textstellen markiert "
          f"({beleg_hits} per Beleg, {heading_fallback} per Überschrift-Fallback, "
          f"{multi} mehrfach belegt, {unresolved} ohne Fundstelle).")
    return len(order)


def _retry_single_question(full_kb: str, num: int, question_text: str) -> str:
    """Beantwortet GENAU EINE Quizfrage gegen eine (große) Wissensbasis – strikt, damit das
    Modell nicht die dokumenteigenen Folien-/Übungsfragen aufzählt. Gibt Roh-QA-Text zurück."""
    prompt = (
        "Rolle:\nDu bist Prüfer für Wirtschaftspsychologie.\n\n"
        "Aufgabe:\nDu erhältst eine Wissensbasis (ein Lerndokument) und GENAU EINE Quizfrage.\n"
        "Beantworte AUSSCHLIESSLICH diese eine Frage anhand der Wissensbasis.\n"
        "- Erfinde KEINE weiteren Fragen und beantworte NICHTS anderes aus dem Dokument.\n"
        "- Wenn die Antwort nicht eindeutig in der Wissensbasis steht, setze Abdeckung auf "
        "'nicht enthalten' und rate NICHT.\n"
        "- Bei Zuordnungs-/Reihenfolge-/Lückenfragen die konkrete Lösung strukturiert angeben "
        "(pro Zuordnung/Lücke/Schritt eine eigene Zeile; Bezeichner als **fett**).\n\n"
        "Gib GENAU EINEN Block in diesem Format aus (keine Einleitung, nichts davor/danach):\n"
        f"Frage {num}\n"
        "Fragetext: [die Frage wörtlich]\n"
        "Antwort: [Lösung; bei mehreren Teilen mehrzeilig]\n"
        "Textgrundlage: [genaues Unterkapitel]\n"
        "Schlüsselbegriffe: [1-3]\n"
        "Beleg: [wörtliches Zitat, 5-15 Wörter]\n"
        "Abdeckung: vollständig / teilweise / nicht enthalten\n\n"
        f"Wissensbasis:\n{full_kb}\n\n"
        f"Die EINE zu beantwortende Frage:\n{num}. {question_text}"
    )
    resp = call_gemini_with_retry(
        model_name='gemini-2.5-pro',
        contents=prompt,
        config=types.GenerateContentConfig(temperature=0.1),
    )
    return resp.text or ""


def _abdeckung_missing(it: dict) -> bool:
    """True, wenn ein QA-Item (oder eines seiner Unterpunkte) 'nicht enthalten' meldet."""
    levels = [it.get('abdeckung', '')]
    levels += [s.get('abdeckung', '') for s in it.get('subs', [])]
    return any('nicht enthalten' in (a or '').lower() for a in levels)


def _analyze_one_quiz(doc, src: Path, quiz_chapter_heading: str, force: bool) -> dict:
    """Analyse-Phase für EIN Quiz-Kapitel – rein lesend, KEINE Mutation des Dokuments:
    Fragen extrahieren, Vision-OCR, Wissensbasis (umschließende Heading-1-Kurseinheit),
    Vollständigkeitsprüfung und Whole-Doc-Retry. Gibt alle Ergebnisse für die spätere
    Mutations-Phase zurück (questions/items_by_num/heading_map/quiz_range)."""
    tag = quiz_chapter_heading[:50]
    # Cache-Ordner MUSS auch vom Quiz-Kapitel abhängen: sonst wiederverwendet ein Lauf
    # für ein anderes Quiz-Kapitel (gleiche .docx, andere Überschrift) fälschlich den
    # Cache eines vorherigen, andersartigen Quiz-Laufs (unterschiedliche Fragenanzahl!).
    chapter_slug = re.sub(r'[^A-Za-z0-9]+', '_', quiz_chapter_heading).strip('_')[:60]
    work_dir = src.parent / f".quizcheck_{src.stem}_{chapter_slug}"
    work_dir.mkdir(parents=True, exist_ok=True)
    if force:
        # load_or_run/OCR-Cache honorieren --force nur teilweise; hier alle Zwischen-
        # ergebnisse gezielt entfernen, damit --force wirklich neu berechnet.
        for f in ("quiz_fragen.json", "quiz_fragen.txt", "quiz_wissensbasis.md",
                  "quiz_qa_ergebnis.md"):
            (work_dir / f).unlink(missing_ok=True)
        for f in work_dir.glob("quiz_retry_*.md"):
            f.unlink(missing_ok=True)
        for f in work_dir.glob("quiz_single_*.md"):
            f.unlink(missing_ok=True)

    # Schritt 1: Fragen + Bilder extrahieren
    questions, quiz_range = extract_quiz_questions(doc, quiz_chapter_heading)
    print(f"   [{tag}] {len(questions)} Quizfragen (Absätze {quiz_range[0]}–{quiz_range[1]}).")

    # Schritt 2: Vision-OCR (Bilder → Fragetext), gecacht
    ocr_map = ocr_quiz_questions(questions, work_dir / "quiz_fragen.json", force=force)
    questions_txt_path = work_dir / "quiz_fragen.txt"
    q_lines = []
    for q in sorted(questions, key=lambda x: x['num']):
        text = ' '.join(ocr_map.get(q['num'], '').split())
        q_lines.append(f"{q['num']}. {text}")
    questions_txt_path.write_text('\n'.join(q_lines), encoding="utf-8")

    # Schritt 3: Wissensbasis auf die umschließende Kurseinheit (Heading 1) begrenzen
    # (thematisch passend: ein Quiz wird gegen die eigene Kurseinheit geprüft, nicht gegen
    # fremde Kurseinheiten) und das Quiz-Kapitel selbst ausschließen.
    kb_path = work_dir / "quiz_wissensbasis.md"
    kb_scope = _enclosing_h1_range(doc, quiz_range[0])
    kb, heading_map = build_quiz_knowledge_base(doc, kb_scope, quiz_range)
    kb_path.write_text(kb, encoding="utf-8")
    print(f"   [{tag}] Wissensbasis (Absätze {kb_scope[0]}–{kb_scope[1]}): {len(kb)} Zeichen, "
          f"{len(heading_map)} Überschrift-Schlüssel.")

    # Schritt 4: Vollständigkeitsprüfung (bestehendes QA-Feature wiederverwenden)
    qa_path = work_dir / "quiz_qa_ergebnis.md"
    qa_text = load_or_run(
        qa_path,
        lambda: verify_with_questions(kb, str(questions_txt_path)),
        f"Quiz-Vollständigkeitsprüfung [{tag}]",
    )
    raw_items = parse_qa_response(qa_text)
    valid_nums = {q['num'] for q in questions}
    # Auf ECHTE Fragenummern klemmen: bei großer Wissensbasis und wenigen Fragen generiert das
    # Modell mitunter zusätzliche Blöcke aus KB-Inhalten (halluzinierte Fragen). Diese dürfen
    # weder als Antwort eingefügt noch als Marker gesetzt werden.
    items = [it for it in raw_items if it['num'] in valid_nums]
    items_by_num = {it['num']: it for it in items}
    print(f"   [{tag}] QA geparst: {len(raw_items)} Blöcke, {len(items_by_num)}/{len(questions)} gültig.")

    # Schritt 4a: Ausrichtungs-/Über-Generierungs-Check. Lieferte der Batch deutlich mehr Blöcke
    # als Fragen (Modell abgedriftet) oder fehlen Fragen, sind die Batch-Antworten unzuverlässig
    # (Fehlausrichtung: 'Frage k' ≠ k-te Eingabefrage). Dann betroffene Fragen EINZELN und STRIKT
    # gegen die Kapitel-Wissensbasis neu beantworten (bewährter Anti-Halluzinations-Pfad).
    overgenerated = len(raw_items) > len(questions) + max(2, len(questions) // 5)
    missing = valid_nums - set(items_by_num)
    if overgenerated or missing:
        targets = sorted(valid_nums) if overgenerated else sorted(missing)
        reason = (f"Über-Generierung (parsed={len(raw_items)} ≫ Fragen={len(questions)})"
                  if overgenerated else f"{len(missing)} fehlende Antwort(en)")
        print(f"   [{tag}] Batch-QA unzuverlässig – {reason}; {len(targets)} Frage(n) einzeln neu.")
        for num in targets:
            qtext = ' '.join(ocr_map.get(num, '').split())
            spath = work_dir / f"quiz_single_{num}.md"
            stext = load_or_run(
                spath,
                lambda: _retry_single_question(kb, num, qtext),
                f"Einzel-QA Frage {num} [{tag}]",
            )
            parsed = parse_qa_response(stext)
            if len(parsed) == 1:
                it = parsed[0]
                it['num'] = num
                items_by_num[num] = it

    # Schritt 4b: Whole-Doc-Retry für 'nicht enthalten' — gegen das GESAMTE Dokument
    # (alle Kurseinheiten) statt nur der umschließenden KE, falls die Antwort außerhalb liegt.
    retry_nums = sorted(num for num, it in items_by_num.items() if _abdeckung_missing(it))
    if retry_nums:
        print(f"   [{tag}] Whole-Doc-Retry für {len(retry_nums)} 'nicht enthalten'-Frage(n): {retry_nums}")
        full_kb, _ = build_quiz_knowledge_base(doc, (0, len(doc.paragraphs)), quiz_range)
        updated = 0
        for num in retry_nums:
            qtext = ' '.join(ocr_map.get(num, '').split())
            rpath = work_dir / f"quiz_retry_{num}.md"
            rtext = load_or_run(
                rpath,
                lambda: _retry_single_question(full_kb, num, qtext),
                f"Quiz-Retry Frage {num} [{tag}]",
            )
            parsed = parse_qa_response(rtext)
            # Guard: genau EIN valider Block, sonst ist die Antwort unzuverlässig (Modell hat
            # z.B. dokumenteigene Fragen aufgezählt) → Original ('nicht enthalten') behalten.
            if len(parsed) == 1 and not _abdeckung_missing(parsed[0]):
                it = parsed[0]
                it['num'] = num
                items_by_num[num] = it
                updated += 1
            else:
                print(f"     Frage {num}: kein zuverlässiger Treffer (bleibt 'nicht enthalten').")
        print(f"   [{tag}] Whole-Doc-Retry: {updated} Frage(n) neu beantwortet.")

    return {
        'heading': quiz_chapter_heading,
        'questions': questions,
        'items_by_num': items_by_num,
        'heading_map': heading_map,
        'quiz_range': quiz_range,
        'kb_scope': kb_scope,
    }


def _mutate_one_quiz(doc, analysis: dict):
    """Mutations-Phase für EIN Quiz-Kapitel: frühere Tool-Antworten entfernen (Idempotenz)
    und Antworten je Frage einfügen (nur Insert). Marker werden separat und global gesetzt
    (siehe _apply_quiz_markers), nicht hier."""
    tag = analysis['heading'][:50]
    questions = analysis['questions']
    items_by_num = analysis['items_by_num']

    # Zuvor evtl. vorhandene frühere Antworten DES TOOLS entfernen (Signatur-basiert;
    # eigene manuelle Notizen des Nutzers bleiben unangetastet).
    stripped = _strip_previous_quiz_answers(doc, analysis['quiz_range'])
    if stripped:
        print(f"   [{tag}] {stripped} frühere Antwort-/Meta-Absätze ersetzt (Idempotenz).")

    inserted = 0
    for q in questions:
        item = items_by_num.get(q['num'])
        if not item:
            continue
        anchor = q['anchor']
        if anchor is None:
            # Letzte Frage am Dokumentende hat keinen nachfolgenden Absatz: einen leeren
            # Sentinel-Absatz am Ende anhängen und die Antwort davor einfügen. Sentinel wird
            # getaggt, damit ein Re-Lauf ihn mitentfernt (sonst wüchse das Dokument je Lauf).
            anchor = doc.add_paragraph()
            try:
                anchor.style = _QUIZ_ANSWER_STYLE
            except Exception:
                pass
        _insert_qa_answer_before(anchor, item)
        inserted += 1
    print(f"   [{tag}] Antworten eingefügt: {inserted}/{len(questions)}.")


def run_quiz_check(docx_path: str, quiz_chapter_headings,
                   output_path: str = None, force: bool = False) -> str:
    """End-to-End: prüft EIN ODER MEHRERE Leitfragen-Quizzes derselben Lernunterlage,
    schreibt Antworten in die jeweiligen Quiz-Kapitel und setzt Kommentar-Marker.
    Speichert EINE neue Kopie. `quiz_chapter_headings` ist ein Kapitel-String oder eine
    Liste von Kapitel-Strings."""
    src = Path(docx_path)
    if not src.exists():
        raise FileNotFoundError(f"Datei nicht gefunden: {docx_path}")
    if isinstance(quiz_chapter_headings, str):
        quiz_chapter_headings = [quiz_chapter_headings]
    if not quiz_chapter_headings:
        raise ValueError("Kein Quiz-Kapitel angegeben.")
    if output_path is None:
        output_path = str(src.with_name(f"{src.stem}_Quizgeprueft.docx"))

    print(f"--- Quiz-Prüfung: {src.name}  ({len(quiz_chapter_headings)} Quiz-Kapitel) ---")
    doc = Document(str(src))
    para_count_before = len(doc.paragraphs)
    _ensure_quiz_answer_style(doc)  # Markierungsstil für Tool-Einfügungen bereitstellen

    # Analyse-Phase: ALLE Kapitel auf dem UNVERÄNDERTEN Dokument prüfen (keine Mutation),
    # damit keine eingefügte Antwort eines Quiz die Wissensbasis eines anderen verunreinigt.
    analyses = [_analyze_one_quiz(doc, src, h, force) for h in quiz_chapter_headings]

    # Marker-Phase: auf dem NOCH UNVERÄNDERTEN Dokument (pristine Absatz-Indizes). Zuvor eigene
    # frühere „Quiz"-Kommentare entfernen (kein Mix aus alt/neu); Nutzer-/Lernfragen-Kommentare
    # bleiben unberührt. Marker sitzen an der konkreten Beleg-Stelle, nie in einem Quiz-Kapitel.
    removed_c = _remove_quiz_comments(doc)
    if removed_c:
        print(f"   {removed_c} frühere „Quiz\"-Kommentare entfernt (Idempotenz).")
    _apply_quiz_markers(doc, analyses)

    # Mutations-Phase: von unten nach oben (größter Startindex zuerst), damit die in der
    # Analyse-Phase bestimmten Absatz-Indizes gültig bleiben – Einfügungen liegen dann stets
    # UNTERHALB bereits verarbeiteter Kapitel und verschieben deren Bereiche nicht.
    for analysis in sorted(analyses, key=lambda a: a['quiz_range'][0], reverse=True):
        _mutate_one_quiz(doc, analysis)

    # Neue Kopie speichern
    doc.save(output_path)
    added = len(doc.paragraphs) - para_count_before
    print(f"Fertig: {output_path}  (+{added} Absätze eingefügt, "
          f"{len(analyses)} Quiz-Kapitel, Originalinhalt unverändert)")
    return output_path


# ---------------------------------------------------------------------------
# Hauptprogramm
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="End-to-End PDF Translation & Learning Pipeline",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=(
            "Beispiele:\n"
            "  python pipeline.py dok.pdf\n"
            "  python pipeline.py dok.pdf --parent-chapter 6.2.4.3 --title-as-parent --questions fragen.txt\n"
            "  python pipeline.py dok.pdf --chapter 4.2 --parent-chapter 2.4.1.2\n"
            "  python pipeline.py dok.pdf --no-translate --no-summary\n"
            "  python pipeline.py dok.pdf --force   # alle Zwischenergebnisse neu berechnen\n"
            "  python pipeline.py --quiz-docx unterlage.docx --quiz-chapter \"Quiz Kapitel 1 …\" \\\n"
            "                     --quiz-chapter \"Quiz Kapitel 2 …\"   # mehrere Quizzes, eine Ausgabe\n"
        ),
    )
    parser.add_argument("pdf_path", type=str, nargs="?", default=None, help="Pfad zur Quell-PDF-Datei")
    parser.add_argument("--questions", type=str, default=None, help="Pfad zu den leseleitenden Fragen (optional)")
    parser.add_argument("--quiz-docx", type=str, default=None,
                        help="Getrennter Modus: prüft ein Leitfragen-Quiz in einer bestehenden "
                             ".docx-Lernunterlage (Fragen als Screenshots) gegen den Dokumenttext, "
                             "schreibt Antworten + Abdeckungsgrad in das Quiz-Kapitel und setzt "
                             "Kommentar-Marker an den Quellkapiteln. Ausgabe: '<name>_Quizgeprueft.docx'.")
    parser.add_argument("--quiz-chapter", type=str, action="append", default=None,
                        help="PFLICHT bei --quiz-docx: Überschrift des Quiz-Kapitels (beliebige "
                             "Ebene/Position). MEHRFACH angebbar, um mehrere Quizzes derselben "
                             "Datei in EINEM Lauf/EINER Ausgabedatei zu prüfen. Der Abgleich ist "
                             "nummern-tolerant – der stabile Textteil ohne Kapitelnummer genügt, "
                             "z.B. --quiz-chapter \"Leitfragen-Quiz zum Kurs Personalauswahl\".")
    parser.add_argument("--force", action="store_true", help="Alle Schritte neu berechnen (kein Resume)")
    parser.add_argument("--parent-chapter", type=str, default=None,
                        help="Elternkapitel im Zieldokument, z.B. '4.2.1.2'. "
                             "Präfixiert alle Kapitelnummern und verschiebt Heading-Level.")
    parser.add_argument("--parent-level", type=int, default=None,
                        help="Heading-Level des Elternkapitels (Standard: automatisch aus --parent-chapter).")
    parser.add_argument("--include-references", action="store_true",
                        help="Referenzen, Literaturverzeichnis etc. einbeziehen (Standard: werden übersprungen).")
    parser.add_argument("--include-front-matter", action="store_true",
                        help="Front-Matter (Impressum, Inhaltsverzeichnis) NICHT überspringen "
                             "(Standard: wird übersprungen; Titel + Titelbild bleiben immer erhalten).")
    parser.add_argument("--chapter", type=str, default=None,
                        help="Nur dieses Kapitel extrahieren, z.B. '4.2' oder '3'. "
                             "Sucht im OCR-Markdown nach der Überschrift und extrahiert das Kapitel "
                             "inkl. aller Unterkapitel bis zur nächsten gleichrangigen Überschrift.")
    parser.add_argument("--from", dest="from_section", type=str, default=None,
                        help="Nur mit --chapter: Extrahiere das Kapitel erst AB diesem Unterkapitel "
                             "(z.B. --chapter 4 --from 4.2). Der Kapiteltitel bleibt erhalten; alles "
                             "davor (Einleitung + frühere Unterkapitel wie 4.1) wird übersprungen, "
                             "bis zum Kapitelende.")
    parser.add_argument("--to", dest="to_section", type=str, default=None,
                        help="Nur mit --chapter: Extrahiere das Kapitel nur BIS EINSCHLIESSLICH "
                             "diesem Unterkapitel (z.B. --chapter 4.6 --to 4.6.3). Dessen eigene "
                             "Unterkapitel (z.B. 4.6.3.1) bleiben erhalten; ab dem nächsten "
                             "gleichrangigen Kapitel (z.B. 4.6.4) wird abgeschnitten. Kombinierbar "
                             "mit --from.")
    parser.add_argument("--no-summary", action="store_true",
                        help="Nur Übersetzung ausgeben – keine Zusammenfassung, kein ausgeblendeter Text, "
                             "kein interleaved-Dokument. Das Übersetzungs-Docx ist das finale Ergebnis.")
    parser.add_argument("--target-language", type=str, default=None,
                        help="Zielsprache als ISO-Code, z.B. 'de' oder 'en'. "
                             "Standard: 'de' (englische Quellen werden automatisch nach Deutsch übersetzt). "
                             "Ist Quell- = Zielsprache, wird NICHT übersetzt.")
    parser.add_argument("--source-language", type=str, default=None,
                        help="Quellsprache als ISO-Code erzwingen (überspringt die automatische Erkennung).")
    parser.add_argument("--no-translate", action="store_true",
                        help="Keine Übersetzung – Originalsprache beibehalten, unabhängig von der Erkennung.")
    parser.add_argument("--title-as-parent", action="store_true",
                        help="Nur mit --parent-chapter (ohne --chapter): Der Artikeltitel trägt direkt "
                             "die Elternkapitelnummer (z.B. 6.2.4.1) statt 6.2.4.1.1; alle Abschnitte "
                             "inkl. Studien werden zu dessen Unterpunkten (6.2.4.1.1, 6.2.4.1.2 …).")
    parser.add_argument("--condensed", action="store_true",
                        help="Verdichteter Modus für nicht-prüfungsrelevante Zusatzliteratur "
                             "(z.B. Sammelbände/Festschriften mit römisch nummerierten Artikeln): "
                             "pro Artikel nur eine Kurzübersicht (Kernaussage) + 2-3 wörtliche Zitate "
                             "im ausgeblendeten Text, statt vollständigem Originaltext. "
                             "Kombinierbar mit --parent-chapter/--title-as-parent zum Einbetten.")

    args = parser.parse_args()
    OUTPUT_BASE = "workspace/output"

    try:
        if not os.getenv("GEMINI_API_KEY"):
            raise ValueError("GEMINI_API_KEY fehlt in der .env-Datei!")

        # --- Getrennter Modus: Leitfragen-Quiz-Prüfung einer bestehenden .docx ---
        if args.quiz_docx:
            if not args.quiz_chapter:
                raise ValueError(
                    "--quiz-docx benötigt --quiz-chapter (Überschrift des Quiz-Kapitels), z.B. "
                    "--quiz-chapter \"Leitfragen-Quiz zum Kurs Personalauswahl\". "
                    "Der Abgleich ist nummern-tolerant; der stabile Textteil ohne Nummer genügt."
                )
            run_quiz_check(args.quiz_docx, args.quiz_chapter, force=args.force)
            sys.exit(0)

        if not args.pdf_path:
            raise ValueError("Kein PDF angegeben. (Für den Quiz-Modus: --quiz-docx <datei>.)")

        if not os.path.exists(args.pdf_path):
            raise FileNotFoundError(f"Die Datei {args.pdf_path} wurde nicht gefunden.")

        if args.from_section and not args.chapter:
            raise ValueError("--from benötigt --chapter (z.B. --chapter 4 --from 4.2).")

        if args.to_section and not args.chapter:
            raise ValueError("--to benötigt --chapter (z.B. --chapter 4.6 --to 4.6.3).")

        if args.condensed and args.chapter:
            raise ValueError("--condensed ist nicht mit --chapter kombinierbar.")

        if args.condensed and args.questions:
            raise ValueError("--condensed ist nicht mit --questions kombinierbar.")

        pdf_stem  = Path(args.pdf_path).stem
        doc_title = pdf_stem.replace('_', ' ')
        out_dir   = Path(OUTPUT_BASE) / pdf_stem      # Dokument-Ordner: enthält die finalen .docx
        out_dir.mkdir(parents=True, exist_ok=True)

        # Arbeits-/Zwischendateien (OCR-Markdown + Bilder, Übersetzung, Zusammenfassung, QA)
        # liegen gebündelt in out_dir/work, getrennt von den finalen Ergebnissen.
        work_dir = out_dir / "work"
        work_dir.mkdir(parents=True, exist_ok=True)

        # Kapitel-spezifisches Cache-Verzeichnis innerhalb des Arbeitsordners
        chapter_safe = args.chapter.replace('.', '_') if args.chapter else None
        cache_dir = work_dir / f"kap{chapter_safe}" if chapter_safe else work_dir
        cache_dir.mkdir(parents=True, exist_ok=True)

        # --- Schritt 1: OCR ---
        # Marker schreibt nach work_dir/<pdf_stem>/<pdf_stem>.md (inkl. Bilder).
        md_path = work_dir / pdf_stem / f"{pdf_stem}.md"
        if args.force or not md_path.exists():
            raw_md_path = run_marker_ocr(args.pdf_path, str(work_dir))
            md_path = Path(raw_md_path)
        else:
            print(f"[SKIP] OCR – Markdown bereits vorhanden: {md_path}")

        raw_md = md_path.read_text(encoding="utf-8")
        # Bilder liegen im selben Ordner wie das OCR-Markdown → base_path für Bildauflösung.
        image_base = str(md_path.parent)

        # OCR-Kolumnentitel ('#### **316** | 4 Schritt 3 …') früh entfernen: solche
        # Artefakte würden sonst die Kapitel-Extraktion vorzeitig abbrechen (Seitenzahl
        # wirkt wie neues Kapitel) und als Stör-Überschrift in der Navigationsleiste landen.
        _lines_before = raw_md.count('\n') + 1
        raw_md = strip_ocr_page_headers(raw_md)
        _removed = _lines_before - (raw_md.count('\n') + 1)
        if _removed > 0:
            print(f"[OCR-CLEANUP] {_removed} Seiten-Kolumnentitel als Überschrift entfernt.")

        # Verlorene OCR-Icon-/Bullet-Glyphen ('. **Abb…**', '#### 4 **X:**', '- 4 **X**')
        # bereinigen: hält Captions sichtbar und verhindert Schein-Unterkapitel/-Nummern.
        raw_md = normalize_ocr_glyph_artifacts(raw_md)
        # Kapitel-Mini-Inhaltsverzeichnis ('| 3.1 | Titel – 40 |') als TOC-Tabelle entfernen.
        raw_md = strip_toc_tables(raw_md)

        # --- Schritt 1a: Front-Matter überspringen (Impressum, Inhaltsverzeichnis) ---
        # Standardmäßig aktiv; Titel + Titelbild bleiben erhalten. Spart außerdem
        # Übersetzungs-/Zusammenfassungs-Kosten für irrelevante Vorseiten.
        if not args.include_front_matter:
            before_len = len(raw_md)
            raw_md = strip_front_matter(raw_md)
            if len(raw_md) < before_len:
                print(f"[FRONT-MATTER] Impressum/Inhaltsverzeichnis übersprungen "
                      f"({before_len - len(raw_md)} Zeichen entfernt).")

        # --- Schritt 1b: Kapitel-Filter (optional) ---
        if args.chapter:
            raw_md = extract_chapter(raw_md, args.chapter, from_section=args.from_section,
                                     to_section=args.to_section)

        # --- Schritt 2: Sprache prüfen & ggf. übersetzen ---
        # Übersetzt wird NUR, wenn Quell- und Zielsprache verschieden sind. Default-Ziel ist
        # Deutsch: englische Quellen werden automatisch übersetzt, deutsche bleiben deutsch.
        # Eine explizite --target-language oder --source-language überschreibt die Automatik.
        transl_path = cache_dir / "de_uebersetzung.md"
        if args.force:
            if transl_path.exists():
                transl_path.unlink()
            for f in cache_dir.glob("de_uebersetzung_chunk_*.md"):
                f.unlink()
        # Legacy-Fallback: ältere Läufe legten die Übersetzung direkt in out_dir ab (vor dem
        # work/-Layout). Eine dort vorhandene, bereits verifizierte Übersetzung wiederverwenden,
        # damit der eingebettete Originaltext deutsch ist statt englisch.
        if not transl_path.exists():
            legacy_transl_path = out_dir / "de_uebersetzung.md"
            if legacy_transl_path.exists():
                transl_path = legacy_transl_path

        target_lang = (args.target_language or "de").lower()
        is_translated = False
        # content_lang = Sprache des working_text → steuert die Ausgabesprache der Zusammenfassung.
        content_lang = "de"

        if args.no_translate:
            print("--no-translate: Übersetzung übersprungen, Originalsprache bleibt erhalten.")
            working_text = raw_md
            content_lang = (args.source_language or "de").lower()
        elif transl_path.exists():
            # Bereits vorhandene Übersetzung (auch aus altem Cache-Ort) immer wiederverwenden –
            # unabhängig von der Spracherkennung. Verhindert englischen Originaltext trotz
            # existierender deutscher Übersetzung.
            print(f"[SKIP] Übersetzung – bereits vorhanden: {transl_path}")
            working_text = transl_path.read_text(encoding="utf-8")
            is_translated = True
            content_lang = target_lang
        else:
            source_lang = (args.source_language or detect_language(raw_md)).lower()
            if source_lang == "unknown":
                print("Sprache unklar – sicherheitshalber keine Übersetzung (Original bleibt erhalten).")
                working_text = raw_md
                content_lang = "de"
            elif source_lang == target_lang:
                print(f"Quelle ist bereits {_language_name(target_lang)}. Keine Übersetzung notwendig.")
                working_text = raw_md
                content_lang = target_lang
            else:
                print(f"Quelle ist {_language_name(source_lang)}. Starte Übersetzung nach "
                      f"{_language_name(target_lang)}...")
                working_text = translate_text(raw_md, source_lang=source_lang, target_lang=target_lang, cache_dir=cache_dir)
                transl_path.write_text(working_text, encoding="utf-8")
                print(f"       Übersetzung gespeichert: {transl_path}")
                is_translated = True
                content_lang = target_lang

        # --- Zwischenschritt: Übersetzungs-Docx (nur bei englischer Quelle) ---
        kap_infix = f"_kap{chapter_safe}" if chapter_safe else ""
        if is_translated:
            transl_docx_path = out_dir / f"{pdf_stem}{kap_infix}_Uebersetzung.docx"
            if args.force or not transl_docx_path.exists():
                # Auch das Übersetzungs-Docx von Publikations-Metadaten bereinigen und
                # flachgewalzte Tabellen als Bild einfügen (lokale Kopie, ohne den Cache zu ändern).
                transl_text = strip_publication_metadata(working_text)
                transl_text = inline_flattened_tables(transl_text, args.pdf_path, image_base)
                build_translation_word_document(transl_text, str(transl_docx_path), base_path=image_base)
            else:
                print(f"[SKIP] Übersetzungs-Docx – bereits vorhanden: {transl_docx_path}")

            if args.no_summary:
                print(f"\n=== PIPELINE ERFOLGREICH BEENDET (nur Übersetzung) ===")
                print(f"Zwischenergebnisse: {cache_dir}")
                print(f"Fertiges Dokument:  {transl_docx_path}")
                sys.exit(0)
        elif args.no_summary:
            print(f"\n=== PIPELINE ERFOLGREICH BEENDET (Text bereits Deutsch – keine Übersetzung erstellt) ===")
            print(f"Zwischenergebnisse: {cache_dir}")
            sys.exit(0)

        # --- Schritt 3: Box-Strukturreparatur (nur wenn Lehrbuch-Kästen vorhanden) ---
        # Trennt eingeschobene Kästen (Fokus/Studie/Definition/Beispiel) vom Fließtext,
        # den die OCR fälschlich in den Kasten gezogen hat. Speist Summary + Interleaved.
        struct_path = cache_dir / "de_strukturiert.md"
        if args.force and struct_path.exists():
            struct_path.unlink()
        if struct_path.exists():
            print(f"[SKIP] Box-Strukturreparatur – bereits vorhanden: {struct_path}")
            working_text = struct_path.read_text(encoding="utf-8")
        else:
            print("--- Schritt 3: Box-Strukturreparatur (Kästen vom Fließtext trennen) ---")
            working_text = repair_box_structure(working_text)
            # Front-Matter entdoppeln (Journal-Banner + doppelter Titel) und die unbeschriftete
            # Artikel-Einleitung mit eigener Überschrift versehen, damit sie zusammengefasst wird.
            working_text = collapse_duplicate_title(working_text)
            working_text = insert_intro_heading(working_text)
            struct_path.write_text(working_text, encoding="utf-8")

        # --- Verdichteter Modus (--condensed): Kurzübersicht je Artikel statt vollständiger,
        # kapitelweiser Zusammenfassung. Umgeht Schritt 4/5 (Zusammenfassung/QA) sowie den
        # interleaved Builder vollständig - eigener, einfacherer Docx-Aufbau ohne vollen
        # Originaltext (bewusste Ausnahme von der 100%-Erhaltungsregel, siehe Epic 6).
        if args.condensed:
            cond_path = cache_dir / "verdichtung.md"
            if args.force:
                if cond_path.exists():
                    cond_path.unlink()
                for f in cache_dir.glob("verdichtung_artikel_*.md"):
                    f.unlink()
            condensed_result = load_or_run(
                cond_path,
                lambda: generate_condensed_summary(working_text, cache_dir, output_lang=content_lang),
                "Verdichtete Zusammenfassung"
            )
            cond_suffix = f"_Verdichtet_Einbetten_{args.parent_chapter.replace('.', '-')}" if args.parent_chapter else "_Verdichtet"
            cond_docx_path = out_dir / f"{pdf_stem}{cond_suffix}.docx"
            build_condensed_word_document(
                condensed_result, str(cond_docx_path), base_path=image_base,
                parent_chapter=args.parent_chapter, parent_level=args.parent_level,
                doc_title=doc_title,
            )
            print(f"\n=== PIPELINE ERFOLGREICH BEENDET (verdichteter Modus) ===")
            print(f"Zwischenergebnisse: {cache_dir}")
            print(f"Fertiges Dokument:  {cond_docx_path}")
            sys.exit(0)

        # --- Schritt 4: Zusammenfassung (kapitelweise) ---
        sum_path = cache_dir / "zusammenfassung.md"
        if args.force and sum_path.exists():
            sum_path.unlink()
            for f in cache_dir.glob("zusammenfassung_kap_*.md"):
                f.unlink()

        if sum_path.exists():
            print(f"[SKIP] Zusammenfassung – bereits vorhanden: {sum_path}")
            summary_result = sum_path.read_text(encoding="utf-8")
        else:
            summary_result = generate_summary_by_chapter(working_text, cache_dir, output_lang=content_lang)

        # --- Schritt 5: Qualitätssicherung (optional) ---
        qa_result = "Keine Leitfragen zur Prüfung übergeben."
        supplement_map: dict = {}
        if args.questions:
            if os.path.exists(args.questions):
                qa_path = cache_dir / "qa_ergebnis.md"
                if args.force and qa_path.exists():
                    qa_path.unlink()
                qa_result = load_or_run(
                    qa_path,
                    lambda: verify_with_questions(summary_result, args.questions),
                    "QA / Leitfragen"
                )

                # --- Schritt 5b: Nacharbeit unvollständiger Antworten ---
                rework_path = cache_dir / "qa_nachbearbeitung.md"
                sup_path    = cache_dir / "qa_ergaenzungen.json"
                if args.force:
                    for p in (rework_path, sup_path):
                        if p.exists():
                            p.unlink()
                if rework_path.exists() and sup_path.exists():
                    print(f"[SKIP] Nacharbeit – bereits vorhanden: {rework_path}")
                    qa_result = rework_path.read_text(encoding="utf-8")
                    supplement_map = json.loads(sup_path.read_text(encoding="utf-8"))
                else:
                    qa_result, supplement_map = rework_partial_answers(
                        qa_result, working_text, args.questions
                    )
                    rework_path.write_text(qa_result, encoding="utf-8")
                    sup_path.write_text(json.dumps(supplement_map, ensure_ascii=False, indent=2),
                                        encoding="utf-8")
            else:
                print(f"Warnung: Fragen-Datei '{args.questions}' nicht gefunden. Überspringe QS.")

        # Publikations-Metadaten (Copyright/ISSN/DOI/Online-First) entfernen, die die OCR
        # zwischen den Inhalt streut und die den Lesefluss stören.
        working_text = strip_publication_metadata(working_text)
        # Zerstückelte Abbildungen (Marker hat eine Grafik in Text + Icons zerlegt) als ganze
        # Figur aus der Quell-PDF rendern und die scattered Icon-Bilder ersetzen.
        working_text = inline_decomposed_figures(working_text, args.pdf_path, image_base)
        # Flachgewalzte Tabellen (von der OCR nicht als Pipe-Tabelle erkannt) als Bild der
        # PDF-Seite einfügen.
        working_text = inline_flattened_tables(working_text, args.pdf_path, image_base)

        # --- Word-Dokument zusammensetzen ---
        suffix = f"_Einbetten_{args.parent_chapter.replace('.', '-')}" if args.parent_chapter else "_Lernskript"
        final_docx_path = out_dir / f"{pdf_stem}{kap_infix}{suffix}.docx"
        build_interleaved_word_document(
            working_text, summary_result, qa_result,
            str(final_docx_path), base_path=image_base,
            parent_chapter=args.parent_chapter,
            parent_level=args.parent_level,
            skip_references=not args.include_references,
            questions_path=args.questions,
            doc_title=doc_title,
            extracted_chapter=args.chapter,
            supplement_map=supplement_map,
            title_as_parent=args.title_as_parent,
        )

        print(f"\n=== PIPELINE ERFOLGREICH BEENDET ===")
        print(f"Zwischenergebnisse:   {cache_dir}")
        if is_translated:
            print(f"Übersetzung (Word):   {transl_docx_path}")
        print(f"Fertiges Dokument:    {final_docx_path}")
        if args.parent_chapter:
            print(f"  → Einfügemodus: Kapitelpräfix '{args.parent_chapter}', "
                  f"Heading-Shift +{args.parent_level or _parent_level_from_chapter(args.parent_chapter)}")

    except Exception as e:
        print(f"\nPipeline abgebrochen wegen: {e}")
